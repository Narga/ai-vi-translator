# -*- coding: utf-8 -*-

"""
OCR reader module: extract text from scanned PDFs or images based on settings in config/config.yaml.

Dependencies:
- pytesseract (Python wrapper for Tesseract OCR)
- pdf2image (convert PDF pages to images)
- Pillow (image processing)
- PyYAML (read YAML config)

Config example in config/config.yaml:
  ocr:
    enabled: true
    tesseract_cmd: "C:/Program Files/Tesseract-OCR/tesseract.exe"
    lang: "vie+eng"
    psm: 3
    dpi: 300
"""

import os
import sys
import subprocess
import logging
import time
import asyncio
import gc
from typing import List, Optional
from pathlib import Path

# Suppress noisy logs từ Google libraries (absl, gRPC) trước khi import
os.environ['GRPC_VERBOSITY'] = 'ERROR'
os.environ['GLOG_minloglevel'] = '2'  # Suppress INFO và WARNING từ GLOG/absl
os.environ['GRPC_PYTHON_LOG_LEVEL'] = 'ERROR'
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '2'  # Suppress TensorFlow logs nếu có

# Define StderrFilter trước để filter stderr ngay từ đầu
class NoisyMessageFilter:
    """Filter để chặn các messages gây nhiễu được in trực tiếp ra stderr/stdout."""
    def __init__(self, original_stream):
        self.original_stream = original_stream
        # Buffer để xử lý multi-line messages
        self.buffer = ""
        # Patterns gây nhiễu (tất cả lowercase để so sánh)
        # Bao gồm cả partial matches để catch variations
        self.noisy_patterns = [
            'e0000',  # gRPC error prefix
            'alts_credentials',
            'alts creds',
            'alts creds ignored',
            'alts creds ignored. not running on gcp',
            'absl::initializelog',
            'not running on gcp',
            'untrusted alts',
            'untrusted alts is not enabled',
            'written to stderr',
            'all log messages before',
            'all log messages before absl',
            'alts_credentials.cc',  # File path pattern
            'alts_credentials.cc:93',  # File path with line number
            'warning: all log messages'  # Warning prefix
        ]
    
    def write(self, text):
        if not text:
            return
        
        # Thêm vào buffer để xử lý multi-line messages
        self.buffer += text
        
        # Kiểm tra buffer có chứa noisy patterns không (case-insensitive)
        buffer_lower = self.buffer.lower()
        
        # Kiểm tra nhanh trước khi split (tối ưu hơn)
        is_noisy = False
        
        # Check toàn bộ buffer trước (faster)
        for pattern in self.noisy_patterns:
            if pattern in buffer_lower:
                is_noisy = True
                break
        
        # Nếu chưa detect, check từng dòng chi tiết
        if not is_noisy:
            lines = self.buffer.split('\n')
            for line in lines:
                line_lower = line.lower().strip()
                # Check các pattern cụ thể
                if any(pattern in line_lower for pattern in self.noisy_patterns):
                    is_noisy = True
                    break
                # Check pattern E0000 ở đầu dòng
                if line_lower.startswith('e0000'):
                    is_noisy = True
                    break
                # Check "WARNING:" prefix với absl messages
                if line_lower.startswith('warning:') and ('absl' in line_lower or 'stderr' in line_lower):
                    is_noisy = True
                    break
        
        # Nếu không noisy, ghi ra stream
        if not is_noisy:
            self.original_stream.write(text)
        # Nếu noisy, không ghi gì cả (suppress hoàn toàn)
        
        # Reset buffer sau mỗi newline hoặc khi buffer quá dài
        if '\n' in text:
            # Giữ lại phần sau newline cuối cùng để check tiếp (cho multi-line messages)
            parts = self.buffer.rsplit('\n', 1)
            self.buffer = parts[-1] if len(parts) > 1 else ""
        
        if len(self.buffer) > 2000:  # Reset nếu buffer quá dài
            self.buffer = ""
    
    def flush(self):
        self.original_stream.flush()
    
    def __getattr__(self, name):
        return getattr(self.original_stream, name)


# Alias cho backward compatibility
StderrFilter = NoisyMessageFilter

# Suppress warnings từ absl trước khi import google libraries
try:
    import absl.logging
    absl.logging.set_verbosity(absl.logging.ERROR)
except Exception:
    pass

# Suppress logging từ các Google libraries
for lib_name in ['google', 'grpc', 'absl', 'google.generativeai', 'google.api_core', 'google.auth']:
    lib_logger = logging.getLogger(lib_name)
    lib_logger.setLevel(logging.ERROR)
    lib_logger.propagate = False

# Filter stderr ngay từ đầu để chặn messages in trực tiếp
# (không filter stdout ở đây để không ảnh hưởng đến user interaction)
_stderr_filter_active = False
_stdout_filter_active = False
try:
    original_stderr = sys.stderr
    sys.stderr = NoisyMessageFilter(original_stderr)
    _stderr_filter_active = True
except Exception:
    pass

# Lazy import yaml to avoid import errors
yaml = None
try:
    import yaml as _yaml
    yaml = _yaml
except ImportError:
    pass  # Will be imported later via lazy_import_and_install if needed

try:
    from PIL import Image
except Exception:
    Image = None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from pdf2image import convert_from_path
except Exception:
    convert_from_path = None

try:
    from tqdm import tqdm  # progress bar (optional)
except Exception:
    tqdm = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except Exception:
    Document = None

# pdf2docx và ocrmypdf sẽ được import lazy trong _ensure_dependencies để tránh treo khi load module
Converter = None

# OCRmyPDF (for fallback when pdf2docx fails)
ocrmypdf_available = False
ocrmypdf = None

logger = logging.getLogger("NovelTranslator")


class GoogleLogFilter(logging.Filter):
    """
    Filter để loại bỏ các log messages gây nhiễu từ Google libraries.
    """
    def filter(self, record):
        msg = str(record.getMessage())
        msg_lower = msg.lower()
        
        # Loại bỏ các messages về absl::InitializeLog
        if 'absl::initializelog' in msg_lower or 'absl::InitializeLog' in msg:
            return False
        
        # Loại bỏ các messages về ALTS creds (nhiều pattern khác nhau)
        if any(pattern in msg for pattern in [
            'ALTS creds',
            'alts_credentials',
            'alts creds ignored',
            'not running on gcp',
            'untrusted alts is not enabled'
        ]):
            return False
        
        # Loại bỏ các messages từ absl logger
        if record.name.startswith('absl.') or 'absl' in record.name.lower():
            return False
        
        # Loại bỏ messages có pattern E0000 từ gRPC/absl
        if msg.startswith('E0000') and ('alts' in msg_lower or 'cred' in msg_lower):
            return False
        
        return True


_stderr_filter_active = False


def _suppress_google_logs():
    """
    Suppress logging từ Google libraries (gRPC, absl, etc.)
    Bao gồm cả việc filter stderr trực tiếp.
    """
    global _stderr_filter_active
    
    # Set environment variables
    os.environ['GRPC_VERBOSITY'] = 'ERROR'
    os.environ['GLOG_minloglevel'] = '2'
    os.environ['GRPC_PYTHON_LOG_LEVEL'] = 'ERROR'
    os.environ['TF_CPP_MIN_LOG_LEVEL'] = '2'
    
    # Suppress absl logging nếu có
    try:
        import absl.logging
        absl.logging.set_verbosity(absl.logging.ERROR)
        # Disable absl handler
        for handler in absl.logging._absl_logger.handlers:
            handler.setLevel(logging.ERROR)
    except Exception:
        pass
    
    # Filter stderr và stdout để chặn messages in trực tiếp (re-apply để đảm bảo)
    global _stderr_filter_active, _stdout_filter_active
    try:
        # Kiểm tra nếu đã là NoisyMessageFilter rồi thì không cần apply lại
        if not isinstance(sys.stderr, NoisyMessageFilter):
            original_stderr = sys.stderr if not isinstance(sys.stderr, NoisyMessageFilter) else sys.stderr.original_stream
            sys.stderr = NoisyMessageFilter(original_stderr)
            _stderr_filter_active = True
    except Exception:
        pass
    
    try:
        if not isinstance(sys.stdout, NoisyMessageFilter):
            original_stdout = sys.stdout if not isinstance(sys.stdout, NoisyMessageFilter) else sys.stdout.original_stream
            sys.stdout = NoisyMessageFilter(original_stdout)
            _stdout_filter_active = True
    except Exception:
        pass
    
    # Apply filter cho root logger và các loggers cụ thể
    google_filter = GoogleLogFilter()
    root_logger = logging.getLogger()
    root_logger.addFilter(google_filter)
    
    # Suppress logging từ các Google libraries
    for lib_name in ['google', 'grpc', 'absl', 'google.generativeai', 'google.api_core', 'google.auth', 'grpc._cython']:
        lib_logger = logging.getLogger(lib_name)
        lib_logger.setLevel(logging.ERROR)
        lib_logger.propagate = False
        lib_logger.addFilter(google_filter)


def _parse_pages(pages_str: str) -> Optional[List[int]]:
    """
    Parse chuỗi pages thành danh sách số trang.
    Hỗ trợ:
    - "1,2,5,7" → [1, 2, 5, 7]
    - "1-7" → [1, 2, 3, 4, 5, 6, 7]
    - "1-3,5,7-9" → [1, 2, 3, 5, 7, 8, 9]
    
    Returns: List[int] hoặc None nếu không hợp lệ
    """
    if not pages_str or not pages_str.strip():
        return None
    
    pages_str = pages_str.strip()
    # Loại bỏ dấu ngoặc vuông hoặc ngoặc tròn nếu có (để tương thích ngược)
    if (pages_str.startswith('[') and pages_str.endswith(']')) or \
       (pages_str.startswith('(') and pages_str.endswith(')')):
        pages_str = pages_str[1:-1].strip()
    
    pages: List[int] = []
    parts = [p.strip() for p in pages_str.split(',')]
    
    for part in parts:
        part = part.strip()
        if not part:
            continue
        
        # Kiểm tra có phải range không (ví dụ: "1-7")
        if '-' in part:
            try:
                start, end = part.split('-', 1)
                start = int(start.strip())
                end = int(end.strip())
                if start > end:
                    logger.warning(f"Range không hợp lệ: {part} (start > end). Bỏ qua.")
                    continue
                pages.extend(range(start, end + 1))
            except ValueError:
                logger.warning(f"Range không hợp lệ: {part}. Bỏ qua.")
                continue
        else:
            # Số trang đơn lẻ
            try:
                page_num = int(part)
                if page_num > 0:
                    pages.append(page_num)
            except ValueError:
                logger.warning(f"Số trang không hợp lệ: {part}. Bỏ qua.")
                continue
    
    # Loại bỏ trùng lặp và sắp xếp
    pages = sorted(list(set(pages)))
    
    if not pages:
        logger.warning("Không có trang hợp lệ nào được parse.")
        return None
    
    return pages


def _ensure_logger_config() -> None:
    """Đảm bảo logger có handler để in ra console và lưu file khi chạy trực tiếp.
    Tránh tình trạng không thấy log do thiếu cấu hình bên ngoài.
    """
    if getattr(_ensure_logger_config, "_configured", False):
        return
    
    # Suppress Google logs trước khi cấu hình logger
    _suppress_google_logs()
    
    logger.setLevel(logging.INFO)
    logger.propagate = False
    # Kiểm tra có StreamHandler/FileHandler chưa
    has_stream = any(isinstance(h, logging.StreamHandler) for h in logger.handlers)
    has_file = any(isinstance(h, logging.FileHandler) for h in logger.handlers)
    # Console handler
    if not has_stream:
        ch = logging.StreamHandler()
        ch.setLevel(logging.INFO)
        ch.setFormatter(logging.Formatter("%(asctime)s - %(levelname)s - %(message)s"))
        ch.addFilter(GoogleLogFilter())  # Apply filter để loại bỏ Google logs
        logger.addHandler(ch)
    # File handler
    if not has_file:
        try:
            os.makedirs("logs", exist_ok=True)
            fh = logging.FileHandler(os.path.join("logs", "ocr_runtime.log"), encoding="utf-8")
            fh.setLevel(logging.INFO)
            fh.setFormatter(logging.Formatter("%(asctime)s - %(levelname)s - %(message)s"))
            fh.addFilter(GoogleLogFilter())  # Apply filter để loại bỏ Google logs
            logger.addHandler(fh)
        except Exception:
            pass
    setattr(_ensure_logger_config, "_configured", True)


def _load_yaml(path: str) -> dict:
    global yaml
    if yaml is None:
        try:
            import yaml as _yaml
            yaml = _yaml
        except ImportError:
            raise ImportError("PyYAML is not installed. Please install: pip install PyYAML")
    
    if not os.path.exists(path):
        raise FileNotFoundError(path)
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f) or {}



def load_ocr_config(config_path: str = "config/config.yaml") -> dict:
    cfg = _load_yaml(config_path)
    ocr_cfg = cfg.get("ocr") or {}
    # Lưu config_path để dùng sau
    ocr_cfg["_config_path"] = config_path
    # Lưu api_keys từ root config để dùng cho AI cleanup
    ocr_cfg["_root_api_keys"] = cfg.get("api_keys", [])
    # Lưu safety_level từ root config (nếu có) để dùng cho AI cleanup/spell check
    # Ưu tiên: ocr.safety_level > root safety_level > default BLOCK_ONLY_HIGH
    if "safety_level" not in ocr_cfg:
        ocr_cfg["safety_level"] = cfg.get("safety_level", "BLOCK_ONLY_HIGH")
    return ocr_cfg


def _build_safety_settings(safety_level: str = "BLOCK_ONLY_HIGH") -> List[dict]:
    """
    Tạo safety settings cho Google Gemini API.
    Học hỏi từ module dịch thuật (model_router.py).
    
    Args:
        safety_level: Safety level từ config (BLOCK_NONE, BLOCK_ONLY_HIGH, BLOCK_MEDIUM_AND_ABOVE, BLOCK_LOW_AND_ABOVE)
    
    Returns:
        List of safety settings dicts cho GenerativeModel
    """
    safety_level = safety_level.upper() if safety_level else "BLOCK_ONLY_HIGH"
    
    # Các levels hợp lệ từ Google Gemini API
    valid_levels = ["BLOCK_NONE", "BLOCK_ONLY_HIGH", "BLOCK_MEDIUM_AND_ABOVE", "BLOCK_LOW_AND_ABOVE"]
    if safety_level not in valid_levels:
        logger.warning(f"Safety level '{safety_level}' không hợp lệ. Dùng default: BLOCK_ONLY_HIGH")
        safety_level = "BLOCK_ONLY_HIGH"
    
    return [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": safety_level},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": safety_level},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": safety_level},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": safety_level},
    ]


def _bundle_base_dir() -> str:
    """Return base dir for bundled resources (PyInstaller) or script dir."""
    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        return getattr(sys, '_MEIPASS')  # PyInstaller temp dir
    # Fallback: repo/script directory
    return os.path.dirname(os.path.abspath(sys.argv[0]))


def _detect_bundled_binaries(ocr_cfg: dict) -> dict:
    """
    If config values are missing, try to detect bundled Tesseract and Poppler paths.
    - Looks for vendor/tesseract/tesseract.exe
    - Looks for vendor/poppler/bin
    """
    cfg = dict(ocr_cfg) if ocr_cfg else {}
    base = _bundle_base_dir()
    # Detect Tesseract
    if not cfg.get('tesseract_cmd'):
        cand = os.path.join(base, 'tesseract', 'tesseract.exe')
        if not os.path.exists(cand):
            cand = os.path.join(base, 'vendor', 'tesseract', 'tesseract.exe')
        if os.path.exists(cand):
            cfg['tesseract_cmd'] = cand.replace('\\', '/')
    # Detect Poppler bin
    if not cfg.get('poppler_path'):
        # 1) Bundled relative paths
        cand_dir = os.path.join(base, 'poppler', 'bin')
        if not os.path.isdir(cand_dir):
            cand_dir = os.path.join(base, 'vendor', 'poppler', 'bin')
        if os.path.isdir(cand_dir):
            cfg['poppler_path'] = cand_dir.replace('\\', '/')
        else:
            # 2) Environment variables
            env_poppler = os.environ.get('POPPLER_PATH') or os.environ.get('POPPLER_BIN')
            if env_poppler and os.path.isdir(env_poppler):
                cfg['poppler_path'] = env_poppler.replace('\\', '/')
            else:
                # 3) Try to find pdftoppm from PATH
                try:
                    import shutil
                    pdftoppm = shutil.which('pdftoppm')
                    if pdftoppm:
                        cfg['poppler_path'] = os.path.dirname(pdftoppm).replace('\\', '/')
                except Exception:
                    pass
    return cfg


# Inline lazy import implementation (no dependency on src.utils.helpers)
def lazy_import_and_install(package_name: str, import_name: str = None, version_spec: str = ""):
    """
    Lazy import a package, installing it if not available.
    
    Args:
        package_name: Package name to install (e.g., 'Pillow')
        import_name: Module name to import (e.g., 'PIL'), defaults to package_name
        version_spec: Version specification (e.g., '>=10.0.0')
    
    Returns:
        Imported module
    """
    if import_name is None:
        import_name = package_name
    
    try:
        # Try to import
        import importlib
        return importlib.import_module(import_name)
    except ImportError:
        # Install if not available
        logger.info(f"Installing {package_name}{version_spec}...")
        try:
            subprocess.check_call([
                sys.executable, '-m', 'pip', 'install', '-q',
                f"{package_name}{version_spec}" if version_spec else package_name
            ])
            import importlib
            return importlib.import_module(import_name)
        except Exception as e:
            raise ImportError(f"Failed to install {package_name}: {e}")

def _pip_install(package: str) -> None:
    # Giữ hàm để tương thích ngược nếu có nơi khác gọi
    try:
        lazy_import_and_install(package)
    except Exception as e:
        raise RuntimeError(f"Không thể cài gói '{package}': {e}")



def _ensure_dependencies(ocr_cfg: dict) -> None:
    global Image, pytesseract, convert_from_path, tqdm, PyPDF2, pdfplumber, fitz, Document, Inches, Pt, WD_PARAGRAPH_ALIGNMENT
    # Pillow
    if Image is None:
        # Import trực tiếp submodule PIL.Image để tránh getattr trên package gốc
        _pil_image_mod = lazy_import_and_install("Pillow", "PIL.Image", ">=10.0.0")
        Image = getattr(_pil_image_mod, 'Image', _pil_image_mod)
    # pdf2image
    if convert_from_path is None:
        pdf2image = lazy_import_and_install("pdf2image", "pdf2image", ">=1.17.0")
        convert_from_path = getattr(pdf2image, 'convert_from_path')
    # pytesseract (runtime still needs system Tesseract)
    if pytesseract is None:
        pytesseract = lazy_import_and_install("pytesseract", "pytesseract", ">=0.3.10")
    # pdfplumber (preferred for text extraction)
    if pdfplumber is None:
        try:
            pdfplumber = lazy_import_and_install("pdfplumber", "pdfplumber", ">=0.9.0")
        except Exception:
            pass
    # PyPDF2 (fallback for text extraction)
    if PyPDF2 is None:
        try:
            PyPDF2 = lazy_import_and_install("PyPDF2", "PyPDF2", ">=3.0.0")
        except Exception:
            pass
    # yaml (PyYAML) was already imported to read config; skip
    # tqdm optional
    if bool(ocr_cfg.get("show_progress", True)) and tqdm is None:
        try:
            _tqdm_mod = lazy_import_and_install("tqdm", "tqdm", ">=4.65.0")
            from tqdm import tqdm as _tqdm
            tqdm = _tqdm
        except Exception:
            tqdm = None
    # PyMuPDF (for extracting images from PDF)
    # Limited to <1.26.5 for pdf2docx compatibility (1.26.5+ removed get_area method)
    if fitz is None:
        try:
            fitz = lazy_import_and_install("PyMuPDF", "fitz", ">=1.23.0,<1.26.5")
        except Exception:
            try:
                fitz = lazy_import_and_install("PyMuPDF==1.26.4", "fitz")
            except Exception:
                fitz = None
    # python-docx (for creating DOCX output)
    if Document is None:
        try:
            docx_mod = lazy_import_and_install("python-docx", "docx", ">=1.0.0")
            from docx import Document as _Document
            from docx.shared import Inches as _Inches, Pt as _Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as _WD_PARAGRAPH_ALIGNMENT
            Document = _Document
            global Inches, Pt, WD_PARAGRAPH_ALIGNMENT
            Inches = _Inches
            Pt = _Pt
            WD_PARAGRAPH_ALIGNMENT = _WD_PARAGRAPH_ALIGNMENT
        except Exception:
            Document = None
            Inches = None
            Pt = None
            WD_PARAGRAPH_ALIGNMENT = None
    # pdf2docx (for hybrid workflow: PDF → DOCX conversion)
    global Converter
    if Converter is None:
        try:
            _pdf2docx = lazy_import_and_install("pdf2docx", "pdf2docx", ">=0.5.0")
            from pdf2docx import Converter as _Converter
            Converter = _Converter
        except Exception:
            Converter = None
    
    # OCRmyPDF (for fallback when pdf2docx fails)
    global ocrmypdf_available
    if not ocrmypdf_available:
        try:
            _ocrmypdf = lazy_import_and_install("ocrmypdf", "ocrmypdf", ">=15.0.0")
            global ocrmypdf
            ocrmypdf = _ocrmypdf
            ocrmypdf_available = True
        except Exception:
            ocrmypdf_available = False
            ocrmypdf = None
            logger.warning("⚠️  OCRmyPDF không khả dụng. Fallback sẽ không hoạt động.")


def _apply_tesseract_cfg(ocr_cfg: dict) -> None:
    if pytesseract is None:
        raise RuntimeError("pytesseract not installed. Please install pytesseract and system Tesseract.")
    # Allow auto-detect of bundled binaries
    _cfg = _detect_bundled_binaries(ocr_cfg)
    tesseract_cmd = _cfg.get("tesseract_cmd")
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd


def _normalize_lang_code(lang: str) -> str:
    """
    Chuyển đổi mã ngôn ngữ từ format ngắn (VN, EN, CN) sang Tesseract format (vie, eng, chi).
    Hỗ trợ backward compatibility với format cũ.
    
    Args:
        lang: Language string có thể là "VN", "EN", "CN", "auto", hoặc format cũ "vie", "eng", "chi"
    
    Returns:
        Tesseract language code hoặc "auto"
    """
    if not lang:
        return "vie"
    
    lang = lang.strip().upper()
    
    # Mapping từ format ngắn sang Tesseract
    lang_map = {
        "VN": "vie",
        "EN": "eng", 
        "CN": "chi",
        "VIE": "vie",  # Backward compatibility
        "ENG": "eng",  # Backward compatibility
        "CHI": "chi",  # Backward compatibility
        "AUTO": "auto"
    }
    
    # Xử lý kết hợp ngôn ngữ (VD: "VN+EN" hoặc "vie+eng")
    if "+" in lang:
        parts = lang.split("+")
        normalized_parts = []
        for part in parts:
            part = part.strip().upper()
            normalized = lang_map.get(part, part.lower())  # Fallback về lowercase nếu không map được
            normalized_parts.append(normalized)
        return "+".join(normalized_parts)
    
    # Xử lý single language
    return lang_map.get(lang, lang.lower())  # Fallback về lowercase nếu không map được


def _get_exif_rotation_degrees(img: "Image.Image") -> int:
    """Đọc EXIF Orientation nếu có và trả về góc xoay cần thiết (0/90/180/270)."""
    try:
        if hasattr(img, "_getexif") and callable(getattr(img, "_getexif")):
            exif = img._getexif() or {}
            orientation = exif.get(274)  # EXIF Orientation tag
            mapping = {3: 180, 6: 270, 8: 90}  # cần xoay để hiển thị đúng
            return mapping.get(orientation, 0)
    except Exception:
        pass
    return 0


def _detect_orientation_degrees_osd(img: "Image.Image") -> int:
    """Dùng Tesseract OSD để phát hiện góc xoay. Trả về 0/90/180/270."""
    try:
        if pytesseract is None:
            return 0
        osd = pytesseract.image_to_osd(img)
        # OSD text chứa dòng: "Rotate: 90"
        for line in str(osd).splitlines():
            line = line.strip()
            if line.lower().startswith("rotate:"):
                val = line.split(":", 1)[1].strip()
                deg = int("".join(ch for ch in val if ch.isdigit()))
                if deg in (0, 90, 180, 270):
                    return deg
                break
    except Exception:
        pass
    return 0


def _auto_rotate_image(img: "Image.Image", ocr_cfg: dict) -> "Image.Image":
    """Tự động xoay ảnh dựa trên EXIF và OSD. Ưu tiên EXIF, sau đó OSD nếu cần.
    - ocr.auto_rotate_exif: bật/tắt xoay theo EXIF (default True)
    - ocr.auto_rotate_osd: bật/tắt xoay theo OSD (default True)
    """
    if not isinstance(img, Image.Image):
        return img
    auto_exif = bool(ocr_cfg.get("auto_rotate_exif", True))
    auto_osd = bool(ocr_cfg.get("auto_rotate_osd", True))

    rotated = False
    try:
        if auto_exif:
            deg = _get_exif_rotation_degrees(img)
            if deg in (90, 180, 270):
                # PIL.rotate: xoay ngược chiều kim đồng hồ; EXIF deg là cần xoay thuận để đúng hướng
                img = img.rotate(360 - deg, expand=True)
                rotated = True
        if auto_osd:
            deg_osd = _detect_orientation_degrees_osd(img)
            # Nếu OSD báo phải xoay, thực hiện xoay để chữ nằm thẳng đứng
            if deg_osd in (90, 180, 270):
                img = img.rotate(360 - deg_osd, expand=True)
                rotated = True
    except Exception:
        return img
    return img


def _is_cjk_character(char: str) -> bool:
    """
    Kiểm tra xem ký tự có phải là CJK (Chinese, Japanese, Korean) không.
    Dựa trên Unicode ranges cho CJK.
    """
    if not char:
        return False
    code = ord(char)
    # CJK Unified Ideographs: U+4E00–U+9FFF
    # CJK Extension A: U+3400–U+4DBF
    # CJK Extension B: U+20000–U+2A6DF
    # CJK Compatibility: U+F900–U+FAFF
    return (
        0x4E00 <= code <= 0x9FFF or  # CJK Unified Ideographs
        0x3400 <= code <= 0x4DBF or  # CJK Extension A
        0xF900 <= code <= 0xFAFF     # CJK Compatibility
    )


def _count_cjk_characters(text: str) -> int:
    """Đếm số ký tự CJK trong text."""
    return sum(1 for char in text if _is_cjk_character(char))


# DEPRECATED FUNCTIONS REMOVED:
# - _detect_language_from_image (removed - không còn được sử dụng)
# - _detect_language_from_multiple_pages (removed - không còn được sử dụng)


def _detect_chinese_variant(img: "Image.Image", ocr_cfg: dict) -> str:
    """
    Tự động nhận biết tiếng Trung giản thể hay phồn thể.
    Returns: "chi_sim" hoặc "chi_tra"
    """
    psm = int(ocr_cfg.get("psm", 3) or 3)
    config = f"--psm {psm}"
    
    try:
        # OCR với cả 2 ngôn ngữ và so sánh confidence
        # Simplified Chinese
        data_sim = pytesseract.image_to_data(img, lang="chi_sim", config=config, output_type=pytesseract.Output.DICT)
        confidences_sim = [int(conf) for conf in data_sim['conf'] if int(conf) > 0]
        avg_conf_sim = sum(confidences_sim) / len(confidences_sim) if confidences_sim else 0
        # Đếm số ký tự được nhận dạng (có confidence > 0)
        char_count_sim = sum(1 for i, text_item in enumerate(data_sim['text']) if text_item.strip() and int(data_sim['conf'][i]) > 0)
        
        # Traditional Chinese
        data_tra = pytesseract.image_to_data(img, lang="chi_tra", config=config, output_type=pytesseract.Output.DICT)
        confidences_tra = [int(conf) for conf in data_tra['conf'] if int(conf) > 0]
        avg_conf_tra = sum(confidences_tra) / len(confidences_tra) if confidences_tra else 0
        # Đếm số ký tự được nhận dạng (có confidence > 0)
        char_count_tra = sum(1 for i, text_item in enumerate(data_tra['text']) if text_item.strip() and int(data_tra['conf'][i]) > 0)
        
        # Quyết định dựa trên confidence và số ký tự
        # Ưu tiên confidence, nếu gần bằng nhau thì ưu tiên số ký tự nhiều hơn
        score_sim = avg_conf_sim * 0.7 + (char_count_sim / max(char_count_sim + char_count_tra, 1)) * 30 * 0.3
        score_tra = avg_conf_tra * 0.7 + (char_count_tra / max(char_count_sim + char_count_tra, 1)) * 30 * 0.3
        
        if score_sim > score_tra:
            detected = "chi_sim"
            logger.debug(f"Chinese variant detected: Simplified (conf: {avg_conf_sim:.1f}, chars: {char_count_sim})")
        else:
            detected = "chi_tra"
            logger.debug(f"Chinese variant detected: Traditional (conf: {avg_conf_tra:.1f}, chars: {char_count_tra})")
        
        return detected
    except Exception as e:
        # Fallback: mặc định là Simplified (phổ biến hơn)
        logger.warning(f"Không thể detect Chinese variant: {e}. Mặc định dùng chi_sim")
        return "chi_sim"


def _resolve_language(lang: str, ocr_cfg: dict, sample_img: Optional["Image.Image"] = None) -> str:
    """
    Resolve language code, chỉ hỗ trợ Chinese variant detection (giản thể/phồn thể).
    Auto-detect ngôn ngữ đã được loại bỏ do kém hiệu quả.
    
    Args:
        lang: Language string từ config (có thể là "VN", "EN", "CN", "VN+EN", "chi", "chi_sim", "chi_tra", etc.)
        ocr_cfg: OCR config
        sample_img: Optional sample image để detect Chinese variant (chỉ khi lang="CN" hoặc "chi")
    
    Returns:
        Resolved language string cho Tesseract (e.g., "chi_sim", "chi_tra", "vie+eng")
    """
    if not lang:
        return "vie"
    
    # Normalize: VN/EN/CN → vie/eng/chi
    lang = _normalize_lang_code(lang)
    
    # Loại bỏ auto-detect: nếu config là "auto", cảnh báo và fallback về "vie"
    if lang == "auto" or lang.startswith("auto+"):
        logger.warning(f"Auto-detect ngôn ngữ đã bị loại bỏ do kém hiệu quả. "
                      f"Config '{lang}' không được hỗ trợ. Vui lòng chỉ định rõ ngôn ngữ (VN/EN/CN). "
                      f"Fallback về 'vie'.")
        lang = "vie"
    
    # Chỉ hỗ trợ detect Chinese variant (giản thể/phồn thể) khi lang="CN" hoặc "chi"
    # Kiểm tra nếu có "chi" (cần detect variant: Simplified vs Traditional)
    if "chi" in lang.lower() and "chi_sim" not in lang and "chi_tra" not in lang:
        # Cần detect variant
        if sample_img is not None:
            detected_variant = _detect_chinese_variant(sample_img, ocr_cfg)
            # Replace "chi" bằng variant detected
            lang = lang.replace("chi", detected_variant).replace("Chi", detected_variant)
            # Clean up duplicate "+" nếu có
            lang = lang.replace(f"{detected_variant}+{detected_variant}", detected_variant)
            logger.info(f"Auto-detected Chinese variant: {detected_variant} → Language: {lang}")
        else:
            # Không có sample image → mặc định Simplified
            detected_variant = "chi_sim"
            lang = lang.replace("chi", detected_variant).replace("Chi", detected_variant)
            logger.info(f"No sample image for detection, defaulting to chi_sim → Language: {lang}")
    
    return lang


def _image_to_text(img: "Image.Image", ocr_cfg: dict, lang_override: Optional[str] = None) -> str:
    """
    OCR một ảnh thành text.
    
    Args:
        img: PIL Image object
        ocr_cfg: OCR config dictionary
        lang_override: Optional resolved language string (đã detect variant nếu cần)
    """
    # Auto-rotate trước khi OCR (dựa vào EXIF và OSD của Tesseract)
    try:
        if bool(ocr_cfg.get("auto_rotate", True)):
            img = _auto_rotate_image(img, ocr_cfg)
    except Exception:
        pass

    lang = lang_override
    if lang is None:
        raw_lang = ocr_cfg.get("lang", "vie+eng")
        lang = _resolve_language(raw_lang, ocr_cfg, sample_img=img)
    
    psm = int(ocr_cfg.get("psm", 3) or 3)
    config = f"--psm {psm}"
    return pytesseract.image_to_string(img, lang=lang, config=config)


# Cache để tránh detect PDF type nhiều lần cho cùng một file
_pdf_type_cache: dict[str, tuple[str, float]] = {}  # {pdf_path: (result, timestamp)}
_cache_timeout = 300.0  # Cache 5 phút

def detect_pdf_type(pdf_path: str, ocr_cfg: Optional[dict] = None) -> str:
    """
    Phát hiện PDF là scan hay text-based.
    Returns: "text" hoặc "scan"
    
    Logic cải thiện:
    - Sample nhiều trang hơn (5 trang đầu + 1 trang giữa + 1 trang cuối) để tăng độ chính xác
    - Threshold thấp hơn (20 ký tự thay vì 100) để không bỏ sót text-based ít text ở đầu
    - Timeout tăng lên 20 giây để xử lý PDF lớn
    - Kiểm tra chất lượng text: Nếu có text có thể extract được (dù ít) → text-based
    - Retry nếu file chưa sẵn sàng (đặc biệt khi chạy trên web)
    - CẢI TIẾN: Cache kết quả để tránh detect lại nhiều lần cho cùng một file
    """
    import threading
    import time
    
    # CẢI TIẾN: Kiểm tra cache trước khi detect lại
    pdf_path_normalized = os.path.abspath(pdf_path) if os.path.exists(pdf_path) else pdf_path
    if pdf_path_normalized in _pdf_type_cache:
        cached_result, cached_time = _pdf_type_cache[pdf_path_normalized]
        if time.time() - cached_time < _cache_timeout:
            logger.debug(f"✅ Dùng kết quả cache cho PDF type: {cached_result} (cache age: {time.time() - cached_time:.1f}s)")
            return cached_result
        else:
            # Cache hết hạn, xóa và detect lại
            del _pdf_type_cache[pdf_path_normalized]
            logger.debug(f"Cache đã hết hạn cho {pdf_path_normalized}, detect lại...")
    
    # Đảm bảo file tồn tại và có thể đọc được (retry nếu cần)
    max_retries = 3
    retry_delay = 0.2
    for attempt in range(max_retries):
        try:
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                # Thử mở file để đảm bảo không bị lock
                with open(pdf_path, 'rb') as test_file:
                    test_file.read(1)
                break
        except (IOError, PermissionError) as e:
            if attempt < max_retries - 1:
                logger.debug(f"File chưa sẵn sàng, retry {attempt + 1}/{max_retries}: {e}")
                time.sleep(retry_delay)
            else:
                logger.warning(f"Không thể đọc file sau {max_retries} lần thử: {e}")
                return "scan"
    
    result_container = {"result": None, "done": False}
    exception_container = {"exception": None}
    
    def _detect_inner():
        """Hàm detect chạy trong thread riêng với timeout"""
        try:
            if pdfplumber is not None:
                logger.debug("Thử dùng pdfplumber...")
                try:
                    with pdfplumber.open(pdf_path) as pdf:
                        total_pages = len(pdf.pages)
                        if total_pages == 0:
                            logger.warning(f"PDF không có trang nào: {pdf_path}")
                            result_container["result"] = "scan"
                            result_container["done"] = True
                            return
                        
                        logger.debug(f"PDF có {total_pages} trang, bắt đầu sample...")
                        total_chars = 0
                        pages_with_text = 0
                        
                        # Sample nhiều điểm: 5 trang đầu + 1 trang giữa + 1 trang cuối
                        sample_indices = list(range(min(5, total_pages)))  # 5 trang đầu
                        if total_pages > 10:
                            sample_indices.append(total_pages // 2)  # Trang giữa
                        if total_pages > 5:
                            sample_indices.append(total_pages - 1)  # Trang cuối
                        
                        # Loại bỏ trùng lặp và sort
                        sample_indices = sorted(set(sample_indices))
                        logger.debug(f"Sample {len(sample_indices)} trang: {[i+1 for i in sample_indices]}")
                        
                        for idx in sample_indices:
                            try:
                                page = pdf.pages[idx]
                                text = page.extract_text()
                                if text and text.strip():
                                    text_len = len(text.strip())
                                    total_chars += text_len
                                    pages_with_text += 1
                                    logger.debug(f"Trang {idx+1}: {text_len} ký tự")
                                    # Nếu có trang nào có > 50 ký tự → chắc chắn text-based
                                    if text_len > 50:
                                        result_container["result"] = "text"
                                        result_container["done"] = True
                                        logger.info(f"✅ Phát hiện text-based: Trang {idx+1} có {text_len} ký tự")
                                        return
                            except Exception as e:
                                logger.debug(f"Lỗi khi extract trang {idx+1}: {e}")
                                continue
                        
                        # Nếu có text ở nhiều trang (>= 2) hoặc tổng > 20 ký tự → text-based
                        if pages_with_text >= 2 or total_chars > 20:
                            result_container["result"] = "text"
                            logger.info(f"✅ Phát hiện text-based: {pages_with_text} trang có text, tổng {total_chars} ký tự")
                        else:
                            result_container["result"] = "scan"
                            logger.info(f"📷 Phát hiện scan: Chỉ {pages_with_text} trang có text, tổng {total_chars} ký tự")
                        result_container["done"] = True
                        return
                except Exception as e:
                    logger.warning(f"pdfplumber failed: {e}")
                    import traceback
                    logger.debug(traceback.format_exc())
                    pass
            
            # Fallback: dùng PyPDF2
            if PyPDF2 is not None:
                try:
                    logger.debug("Thử dùng PyPDF2 (fallback)...")
                    with open(pdf_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        total_pages = len(reader.pages)
                        if total_pages == 0:
                            logger.warning(f"PDF không có trang nào (PyPDF2): {pdf_path}")
                            result_container["result"] = "scan"
                            result_container["done"] = True
                            return
                        
                        logger.debug(f"PDF có {total_pages} trang (PyPDF2), bắt đầu sample...")
                        total_chars = 0
                        pages_with_text = 0
                        
                        # Sample nhiều điểm
                        sample_indices = list(range(min(5, total_pages)))
                        if total_pages > 10:
                            sample_indices.append(total_pages // 2)
                        if total_pages > 5:
                            sample_indices.append(total_pages - 1)
                        
                        sample_indices = sorted(set(sample_indices))
                        logger.debug(f"Sample {len(sample_indices)} trang (PyPDF2): {[i+1 for i in sample_indices]}")
                        
                        for idx in sample_indices:
                            try:
                                page = reader.pages[idx]
                                text = page.extract_text()
                                if text and text.strip():
                                    text_len = len(text.strip())
                                    total_chars += text_len
                                    pages_with_text += 1
                                    logger.debug(f"Trang {idx+1} (PyPDF2): {text_len} ký tự")
                                    if text_len > 50:
                                        result_container["result"] = "text"
                                        result_container["done"] = True
                                        logger.info(f"✅ Phát hiện text-based (PyPDF2): Trang {idx+1} có {text_len} ký tự")
                                        return
                            except Exception as e:
                                logger.debug(f"Lỗi khi extract trang {idx+1} (PyPDF2): {e}")
                                continue
                        
                        if pages_with_text >= 2 or total_chars > 20:
                            result_container["result"] = "text"
                            logger.info(f"✅ Phát hiện text-based (PyPDF2): {pages_with_text} trang có text, tổng {total_chars} ký tự")
                        else:
                            result_container["result"] = "scan"
                            logger.info(f"📷 Phát hiện scan (PyPDF2): Chỉ {pages_with_text} trang có text, tổng {total_chars} ký tự")
                        result_container["done"] = True
                        return
                except Exception as e:
                    logger.warning(f"PyPDF2 failed: {e}")
                    import traceback
                    logger.debug(traceback.format_exc())
                    pass
            
            # Nếu không thể detect → giả định là scan
            result_container["result"] = "scan"
            result_container["done"] = True
        except Exception as e:
            exception_container["exception"] = e
            logger.debug(f"Exception trong _detect_inner: {e}")
            result_container["result"] = "scan"
            result_container["done"] = True
    
    # Chạy trong thread với timeout tăng lên 20 giây
    thread = threading.Thread(target=_detect_inner, daemon=True)
    thread.start()
    thread.join(timeout=20.0)  # Timeout 20 giây (tăng từ 10s)
    
    if not result_container["done"]:
        # Timeout xảy ra - thử cách đơn giản hơn: chỉ check 1 trang đầu
        logger.warning(f"Timeout khi detect PDF type sau 20 giây, thử cách đơn giản hơn: {pdf_path}")
        try:
            if pdfplumber is not None:
                with pdfplumber.open(pdf_path) as pdf:
                    if len(pdf.pages) > 0:
                        text = pdf.pages[0].extract_text()
                        if text and len(text.strip()) > 10:
                            logger.info("Phát hiện text-based (quick check trang đầu)")
                            return "text"
            elif PyPDF2 is not None:
                with open(pdf_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    if len(reader.pages) > 0:
                        text = reader.pages[0].extract_text()
                        if text and len(text.strip()) > 10:
                            logger.info("Phát hiện text-based (quick check trang đầu)")
                            return "text"
        except Exception:
            pass
        logger.warning("Giả định là scan sau timeout")
        return "scan"
    
    if exception_container["exception"]:
        logger.debug(f"Exception khi detect PDF type: {exception_container['exception']}")
    
    if result_container["result"]:
        result = result_container["result"]
        # CẢI TIẾN: Lưu kết quả vào cache
        pdf_path_normalized = os.path.abspath(pdf_path) if os.path.exists(pdf_path) else pdf_path
        _pdf_type_cache[pdf_path_normalized] = (result, time.time())
        return result
    
    # Fallback cuối cùng
    logger.warning(f"Không thể detect PDF type, giả định là scan: {pdf_path}")
    result = "scan"
    # CẢI TIẾN: Lưu kết quả vào cache
    pdf_path_normalized = os.path.abspath(pdf_path) if os.path.exists(pdf_path) else pdf_path
    _pdf_type_cache[pdf_path_normalized] = (result, time.time())
    return result
def extract_text_from_pdf(pdf_path: str, ocr_cfg: dict, pages: Optional[List[int]] = None) -> str:
    """
    Extract text từ PDF có text layer (không cần OCR).
    
    Args:
        pdf_path: Đường dẫn file PDF
        ocr_cfg: Config dictionary
        pages: Danh sách số trang cần extract (1-indexed). None = tất cả trang.
    """
    texts: List[str] = []
    
    # Ưu tiên pdfplumber (chính xác hơn)
    if pdfplumber is not None:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                total = len(pdf.pages)
                logger.info(f"Extract text: Tổng số trang: {total}")
                
                # Filter pages nếu có chỉ định
                if pages:
                    # Validate pages (phải trong khoảng [1, total])
                    valid_pages = [p for p in pages if 1 <= p <= total]
                    invalid_pages = [p for p in pages if p < 1 or p > total]
                    if invalid_pages:
                        logger.warning(f"Các trang không hợp lệ (nằm ngoài 1-{total}): {invalid_pages}. Bỏ qua.")
                    if not valid_pages:
                        logger.error("Không có trang hợp lệ nào để extract.")
                        return ""
                    logger.info(f"Extract text: Chỉ extract {len(valid_pages)} trang: {valid_pages}")
                    pages_to_extract = sorted(set(valid_pages))
                else:
                    pages_to_extract = list(range(1, total + 1))
                
                show_progress = bool(ocr_cfg.get("show_progress", True))
                
                if show_progress and tqdm is not None and len(pages_to_extract) > 1:
                    for page_num in tqdm(pages_to_extract, desc="Extract text", unit="trang"):
                        page = pdf.pages[page_num - 1]  # pdfplumber dùng 0-indexed
                        text = page.extract_text()
                        if text:
                            texts.append(text.strip())
                else:
                    for page_num in pages_to_extract:
                        page = pdf.pages[page_num - 1]  # pdfplumber dùng 0-indexed
                        text = page.extract_text()
                        if text:
                            texts.append(text.strip())
                        if len(texts) % 50 == 0:
                            logger.info(f"Extract text: {len(texts)}/{len(pages_to_extract)} trang")
                return "\n\n".join(texts)
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyPDF2...")
    
    # Fallback: PyPDF2
    if PyPDF2 is not None:
        try:
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                total = len(reader.pages)
                logger.info(f"Extract text: Tổng số trang: {total}")
                
                # Filter pages nếu có chỉ định
                if pages:
                    valid_pages = [p for p in pages if 1 <= p <= total]
                    invalid_pages = [p for p in pages if p < 1 or p > total]
                    if invalid_pages:
                        logger.warning(f"Các trang không hợp lệ (nằm ngoài 1-{total}): {invalid_pages}. Bỏ qua.")
                    if not valid_pages:
                        logger.error("Không có trang hợp lệ nào để extract.")
                        return ""
                    logger.info(f"Extract text: Chỉ extract {len(valid_pages)} trang: {valid_pages}")
                    pages_to_extract = sorted(set(valid_pages))
                else:
                    pages_to_extract = list(range(1, total + 1))
                
                show_progress = bool(ocr_cfg.get("show_progress", True))
                
                if show_progress and tqdm is not None and len(pages_to_extract) > 1:
                    for page_num in tqdm(pages_to_extract, desc="Extract text", unit="trang"):
                        page = reader.pages[page_num - 1]  # PyPDF2 dùng 0-indexed
                        text = page.extract_text()
                        if text:
                            texts.append(text.strip())
                else:
                    for page_num in pages_to_extract:
                        page = reader.pages[page_num - 1]  # PyPDF2 dùng 0-indexed
                        text = page.extract_text()
                        if text:
                            texts.append(text.strip())
                        if len(texts) % 50 == 0:
                            logger.info(f"Extract text: {len(texts)}/{len(pages_to_extract)} trang")
                return "\n\n".join(texts)
        except Exception as e:
            logger.error(f"PyPDF2 failed: {e}")
            raise
    
    raise RuntimeError("Không có thư viện extract PDF text. Cài pdfplumber hoặc PyPDF2.")


def extract_text_blocks_with_position(pdf_path: str, ocr_cfg: dict, pages: Optional[List[int]] = None) -> tuple[List[dict], int]:
    """
    Extract text blocks với Y-position từ PDF (dùng pdfplumber hoặc PyMuPDF).
    
    Returns:
        tuple: (text_blocks_by_page, total_pages) trong đó text_blocks_by_page là dict:
            {page_num: [{"text": str, "y_position": float, "x_position": float, "bbox": tuple}, ...]}
    """
    text_blocks_by_page = {}
    total_pages = 0
    
    # Ưu tiên pdfplumber (có bbox chính xác hơn)
    if pdfplumber is not None:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                total_pages = len(pdf.pages)
                
                if pages:
                    pages_to_extract = sorted(set([p for p in pages if 1 <= p <= total_pages]))
                else:
                    pages_to_extract = list(range(1, total_pages + 1))
                
                for page_num in pages_to_extract:
                    page = pdf.pages[page_num - 1]
                    text_blocks = []
                    
                    # Extract text với bbox từ pdfplumber
                    words = page.extract_words()
                    if words:
                        # Group words thành blocks dựa trên Y-position (same line)
                        current_block = []
                        current_y = None
                        
                        for word in words:
                            word_y = word.get("top", 0)  # Y-position của word
                            word_text = word.get("text", "")
                            word_x = word.get("x0", 0)
                            
                            # Nếu Y khác biệt nhiều (> threshold), tạo block mới
                            if current_y is None or abs(word_y - current_y) > 5:  # 5px threshold
                                if current_block:
                                    # Save previous block
                                    block_text = " ".join([w["text"] for w in current_block])
                                    block_y = current_block[0].get("top", 0)
                                    block_x = min([w.get("x0", 0) for w in current_block])
                                    block_bbox = (
                                        min([w.get("x0", 0) for w in current_block]),
                                        current_block[0].get("top", 0),
                                        max([w.get("x1", 0) for w in current_block]),
                                        max([w.get("bottom", 0) for w in current_block])
                                    )
                                    text_blocks.append({
                                        "text": block_text,
                                        "y_position": block_y,
                                        "x_position": block_x,
                                        "bbox": block_bbox
                                    })
                                current_block = [word]
                                current_y = word_y
                            else:
                                current_block.append(word)
                        
                        # Save last block
                        if current_block:
                            block_text = " ".join([w["text"] for w in current_block])
                            block_y = current_block[0].get("top", 0)
                            block_x = min([w.get("x0", 0) for w in current_block])
                            block_bbox = (
                                min([w.get("x0", 0) for w in current_block]),
                                current_block[0].get("top", 0),
                                max([w.get("x1", 0) for w in current_block]),
                                max([w.get("bottom", 0) for w in current_block])
                            )
                            text_blocks.append({
                                "text": block_text,
                                "y_position": block_y,
                                "x_position": block_x,
                                "bbox": block_bbox
                            })
                    
                    text_blocks_by_page[page_num] = text_blocks
                
                return text_blocks_by_page, total_pages
        except Exception as e:
            logger.warning(f"pdfplumber extract text blocks failed: {e}, fallback to PyMuPDF...")
    
    # Fallback: PyMuPDF (extract text blocks với bbox)
    if fitz is not None:
        try:
            doc = fitz.open(pdf_path)
            total_pages = len(doc)
            
            if pages:
                pages_to_extract = sorted(set([p for p in pages if 1 <= p <= total_pages]))
            else:
                pages_to_extract = list(range(1, total_pages + 1))
            
            for page_num in pages_to_extract:
                page = doc[page_num - 1]
                text_blocks = []
                
                # Extract text blocks với bbox từ PyMuPDF
                blocks = page.get_text("dict")  # Get text as dict with bbox info
                if blocks and "blocks" in blocks:
                    for block in blocks["blocks"]:
                        if "lines" in block:  # Text block
                            block_text = ""
                            min_y = float('inf')
                            min_x = float('inf')
                            max_x = float('-inf')
                            max_y = float('-inf')
                            
                            for line in block["lines"]:
                                for span in line.get("spans", []):
                                    span_text = span.get("text", "")
                                    bbox = span.get("bbox", [0, 0, 0, 0])
                                    
                                    if span_text:
                                        block_text += span_text + " "
                                    
                                    # Update bbox
                                    min_y = min(min_y, bbox[1])  # top (Y)
                                    min_x = min(min_x, bbox[0])  # left (X)
                                    max_x = max(max_x, bbox[2])  # right
                                    max_y = max(max_y, bbox[3])  # bottom
                            
                            if block_text.strip():
                                text_blocks.append({
                                    "text": block_text.strip(),
                                    "y_position": min_y,
                                    "x_position": min_x,
                                    "bbox": (min_x, min_y, max_x, max_y)
                                })
                
                text_blocks_by_page[page_num] = text_blocks
            
            doc.close()
            return text_blocks_by_page, total_pages
        except Exception as e:
            logger.warning(f"PyMuPDF extract text blocks failed: {e}")
    
    # Fallback: không có position, return empty
    logger.warning("Không thể extract text blocks với position. Trả về empty.")
    return {}, total_pages if total_pages > 0 else 0


def extract_format_hints(para, para_index: int, total_paragraphs: int) -> dict:
    """
    Extract format hints chi tiết từ paragraph.
    
    Args:
        para: python-docx Paragraph object
        para_index: Index của paragraph trong document
        total_paragraphs: Tổng số paragraphs
    
    Returns:
        dict: {
            "style": str,
            "font_size": float,
            "is_bold": bool,
            "is_italic": bool,
            "alignment": str,
            "position_hint": str  # "top", "middle", "bottom"
        }
    """
    hints = {
        "style": para.style.name if para.style else "Normal",
        "font_size": None,
        "is_bold": False,
        "is_italic": False,
        "alignment": "left",
        "position_hint": "middle"
    }
    
    # Get alignment
    if para.alignment is not None and WD_PARAGRAPH_ALIGNMENT is not None:
        if para.alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
            hints["alignment"] = "left"
        elif para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            hints["alignment"] = "center"
        elif para.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            hints["alignment"] = "right"
        elif para.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            hints["alignment"] = "justify"
    
    # Get font info từ runs
    if para.runs:
        # Use first run (hoặc most common format)
        first_run = para.runs[0]
        
        if first_run.font.size:
            hints["font_size"] = first_run.font.size.pt
        
        hints["is_bold"] = first_run.font.bold is True
        hints["is_italic"] = first_run.font.italic is True
    
    # Estimate position: first 20% = top, last 20% = bottom, else = middle
    if total_paragraphs > 0:
        position_ratio = para_index / total_paragraphs
        if position_ratio < 0.2:
            hints["position_hint"] = "top"
        elif position_ratio > 0.8:
            hints["position_hint"] = "bottom"
        else:
            hints["position_hint"] = "middle"
    
    return hints


def is_in_table(para) -> bool:
    """Check nếu paragraph nằm trong table."""
    try:
        parent = para._element.getparent()
        if parent is not None:
            # Check nếu parent là table element
            return parent.tag.endswith('tbl')
    except Exception:
        pass
    return False


def extract_images_from_paragraph(para) -> List[dict]:
    """
    Extract images từ paragraph.
    
    Returns:
        List[dict]: [
            {
                "run_index": int,
                "image_data": bytes,
                "width": float,  # inches (estimate)
                "height": float,  # inches (estimate)
                "run": Run  # Reference để re-insert sau
            },
            ...
        ]
    """
    images = []
    
    try:
        for run_idx, run in enumerate(para.runs):
            # Check nếu run có image (check for blip element)
            blips = run._element.xpath('.//a:blip')
            if blips:
                try:
                    # Get image relationship ID
                    blip = blips[0]
                    rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    
                    if rId and hasattr(para.part, 'rels') and rId in para.part.rels:
                        # Get image blob từ relationship
                        image_part_rel = para.part.rels[rId]
                        image_blob = image_part_rel.target_part.blob
                        
                        # Estimate image dimensions từ run
                        width_inches = 4.0  # Default
                        height_inches = 3.0  # Default
                        
                        # Try to get actual dimensions từ image
                        try:
                            if Image is not None:
                                import io
                                img = Image.open(io.BytesIO(image_blob))
                                width_inches = img.width / 96.0  # Convert pixels to inches (96 DPI)
                                height_inches = img.height / 96.0
                                img.close()
                        except Exception:
                            pass
                        
                        images.append({
                            "run_index": run_idx,
                            "image_data": image_blob,
                            "width": width_inches,
                            "height": height_inches,
                            "run": run
                        })
                except Exception as e:
                    logger.debug(f"Không thể extract image từ run {run_idx}: {e}")
                    continue
    except Exception as e:
        logger.debug(f"Lỗi khi extract images từ paragraph: {e}")
    
    return images


def extract_paragraphs_with_hints(docx_path: str) -> List[dict]:
    """
    Extract paragraphs từ DOCX với format hints chi tiết và images.
    Skip tables.
    
    Returns:
        List[dict]: [
            {
                "index": int,
                "text": str,
                "hints": dict,
                "images": List[dict],
                "has_images": bool,
                "para_object": Paragraph
            },
            ...
        ]
    """
    if Document is None:
        raise RuntimeError("python-docx chưa được cài đặt.")
    
    doc = Document(docx_path)
    paragraphs_data = []
    
    # Get all paragraphs (not in tables)
    all_paragraphs = [para for para in doc.paragraphs if not is_in_table(para)]
    total_paragraphs = len(all_paragraphs)
    
    for para_idx, para in enumerate(all_paragraphs):
        # Extract format hints
        hints = extract_format_hints(para, para_idx, total_paragraphs)
        
        # Extract images
        images = extract_images_from_paragraph(para)
        
        paragraphs_data.append({
            "index": len(paragraphs_data),
            "text": para.text,
            "hints": hints,
            "images": images,
            "has_images": len(images) > 0,
            "para_object": para
        })
    
    logger.info(f"📝 Đã extract {len(paragraphs_data)} paragraphs từ DOCX (đã skip tables)")
    return paragraphs_data


def batch_small_paragraphs(paragraphs_data: List[dict], min_chars: int = 50) -> List[dict]:
    """
    Batch các paragraphs nhỏ lại với nhau để tối ưu API calls.
    
    Args:
        paragraphs_data: List các paragraph dicts
        min_chars: Ngưỡng để xem là paragraph nhỏ (default: 50)
    
    Returns:
        List[dict]: Batched paragraphs với type "batch" hoặc "single"
    """
    batched = []
    current_batch = []
    
    for para in paragraphs_data:
        # Batch nếu paragraph nhỏ và không có images
        if len(para["text"]) < min_chars and not para.get("has_images", False):
            current_batch.append(para)
        else:
            # Paragraph lớn hoặc có images → process riêng
            if current_batch:
                # Merge batch
                batched.append({
                    "type": "batch",
                    "text": "\n\n".join([p["text"] for p in current_batch]),
                    "original_indices": [p["index"] for p in current_batch],
                    "para_objects": [p["para_object"] for p in current_batch],
                    "images_list": [p["images"] for p in current_batch],
                    "hints": current_batch[0]["hints"]
                })
                current_batch = []
            
            batched.append({
                "type": "single",
                **para
            })
    
    # Handle remaining batch
    if current_batch:
        batched.append({
            "type": "batch",
            "text": "\n\n".join([p["text"] for p in current_batch]),
            "original_indices": [p["index"] for p in current_batch],
            "para_objects": [p["para_object"] for p in current_batch],
            "images_list": [p["images"] for p in current_batch],
            "hints": current_batch[0]["hints"]
        })
    
    if batched:
        batch_count = sum(1 for b in batched if b["type"] == "batch")
        single_count = sum(1 for b in batched if b["type"] == "single")
        logger.info(f"📦 Đã batch: {batch_count} batches, {single_count} single paragraphs")
    
    return batched
def build_cleanup_prompt_with_hints(text: str, hints: dict) -> str:
    """
    Build cleanup prompt với format hints chi tiết.
    
    Args:
        text: Paragraph text
        hints: Format hints dict
    
    Returns:
        str: Prompt với hints
    """
    hint_descriptions = []
    
    # Style hints
    style = hints.get("style", "")
    if style.startswith("Heading"):
        hint_descriptions.append(f"Style: {style} (có thể là header/title)")
    elif style == "Normal":
        hint_descriptions.append("Style: Normal (nội dung chính)")
    
    # Font size hints
    font_size = hints.get("font_size")
    if font_size:
        if font_size < 10:
            hint_descriptions.append(f"Font rất nhỏ ({font_size}pt) - có thể là footer/page number")
        elif font_size > 14:
            hint_descriptions.append(f"Font lớn ({font_size}pt) - có thể là header/title")
    
    # Bold hints
    if hints.get("is_bold"):
        hint_descriptions.append("Bold - có thể là header/title")
    
    # Position hints
    position = hints.get("position_hint", "middle")
    if position == "top":
        hint_descriptions.append("Vị trí: Đầu trang - có thể là header")
    elif position == "bottom":
        hint_descriptions.append("Vị trí: Cuối trang - có thể là footer/page number")
    
    # Alignment hints
    alignment = hints.get("alignment", "left")
    if alignment == "center":
        hint_descriptions.append("Căn giữa - có thể là title/header")
    
    prompt = f"""Bạn là AI chuyên dọn dẹp văn bản OCR/scan.

THÔNG TIN FORMATTING:
{chr(10).join('- ' + d for d in hint_descriptions) if hint_descriptions else '- Không có thông tin đặc biệt'}

Dựa trên formatting này, xác định:
- Nếu là header/footer/page number → XÓA
- Nếu là nội dung chính → GIỮ LẠI và cleanup noise

Nhiệm vụ:
1. Loại bỏ header/footer lặp lại ở mỗi trang
2. Loại bỏ số trang, watermark
3. Loại bỏ các ký tự rác, vệt đen vô nghĩa từ quá trình scan
4. Chuẩn hóa khoảng trắng thừa
5. Giữ nguyên nội dung chính của văn bản
6. Giữ nguyên định dạng đoạn văn

Trả về chỉ văn bản đã được dọn dẹp, không giải thích thêm.

Văn bản cần dọn dẹp:
{text}"""
    
    return prompt


def cleanup_paragraph_with_hints(para_data: dict, ocr_cfg: dict) -> dict:
    """
    Cleanup một paragraph/batch với format hints.
    
    Args:
        para_data: Paragraph dict với type "single" hoặc "batch"
        ocr_cfg: Config dictionary
    
    Returns:
        dict: {
            "cleaned_text": str,
            "should_merge_with_next": bool  # Nếu AI merge với paragraph sau
        }
    """
    cleanup_cfg = ocr_cfg.get("ai_cleanup", {})
    if not cleanup_cfg.get("enabled", False):
        return {
            "cleaned_text": para_data["text"],
            "should_merge_with_next": False
        }
    
    # Get API keys
    api_keys = cleanup_cfg.get("api_keys", [])
    if not api_keys:
        api_keys = ocr_cfg.get("_root_api_keys", [])
    if not api_keys:
        logger.warning("Không có API keys cho cleanup, bỏ qua")
        return {
            "cleaned_text": para_data["text"],
            "should_merge_with_next": False
        }
    
    model_name = cleanup_cfg.get("model", "gemini-2.5-flash")
    timeout_s = cleanup_cfg.get("timeout", 60.0)
    
    # Get safety settings
    safety_level = cleanup_cfg.get("safety_level") or ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
    safety_settings = _build_safety_settings(safety_level)
    
    # Build prompt với hints
    text = para_data["text"]
    hints = para_data.get("hints", {})
    prompt = build_cleanup_prompt_with_hints(text, hints)
    
    # Call AI cleanup (dùng async function)
    try:
        # Use first API key (có thể parallelize sau nếu cần)
        cleaned_text = asyncio.run(_cleanup_chunk_async(
            text, api_keys[0], model_name, prompt, 0, 1, timeout_s, safety_settings
        ))
        
        # Simple heuristic: Nếu cleaned text ngắn hơn nhiều → có thể đã merge hoặc xóa
        # Không có cách chính xác để detect merge, tạm thời return False
        # Có thể cải thiện bằng cách prompt AI explicit về merge
        
        return {
            "cleaned_text": cleaned_text,
            "should_merge_with_next": False  # TODO: Implement merge detection
        }
    except Exception as e:
        logger.warning(f"Cleanup paragraph thất bại: {e}")
        return {
            "cleaned_text": text,
            "should_merge_with_next": False
        }


def spell_check_paragraph(para_data: dict, ocr_cfg: dict) -> str:
    """
    Spell check một paragraph/batch.
    
    Args:
        para_data: Paragraph dict (có thể là processed sau cleanup)
        ocr_cfg: Config dictionary
    
    Returns:
        str: Spell-checked text
    """
    spell_check_cfg = ocr_cfg.get("ai_spell_check", {})
    if not spell_check_cfg.get("enabled", False):
        return para_data.get("cleaned_text", para_data["text"])
    
    # Get API keys
    api_keys = spell_check_cfg.get("api_keys", [])
    if not api_keys:
        api_keys = ocr_cfg.get("_root_api_keys", [])
    if not api_keys:
        logger.warning("Không có API keys cho spell check, bỏ qua")
        return para_data.get("cleaned_text", para_data["text"])
    
    model_name = spell_check_cfg.get("model", "gemini-2.5-flash")
    timeout_s = spell_check_cfg.get("timeout", 60.0)
    
    # Get safety settings
    safety_level = spell_check_cfg.get("safety_level") or ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
    safety_settings = _build_safety_settings(safety_level)
    
    text = para_data.get("cleaned_text", para_data["text"])
    
    # Build spell check prompt (giống như ai_spell_check_and_paragraph_restore)
    # Lấy prompt từ existing function hoặc tạo mới
    prompt = """Bạn là AI chuyên soát lỗi chính tả và phục hồi cấu trúc paragraph cho văn bản OCR/scan.

Nhiệm vụ:
1. Soát lỗi chính tả do OCR
2. Phục hồi cấu trúc paragraph hợp lý
3. Nối các câu bị ngắt paragraph (nếu cần)
4. Giữ nguyên nội dung và ý nghĩa

Trả về chỉ văn bản đã được soát và phục hồi, không giải thích thêm.

Văn bản cần phân tích và xử lý:
""" + text
    
    try:
        # Use async spell check function (cần check xem có sẵn không)
        # Tạm thời dùng _cleanup_chunk_async với spell check prompt
        spell_checked_text = asyncio.run(_cleanup_chunk_async(
            text, api_keys[0], model_name, prompt, 0, 1, timeout_s, safety_settings
        ))
        return spell_checked_text
    except Exception as e:
        logger.warning(f"Spell check paragraph thất bại: {e}")
        return text


def convert_pdf_with_ocrmypdf(pdf_path: str, output_path: str, ocr_cfg: dict, pages: Optional[List[int]] = None) -> str:
    """
    Convert PDF → PDF searchable bằng OCRmyPDF (thêm OCR layer).
    
    Args:
        pdf_path: Đường dẫn file PDF input
        output_path: Đường dẫn file PDF output (searchable)
        ocr_cfg: Config dictionary
        pages: Danh sách số trang cần OCR (1-indexed). None = tất cả trang.
    
    Returns:
        str: Đường dẫn file PDF đã tạo (searchable)
    
    Raises:
        RuntimeError: Nếu OCRmyPDF chưa được cài đặt hoặc conversion fail
    """
    global ocrmypdf, ocrmypdf_available
    
    if not ocrmypdf_available or ocrmypdf is None:
        raise RuntimeError("OCRmyPDF chưa được cài đặt. Cài ocrmypdf để dùng fallback workflow.")
    
    logger.info(f"🔍 Đang dùng OCRmyPDF để tạo PDF searchable: {pdf_path}")
    
    # Validate input PDF
    if not os.path.exists(pdf_path):
        raise RuntimeError(f"File PDF không tồn tại: {pdf_path}")
    
    file_size = os.path.getsize(pdf_path)
    if file_size < 100:
        raise RuntimeError(f"File PDF quá nhỏ hoặc không hợp lệ: {pdf_path} ({file_size} bytes)")
    
    # Kiểm tra magic bytes (PDF signature)
    try:
        with open(pdf_path, 'rb') as f:
            header = f.read(4)
            if header != b'%PDF':
                raise RuntimeError(f"File không phải PDF hợp lệ (magic bytes: {header}): {pdf_path}")
    except Exception as e:
        raise RuntimeError(f"Không thể đọc file PDF: {e}")
    
    try:
        # Tạo temp file nếu output_path trùng với input_path
        temp_output = None
        if os.path.abspath(pdf_path) == os.path.abspath(output_path):
            temp_output = output_path + ".tmp"
            final_output = output_path
        else:
            final_output = output_path
        
        # Chuẩn bị command cho OCRmyPDF (gọi qua subprocess)
        cmd = ["ocrmypdf"]
        
        # Language settings từ config
        # Normalize language code từ config format (VN/EN/CN) sang OCRmyPDF format (vie/eng/chi_sim)
        raw_lang = ocr_cfg.get("lang", "eng")
        if raw_lang:
            # Sử dụng hàm normalize đã có để convert format
            lang_normalized = _normalize_lang_code(raw_lang)
            # OCRmyPDF hỗ trợ multiple languages với dấu +, ví dụ: "vie+eng"
            cmd.extend(["-l", lang_normalized])
            if lang_normalized != raw_lang:
                logger.debug(f"Normalized language code: '{raw_lang}' → '{lang_normalized}'")
        
        # Deskew option (làm thẳng trang nghiêng)
        if ocr_cfg.get("deskew", False):
            cmd.append("--deskew")
        
        # Rotate pages option
        if ocr_cfg.get("rotate_pages", False):
            cmd.append("--rotate-pages")
        
        # Jobs (multi-core)
        jobs = ocr_cfg.get("jobs", 1)
        if jobs and jobs > 1:
            cmd.extend(["--jobs", str(jobs)])
        
        # Skip text pages (tối ưu cho PDF đã có text layer)
        if ocr_cfg.get("skip_text", False):
            cmd.append("--skip-text")
            logger.debug("OCRmyPDF: Bật --skip-text (bỏ qua pages đã có text)")
        
        # Force OCR (khi cần OCR lại cả text layer)
        if ocr_cfg.get("force_ocr", False):
            cmd.append("--force-ocr")
            logger.debug("OCRmyPDF: Bật --force-ocr (OCR lại cả text layer)")
        
        # Optimize output file
        optimize_level = ocr_cfg.get("optimize_level", None)
        optimize_flag_added = False
        if optimize_level is not None:
            try:
                lvl = int(optimize_level)
                if lvl > 0:
                    cmd.extend(["--optimize", str(lvl)])
                    optimize_flag_added = True
                    logger.debug(f"OCRmyPDF: --optimize {lvl}")
                else:
                    logger.debug("OCRmyPDF: optimize_level=0 → không tối ưu để tránh lỗi extract_images")
            except Exception:
                pass
        elif ocr_cfg.get("optimize", True):  # Mặc định bật optimize nếu không có optimize_level
            cmd.append("--optimize")
            optimize_flag_added = True
            logger.debug("OCRmyPDF: Bật --optimize (tối ưu kích thước file)")
        
        # Extra args cho OCRmyPDF (nếu có)
        extra_args = ocr_cfg.get("ocrmypdf_extra_args", [])
        if isinstance(extra_args, (list, tuple)) and extra_args:
            cmd.extend([str(a) for a in extra_args])
            logger.debug(f"OCRmyPDF: Thêm extra args: {extra_args}")
        
        # Pages (OCRmyPDF hỗ trợ pages thông qua --pages)
        # Format: "1,3,5" hoặc "1-5" hoặc "1,3-5"
        if pages:
            valid_pages = sorted(set(pages))
            # Tối ưu format: "1-3" thay vì "1,2,3" nếu liên tục
            pages_ranges = []
            i = 0
            while i < len(valid_pages):
                start = valid_pages[i]
                end = start
                # Tìm chuỗi liên tục
                while i + 1 < len(valid_pages) and valid_pages[i + 1] == end + 1:
                    i += 1
                    end = valid_pages[i]
                
                if start == end:
                    pages_ranges.append(str(start))
                else:
                    pages_ranges.append(f"{start}-{end}")
                i += 1
            
            pages_str = ",".join(pages_ranges)
            cmd.extend(["--pages", pages_str])
            logger.info(f"Chỉ OCR {len(valid_pages)} trang: {valid_pages} (format: {pages_str})")
        
        # Add input và output paths
        cmd.append(pdf_path)
        cmd.append(temp_output if temp_output else final_output)
        
        # Gọi OCRmyPDF qua subprocess (command line)
        logger.info(f"Chạy OCRmyPDF: {' '.join(cmd)}")
        
        # Đảm bảo Ghostscript trong PATH cho subprocess
        env = os.environ.copy()
        current_path = env.get("PATH", "")
        
        # Tìm Ghostscript tự động
        gs_bin_path = None
        # Thử các đường dẫn phổ biến
        possible_paths = [
            r"C:\Program Files\gs\gs10.06.0\bin",
            r"C:\Program Files\gs\gs10.05.0\bin",
            r"C:\Program Files\gs\gs10.04.0\bin",
            r"C:\Program Files (x86)\gs\gs10.06.0\bin",
            r"C:\Program Files (x86)\gs\gs10.05.0\bin",
        ]
        
        # Tìm trong Program Files
        if sys.platform == "win32":
            import glob
            for pattern in [r"C:\Program Files\gs\gs*\bin", r"C:\Program Files (x86)\gs\gs*\bin"]:
                matches = glob.glob(pattern)
                if matches:
                    # Sort để lấy version mới nhất
                    matches.sort(reverse=True)
                    gs_bin_path = matches[0]
                    logger.info(f"✅ Tìm thấy Ghostscript: {gs_bin_path}")
                    break
        
        # Nếu không tìm thấy, thử các đường dẫn phổ biến
        if not gs_bin_path:
            for path in possible_paths:
                if os.path.exists(path) and os.path.exists(os.path.join(path, "gswin64c.exe")):
                    gs_bin_path = path
                    logger.info(f"✅ Tìm thấy Ghostscript: {gs_bin_path}")
                    break
        
        # Thêm vào PATH nếu tìm thấy (luôn thêm vào đầu để đảm bảo ưu tiên)
        if gs_bin_path:
            # Luôn thêm vào đầu PATH để đảm bảo subprocess tìm thấy
            env["PATH"] = gs_bin_path + os.pathsep + current_path
            logger.info(f"✅ Đã thêm Ghostscript vào PATH cho subprocess: {gs_bin_path}")
            
            # Verify Ghostscript có hoạt động trong env này không
            try:
                test_result = subprocess.run(
                    ["gswin64c", "--version"],
                    capture_output=True,
                    text=True,
                    env=env,
                    timeout=5
                )
                if test_result.returncode == 0:
                    logger.info(f"✅ Verified Ghostscript: {test_result.stdout.strip()}")
                else:
                    logger.warning(f"⚠️  Ghostscript test failed trong subprocess env")
            except Exception as e:
                logger.debug(f"Không thể verify Ghostscript: {e}")
        else:
            logger.warning("⚠️  Không tìm thấy Ghostscript, OCRmyPDF có thể fail")
        
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=3600,  # Timeout 1 giờ
            env=env  # Pass environment với PATH đã cập nhật
        )
        
        if result.returncode != 0:
            error_msg = result.stderr or result.stdout or "Unknown error"
            # Log chi tiết lỗi để debug
            logger.debug(f"OCRmyPDF stderr: {result.stderr[:500] if result.stderr else 'None'}")
            logger.debug(f"OCRmyPDF stdout: {result.stdout[:500] if result.stdout else 'None'}")
            # Retry nếu lỗi do extract_images/optimize
            em_lower = (error_msg or "").lower()
            if ("extract_images" in em_lower) or ("optimize.py" in em_lower):
                logger.warning("⚠️  OCRmyPDF lỗi khi tối ưu ảnh (extract_images). Thử lại với --optimize 0...")
                # Loại bỏ các cờ optimize khỏi cmd và đặt optimize 0
                cmd_no_opt = [a for a in cmd if a != "--optimize"]
                # Nếu có dạng --optimize <level>, loại bỏ cả level đi
                try:
                    while "--optimize" in cmd_no_opt:
                        idx = cmd_no_opt.index("--optimize")
                        # Bỏ cả tham số tiếp theo nếu là số
                        del cmd_no_opt[idx]
                        if idx < len(cmd_no_opt) and str(cmd_no_opt[idx]).isdigit():
                            del cmd_no_opt[idx]
                except Exception:
                    pass
                # Thêm optimize 0 để tắt tối ưu
                cmd_retry = []
                for a in cmd_no_opt:
                    cmd_retry.append(a)
                cmd_retry.insert(1, "--optimize")
                cmd_retry.insert(2, "0")
                logger.info(f"Thử lại: {' '.join(cmd_retry)}")
                result_retry = subprocess.run(
                    cmd_retry,
                    capture_output=True,
                    text=True,
                    timeout=3600,
                    env=env
                )
                if result_retry.returncode == 0:
                    # Gán result để tiếp tục các bước sau
                    result = result_retry
                else:
                    logger.debug(f"Retry stderr: {result_retry.stderr[:500] if result_retry.stderr else 'None'}")
                    raise subprocess.CalledProcessError(result_retry.returncode, cmd_retry, output=result_retry.stdout, stderr=result_retry.stderr)
            else:
                raise subprocess.CalledProcessError(result.returncode, cmd, output=result.stdout, stderr=result.stderr)
        
        # Move temp file nếu cần
        if temp_output and os.path.exists(temp_output):
            if os.path.exists(final_output):
                os.remove(final_output)
            os.rename(temp_output, final_output)
        
        # Validate output file
        if not os.path.exists(final_output):
            raise RuntimeError(f"OCRmyPDF conversion thất bại: File output không tồn tại: {final_output}")
        
        file_size = os.path.getsize(final_output)
        if file_size < 100:  # PDF tối thiểu phải > 100 bytes
            raise RuntimeError(f"OCRmyPDF conversion thất bại: File output quá nhỏ ({file_size} bytes)")
        
        logger.info(f"✅ Đã tạo PDF searchable bằng OCRmyPDF: {final_output} ({file_size} bytes)")
        return final_output
        
    except FileNotFoundError:
        raise RuntimeError(
            "OCRmyPDF không được tìm thấy trên PATH.\n"
            "Cài đặt: pip install ocrmypdf\n"
            "Đảm bảo OCRmyPDF đã được cài đặt và có trong PATH."
        )
    except subprocess.TimeoutExpired:
        raise RuntimeError("OCRmyPDF timeout sau 1 giờ. File có thể quá lớn hoặc có vấn đề.")
    except subprocess.CalledProcessError as e:
        error_output = e.stderr if e.stderr else (e.stdout if e.stdout else str(e))
        # Log chi tiết để debug
        logger.debug(f"OCRmyPDF error (exit code {e.returncode}): {error_output[:1000]}")
        if 'ghostscript' in error_output.lower() or 'gswin64c' in error_output.lower() or 'gs' in error_output.lower():
            raise RuntimeError(
                "OCRmyPDF cần Ghostscript nhưng không tìm thấy trên PATH.\n"
                "💡 Cài đặt Ghostscript:\n"
                "   - Windows: choco install ghostscript\n"
                "   - Hoặc tải từ: https://www.ghostscript.com/download/gsdnld.html\n"
                "   - Đảm bảo thêm Ghostscript vào PATH sau khi cài đặt"
            )
        elif 'tesseract' in error_output.lower():
            raise RuntimeError(
                "OCRmyPDF cần Tesseract OCR nhưng không tìm thấy.\n"
                "💡 Cài đặt Tesseract:\n"
                "   - Windows: choco install tesseract\n"
                "   - Hoặc tải từ: https://github.com/UB-Mannheim/tesseract/wiki\n"
                "   - Đảm bảo thêm Tesseract vào PATH sau khi cài đặt"
            )
        else:
            raise RuntimeError(f"OCRmyPDF thất bại (exit code {e.returncode}): {error_output}")
    except Exception as e:
        error_msg = str(e)
        if 'ghostscript' in error_msg.lower() or 'gswin64c' in error_msg.lower() or 'gs' in error_msg.lower():
            raise RuntimeError(
                "OCRmyPDF cần Ghostscript nhưng không tìm thấy trên PATH.\n"
                "💡 Cài đặt Ghostscript:\n"
                "   - Windows: choco install ghostscript\n"
                "   - Hoặc tải từ: https://www.ghostscript.com/download/gsdnld.html\n"
                "   - Đảm bảo thêm Ghostscript vào PATH sau khi cài đặt"
            )
        logger.error(f"❌ Lỗi khi dùng OCRmyPDF: {error_msg}")
        import traceback
        logger.debug(traceback.format_exc())
        raise
def _extract_tables_with_unstructured(pdf_path: str, output_path: str, ocr_cfg: dict, pages: Optional[List[int]] = None) -> dict:
    """
    Extract bảng từ PDF bằng unstructured.io (độ chính xác 95-98%).
    
    Đây là phương pháp được khuyến cáo nhất theo Grok review:
    - Sử dụng unstructured.io với OCRmyPDF để tự động detect layout
    - Tự động gộp multi-line cell
    - Phát hiện chính xác bảng không có đường kẻ
    - Không cần tự viết alignment phức tạp
    
    Returns:
        dict: {page_num: {"rows": [[cell1, cell2, ...], ...], "num_cols": int}}
    """
    try:
        # Lazy import unstructured
        try:
            from unstructured.partition.pdf import partition_pdf
            from unstructured.documents.elements import Table
        except ImportError:
            logger.warning("⚠️  unstructured.io chưa được cài đặt. Cài bằng: pip install unstructured[pdf]")
            logger.warning("💡 Hoặc dùng phương pháp fallback (pytesseract + DBSCAN)")
            return {}
        
        logger.info("🔍 Đang extract bảng bằng unstructured.io (độ chính xác cao)...")
        
        tables_by_page = {}
        
        # Partition PDF với strategy hi_res (OCR + layout detection)
        try:
            elements = partition_pdf(
                filename=pdf_path,
                strategy="hi_res",  # Quan trọng: dùng OCR + layout detection
                infer_table_structure=True,  # Bật nhận diện bảng
                languages=["vie", "eng"],  # Hỗ trợ tiếng Việt và Anh
            )
        except Exception as e:
            logger.warning(f"⚠️  unstructured.io partition thất bại: {e}")
            logger.warning("💡 Fallback về phương pháp khác...")
            return {}
        
        # Extract tables từ elements
        current_page = 1
        for element in elements:
            if isinstance(element, Table):
                # Convert table thành list of rows
                rows = []
                if hasattr(element, 'metadata') and element.metadata.page_number:
                    current_page = element.metadata.page_number
                
                # Lấy text từ table (unstructured đã gộp multi-line cell)
                if hasattr(element, 'text_as_html'):
                    # Parse HTML table
                    import re
                    html_text = element.text_as_html
                    # Simple HTML table parser (có thể cải thiện)
                    # Tìm tất cả <tr>...</tr>
                    tr_pattern = r'<tr[^>]*>(.*?)</tr>'
                    tr_matches = re.findall(tr_pattern, html_text, re.DOTALL | re.IGNORECASE)
                    
                    for tr_match in tr_matches:
                        # Tìm tất cả <td>...</td> hoặc <th>...</th>
                        td_pattern = r'<t[dh][^>]*>(.*?)</t[dh]>'
                        td_matches = re.findall(td_pattern, tr_match, re.DOTALL | re.IGNORECASE)
                        
                        # Clean HTML tags và whitespace
                        cells = []
                        for td_text in td_matches:
                            # Loại bỏ HTML tags
                            cell_text = re.sub(r'<[^>]+>', '', td_text)
                            # Clean whitespace
                            cell_text = " ".join(cell_text.split())
                            cells.append(cell_text)
                        
                        if cells:
                            rows.append(cells)
                elif hasattr(element, 'text'):
                    # Fallback: parse từ text (ít chính xác hơn)
                    lines = element.text.split('\n')
                    for line in lines:
                        if line.strip():
                            # Giả định delimiter là tab hoặc nhiều spaces
                            cells = [c.strip() for c in re.split(r'\t+|\s{2,}', line) if c.strip()]
                            if cells:
                                rows.append(cells)
                
                if rows:
                    # Tìm số cột tối đa
                    max_cols = max(len(row) for row in rows) if rows else 0
                    
                    # Pad các hàng để có cùng số cột
                    normalized_rows = []
                    for row in rows:
                        normalized_row = row + [""] * (max_cols - len(row)) if len(row) < max_cols else row[:max_cols]
                        normalized_rows.append(normalized_row)
                    
                    tables_by_page[current_page] = {
                        "page": current_page,
                        "rows": normalized_rows,
                        "num_cols": max_cols
                    }
                    logger.info(f"✅ unstructured.io: Đã extract bảng trang {current_page}: {len(normalized_rows)} hàng, {max_cols} cột")
        
        if tables_by_page:
            logger.info(f"✅ unstructured.io: Hoàn tất extract {len(tables_by_page)} bảng")
            return tables_by_page
        else:
            logger.info("ℹ️  unstructured.io: Không tìm thấy bảng")
            return {}
            
    except Exception as e:
        logger.warning(f"⚠️  Lỗi khi extract bảng bằng unstructured.io: {e}")
        import traceback
        logger.debug(traceback.format_exc())
        return {}


def _extract_tables_pytesseract_advanced(img: "Image.Image", ocr_cfg: dict, page_num: int = 1) -> List[List[str]]:
    """
    Extract bảng từ ảnh bằng pytesseract với DBSCAN clustering (cải thiện từ Grok).
    
    Cải thiện so với cách cũ:
    - Dùng image_to_data với level=5 (word)
    - Cluster theo X để tìm cột (DBSCAN tự động tìm số cột)
    - Cluster theo Y để tìm hàng
    - Gộp các word thuộc cùng ô (multi-line, multi-word)
    
    Returns:
        List[List[str]]: List các hàng, mỗi hàng là list các ô
    """
    try:
        import pytesseract
        import numpy as np
        from sklearn.cluster import DBSCAN
    except ImportError as e:
        logger.warning(f"⚠️  Thiếu dependency cho advanced pytesseract: {e}")
        return []
    
    # Lấy config
    lang = ocr_cfg.get("lang", "vie")
    psm = ocr_cfg.get("psm", 6)  # PSM 6 tốt hơn cho bảng
    
    # Extract data với pytesseract
    try:
        data = pytesseract.image_to_data(
            img, lang=lang, output_type=pytesseract.Output.DICT,
            config=f"--psm {psm} --oem 3"
        )
    except Exception as e:
        logger.warning(f"⚠️  pytesseract image_to_data thất bại: {e}")
        return []
    
    # Lọc chỉ word có text và conf > 30
    words = []
    for i in range(len(data.get("text", []))):
        conf = int(data.get("conf", [0])[i]) if i < len(data.get("conf", [])) else 0
        text = data.get("text", [""])[i] if i < len(data.get("text", [])) else ""
        
        if conf > 30 and text.strip():
            left = data.get("left", [0])[i] if i < len(data.get("left", [])) else 0
            top = data.get("top", [0])[i] if i < len(data.get("top", [])) else 0
            width = data.get("width", [0])[i] if i < len(data.get("width", [])) else 0
            height = data.get("height", [0])[i] if i < len(data.get("height", [])) else 0
            
            words.append({
                "text": text.strip(),
                "left": left,
                "top": top,
                "width": width,
                "height": height,
                "right": left + width,
                "bottom": top + height,
                "x_center": left + width // 2,
                "y_center": top + height // 2,
            })
    
    if not words:
        logger.debug(f"Trang {page_num}: Không có word nào được detect")
        return []
    
    # 1. Cluster cột theo X (DBSCAN tự động tìm số cột)
    X_coords = np.array([[w["x_center"]] for w in words])
    clustering = DBSCAN(eps=30, min_samples=2).fit(X_coords)  # eps=30 pixels
    col_labels = clustering.labels_
    
    # 2. Gán từng word vào cột
    col_to_words = {}
    for label, word in zip(col_labels, words):
        if label != -1:  # Không phải noise
            col_to_words.setdefault(label, []).append(word)
    
    if not col_to_words:
        logger.debug(f"Trang {page_num}: Không tìm thấy cột hợp lệ")
        return []
    
    # Sắp xếp cột từ trái sang phải
    sorted_cols = sorted(col_to_words.items(), key=lambda x: np.mean([w["x_center"] for w in x[1]]))
    
    # 3. Với mỗi cột → cluster hàng theo Y
    row_groups_by_col = []
    
    for col_idx, (label, col_words) in enumerate(sorted_cols):
        Y_coords = np.array([[w["y_center"]] for w in col_words])
        row_clustering = DBSCAN(eps=20, min_samples=1).fit(Y_coords)
        row_labels = row_clustering.labels_
        
        # Gán word vào hàng trong cột này
        row_dict = {}
        for row_label, word in zip(row_labels, col_words):
            row_dict.setdefault(row_label, []).append(word)
        
        # Sắp xếp các hàng theo Y
        sorted_rows = sorted(row_dict.items(), key=lambda x: np.mean([w["y_center"] for w in x[1]]))
        row_groups_by_col.append(sorted_rows)
    
    # 4. Ghép các hàng tương ứng giữa các cột → tạo bảng
    if not row_groups_by_col:
        return []
    
    max_rows = max(len(rows) for rows in row_groups_by_col)
    table_rows = []
    
    for row_idx in range(max_rows):
        row_cells = []
        for col_group in row_groups_by_col:
            if row_idx < len(col_group):
                cell_words = sorted(col_group[row_idx][1], key=lambda w: w["left"])
                cell_text = " ".join(w["text"] for w in cell_words).strip()
            else:
                cell_text = ""
            row_cells.append(cell_text)
        
        # Chỉ thêm hàng nếu có ít nhất một ô có nội dung
        if any(cell.strip() for cell in row_cells):
            table_rows.append(row_cells)
    
    return table_rows


def _try_extract_tables_from_pdf_via_ocrmypdf(pdf_path: str, output_path: str, ocr_cfg: dict, pages: Optional[List[int]] = None) -> dict:
    """Tạo searchable PDF bằng OCRmyPDF rồi thử extract bảng bằng pdfplumber.
    
    Trả về dict {page_num: {"rows": [...], "num_cols": int}} thay vì tạo CSV file.
    
    Lưu ý: Đây là bước nhẹ, chỉ thực hiện khi tables.reconstruct = true.
    
    Returns:
        dict: {page_num: {"rows": [[cell1, cell2, ...], ...], "num_cols": int}}
    """
    try:
        if pdfplumber is None:
            logger.warning("⚠️  pdfplumber không khả dụng → bỏ qua extract bảng")
            return {}
        
        # Strategy 1: Thử extract từ PDF gốc trước (có thể có text layer ẩn)
        source_pdf = None
        use_searchable = False
        
        logger.info("🔍 Thử extract bảng từ PDF gốc trước...")
        try:
            with pdfplumber.open(pdf_path) as pdf:
                test_page = pdf.pages[0] if len(pdf.pages) > 0 else None
                if test_page:
                    test_tables = test_page.extract_tables()
                    if test_tables:
                        logger.info("✅ PDF gốc có text layer → extract trực tiếp từ PDF gốc")
                        source_pdf = pdf_path
                        use_searchable = False
                    else:
                        raise ValueError("PDF gốc không có text layer")
                else:
                    raise ValueError("PDF không có trang")
        except Exception as e:
            logger.debug(f"PDF gốc không có text layer: {e}")
            # Strategy 2: Tạo searchable PDF bằng OCRmyPDF
            if not ocrmypdf_available:
                logger.warning("⚠️  OCRmyPDF không khả dụng và PDF gốc không có text layer")
                logger.warning("💡 Để extract bảng từ PDF scan, cần cài Ghostscript:")
                logger.warning("   - Windows: choco install ghostscript")
                logger.warning("   - Hoặc tải từ: https://www.ghostscript.com/download/gsdnld.html")
                # Không return, để tiếp tục đến fallback OpenCV
                source_pdf = None
            else:
                temp_searchable_pdf = os.path.splitext(output_path)[0] + "_searchable_for_tables.pdf"
                logger.info(f"📄 Tạo searchable PDF để extract bảng: {temp_searchable_pdf}")
                try:
                    # Tạo searchable PDF (nhanh vì chỉ phục vụ table detect)
                    # Override optimize=0 để tránh lỗi extract_images trong optimize.py
                    ocr_cfg_no_opt = dict(ocr_cfg)
                    ocr_cfg_no_opt["optimize"] = False
                    ocr_cfg_no_opt["optimize_level"] = 0
                    convert_pdf_with_ocrmypdf(pdf_path, temp_searchable_pdf, ocr_cfg_no_opt, pages)
                    source_pdf = temp_searchable_pdf
                    use_searchable = True
                except Exception as e:
                    logger.warning(f"⚠️  Không thể tạo searchable PDF: {e}")
                    logger.warning("💡 Có thể cần cài Ghostscript để extract bảng từ PDF scan")
                    # Không return, để tiếp tục đến fallback OpenCV
                    source_pdf = None
        
        if source_pdf is None:
            # Fallback cuối: trích bảng từ ảnh bằng OpenCV nếu được cấu hình
            tables_mode = (ocr_cfg.get("tables") or {}).get("mode", "auto")
            if tables_mode in ("auto", "opencv_grid"):
                logger.info("🧭 Fallback: Dùng OpenCV để trích xuất bảng trực tiếp từ ảnh")
                try:
                    tables_dict = _extract_tables_from_images_cv(pdf_path, output_path, ocr_cfg, pages)
                    return tables_dict  # Trả về dict thay vì None
                except Exception as cv_err:
                    logger.warning(f"⚠️  Fallback OpenCV thất bại: {cv_err}")
                    return {}
            logger.warning("⚠️  Không thể xác định source PDF để extract bảng")
            return {}
        
        logger.info("🔍 Bắt đầu extract bảng từ PDF...")
        # Extract tables per page - lưu vào dict thay vì CSV
        base = os.path.splitext(output_path)[0]
        found_any = False
        total_tables = 0
        tables_by_page = {}  # {page_num: {"rows": [...], "num_cols": int}}
        
        if not os.path.exists(source_pdf):
            logger.warning(f"⚠️  PDF không tồn tại: {source_pdf}")
            return {}
        
        with pdfplumber.open(source_pdf) as pdf:
            total_pages = len(pdf.pages)
            logger.info(f"📖 Đang scan {total_pages} trang để tìm bảng...")
            
            # Nếu có pages chỉ định, chỉ xử lý những trang đó
            page_indices = list(range(total_pages))
            if pages:
                # Convert 1-indexed to 0-indexed
                page_indices = [p - 1 for p in pages if 1 <= p <= total_pages]
                logger.info(f"📄 Chỉ extract bảng từ {len(page_indices)} trang: {pages}")
            
            for page_idx in page_indices:
                try:
                    page = pdf.pages[page_idx]
                    # Thử cả lattice và stream strategy
                    tables = page.extract_tables()
                    if not tables:
                        # Thử với strategy khác
                        try:
                            tables = page.extract_tables(strategy="lattice")
                        except Exception:
                            try:
                                tables = page.extract_tables(strategy="stream")
                            except Exception:
                                tables = []
                    
                    if not tables:
                        logger.debug(f"Trang {page_idx + 1}: Không tìm thấy bảng")
                        continue
                    
                    found_any = True
                    table_count = len(tables)
                    total_tables += table_count
                    logger.info(f"📊 Trang {page_idx + 1}: Tìm thấy {table_count} bảng")
                    
                    # Gộp tất cả tables trên trang thành một bảng lớn
                    all_rows = []
                    max_cols = 0
                    for tbl_idx, tbl in enumerate(tables, start=1):
                        if tbl_idx > 1:
                            # Thêm hàng trống giữa các bảng
                            all_rows.append([""] * max_cols)
                        for row in tbl:
                            if row:  # Skip empty rows
                                cleaned_row = [(cell or "").strip() for cell in row]
                                all_rows.append(cleaned_row)
                                max_cols = max(max_cols, len(cleaned_row))
                    
                    # Pad các hàng để có cùng số cột
                    for row in all_rows:
                        while len(row) < max_cols:
                            row.append("")
                    
                    tables_by_page[page_idx + 1] = {
                        "page": page_idx + 1,
                        "rows": all_rows,
                        "num_cols": max_cols
                    }
                    logger.info(f"🗂️  Đã extract bảng trang {page_idx + 1}: {len(all_rows)} hàng, {max_cols} cột")
                except Exception as e:
                    logger.warning(f"⚠️  Extract bảng lỗi ở trang {page_idx + 1}: {e}")
                    import traceback
                    logger.debug(traceback.format_exc())
                    continue
        
        # Cleanup temp và trả về kết quả
        if found_any:
            logger.info(f"✅ Hoàn tất extract bảng: {total_tables} bảng từ {len(tables_by_page)} trang")
            # Cleanup temp (chỉ nếu dùng searchable PDF)
            if use_searchable:
                try:
                    if os.path.exists(source_pdf) and ocr_cfg.get("cleanup_temp_searchable_pdf", True):
                        os.remove(source_pdf)
                        logger.debug(f"🗑️  Đã xóa temp searchable PDF: {source_pdf}")
                except Exception as e:
                    logger.debug(f"Không thể xóa temp file: {e}")
            return tables_by_page
        else:
            logger.info("ℹ️  Không tìm thấy bảng trong PDF (có thể PDF không có bảng hoặc OCR chưa đủ tốt)")
            return {}
    except Exception as e:
        logger.error(f"❌ _try_extract_tables_from_pdf_via_ocrmypdf lỗi: {e}")
        import traceback
        logger.debug(traceback.format_exc())
        return {}


def _extract_tables_from_images_cv(pdf_path: str, output_path: str, ocr_cfg: dict, pages: Optional[List[int]] = None) -> dict:
    """Fallback: Trích xuất bảng từ ảnh PDF bằng OpenCV (phát hiện lưới).

    - Render PDF → ảnh bằng pdf2image (dựa vào poppler_path trong config)
    - Dùng OpenCV morphology để tìm đường kẻ dọc/ngang → xác định ô
    - Trả về dict {page_num: {rows, num_cols, metadata}} thay vì tạo CSV file
    
    Returns:
        dict: {page_num: {"rows": [[cell1, cell2, ...], ...], "num_cols": int, "metadata": [...]}}
    """
    try:
        from pdf2image import convert_from_path as _convert_from_path
    except Exception as e:
        logger.warning(f"⚠️  Thiếu pdf2image: {e}")
        return {}
    try:
        import cv2
        import numpy as np
    except Exception as e:
        logger.warning(f"⚠️  Thiếu OpenCV/numpy: {e}")
        return

    poppler_path = ocr_cfg.get("poppler_path") or os.environ.get("POPPLER_PATH")
    dpi = int(ocr_cfg.get("dpi", 250))

    # Xác định trang cần render
    try:
        from PyPDF2 import PdfReader as _PdfReader
        total = len(_PdfReader(open(pdf_path, 'rb')).pages)
    except Exception:
        total = None
    if pages:
        page_indices = [p for p in pages if p >= 1 and (total is None or p <= total)]
    else:
        page_indices = [1]

    images = _convert_from_path(pdf_path, dpi=dpi, poppler_path=poppler_path, first_page=min(page_indices),
                                last_page=max(page_indices))

    base = os.path.splitext(output_path)[0]
    exported = 0
    tables_by_page = {}  # {page_num: {"rows": [...], "num_cols": int, "metadata": [...]}}
    for i, pil_img in enumerate(images, start=min(page_indices)):
        # Chuyển PIL → OpenCV BGR
        img = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        # Nhị phân hoá & đảo màu (để line rõ hơn)
        bw = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C,
                                   cv2.THRESH_BINARY, 15, -2)
        bw_inv = 255 - bw

        # Tách đường kẻ dọc
        vertical = bw_inv.copy()
        rows = vertical.shape[0]
        vertical_size = max(1, rows // 40)
        verticalStructure = cv2.getStructuringElement(cv2.MORPH_RECT, (1, vertical_size))
        vertical = cv2.erode(vertical, verticalStructure)
        vertical = cv2.dilate(vertical, verticalStructure)

        # Tách đường kẻ ngang
        horizontal = bw_inv.copy()
        cols = horizontal.shape[1]
        horizontal_size = max(1, cols // 40)
        horizontalStructure = cv2.getStructuringElement(cv2.MORPH_RECT, (horizontal_size, 1))
        horizontal = cv2.erode(horizontal, horizontalStructure)
        horizontal = cv2.dilate(horizontal, horizontalStructure)

        # Kết hợp để được lưới
        grid = cv2.addWeighted(vertical, 0.5, horizontal, 0.5, 0.0)
        grid = cv2.threshold(grid, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]

        # Tìm contours để suy ra các ô
        contours, _ = cv2.findContours(grid, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
        # Lọc bbox lớn nhỏ, gom theo hàng/cột bằng heuristic
        boxes = []
        h, w = grid.shape
        for c in contours:
            x, y, cw, ch = cv2.boundingRect(c)
            if cw * ch < (h * w) * 0.0005:
                continue
            if cw < 20 or ch < 15:
                continue
            boxes.append((x, y, cw, ch))
        if not boxes:
            logger.info(f"Trang {i}: Không phát hiện bảng bằng OpenCV")
            continue

        # Sắp xếp theo Y trước để tiện xử lý
        boxes.sort(key=lambda b: (b[1], b[0]))
        heights = [ch for (_, _, _, ch) in boxes]
        widths = [cw for (_, _, cw, _) in boxes]
        med_h = int(np.median(heights)) if heights else 20
        med_w = int(np.median(widths)) if widths else 40
        # Tham số từ config
        ocr_tables_cfg = (ocr_cfg.get("tables") or {}).get("ocr", {})
        row_merge_factor = float(ocr_tables_cfg.get("row_merge_factor", 0.5))
        col_merge_factor = float(ocr_tables_cfg.get("col_merge_factor", 0.6))
        min_cell_area_ratio = float(ocr_tables_cfg.get("min_cell_area_ratio", 0.0005))
        # Lọc box nhỏ theo tỉ lệ trang
        page_area = h * w
        boxes = [(x, y, cw, ch) for (x, y, cw, ch) in boxes if (cw * ch) >= page_area * min_cell_area_ratio]
        # Ngưỡng gom hàng/cột
        y_threshold = max(8, int(med_h * row_merge_factor))

        # Ổn định cột: lấy từ hàng đại diện có nhiều ô nhất để cố định số cột
        # Nếu không có rows_list (vì gom theo boxes), dùng toàn bộ boxes
        sorted_by_row = []
        temp_row = []
        last_y_center = None
        for box in boxes:
            y_center = box[1] + box[3] // 2
            if last_y_center is None or abs(y_center - last_y_center) <= y_threshold:
                temp_row.append(box)
            else:
                sorted_by_row.append(temp_row)
                temp_row = [box]
            last_y_center = y_center
        if temp_row:
            sorted_by_row.append(temp_row)

        # Phát hiện cột: sử dụng hàng đại diện để có số cột chính xác
        representative = max(sorted_by_row, key=lambda r: len(r)) if sorted_by_row else []
        if representative:
            rep_centers = sorted([x + cw // 2 for (x, _, cw, _) in representative])
            col_bins = []
            x_thresh = max(12, int(med_w * col_merge_factor))
            for xc in rep_centers:
                if not col_bins or abs(xc - col_bins[-1]) > x_thresh:
                    col_bins.append(xc)
        else:
            # Fallback: dùng toàn bộ boxes
            x_centers = sorted([x + cw // 2 for (x, _, cw, _) in boxes])
            if not x_centers:
                logger.info(f"Trang {i}: Không có cột hợp lệ sau phát hiện")
                continue
            col_bins = []
            x_thresh = max(12, int(med_w * col_merge_factor))
            for xc in x_centers:
                if not col_bins or abs(xc - col_bins[-1]) > x_thresh:
                    col_bins.append(xc)
        num_cols = len(col_bins)

        # Tạo row bins dựa trên center Y của boxes (trước khi merge)
        row_centers = sorted([y + ch // 2 for (_, y, _, ch) in boxes])
        row_bins = []
        for yc in row_centers:
            if not row_bins or abs(yc - row_bins[-1]) > y_threshold:
                row_bins.append(yc)
        num_rows = len(row_bins)
        if num_rows == 0 or num_cols == 0:
            logger.info(f"Trang {i}: Không tạo được lưới hàng/cột")
            continue

        # OCR từng ô để lấy nội dung text
        try:
            import pytesseract
            tesseract_cmd = ocr_cfg.get("tesseract_cmd")
            if tesseract_cmd:
                pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
            lang = ocr_cfg.get("lang", "vie")
            lang_normalized = _normalize_lang_code(lang)
            psm = ocr_cfg.get("psm", 6)
            # Tham số cải thiện cho bảng kém chất lượng
            ocr_tables_cfg = (ocr_cfg.get("tables") or {}).get("ocr", {})
            upscale_factor = int(ocr_tables_cfg.get("upscale_factor", 2))  # Phóng to ảnh để OCR tốt hơn
            use_clahe = bool(ocr_tables_cfg.get("use_clahe", True))        # Tăng tương phản cục bộ
            primary_psm = int(ocr_tables_cfg.get("primary_psm", psm))      # PSM chính (mặc định lấy từ ocr.psm)
            fallback_psm = int(ocr_tables_cfg.get("fallback_psm", 7))      # PSM fallback (single line)
            try_numeric_whitelist = bool(ocr_tables_cfg.get("try_numeric_whitelist", True))
        except Exception as e:
            logger.warning(f"⚠️  Không thể import pytesseract: {e}")
            pytesseract = None

        # Bước quan trọng: Nhóm các box có overlap hoặc gần nhau thành một ô
        # Điều này xử lý trường hợp text trong một ô bị ngắt dòng và được detect thành nhiều box
        def _boxes_overlap_or_near(box1, box2, overlap_threshold=0.4, near_threshold=2.0):
            """Kiểm tra hai box có overlap hoặc gần nhau không (thuộc cùng một ô)
            
            Chỉ merge khi:
            - Overlap đáng kể (>40%) theo cả hai chiều X và Y (chắc chắn cùng ô)
            - HOẶC cùng cột (dist_x rất nhỏ <30%) và khoảng cách Y nhỏ (<2.0x chiều cao trung bình)
            - HOẶC overlap X đáng kể (>40%) và khoảng cách Y nhỏ (text nhiều dòng trong cùng ô)
            """
            x1, y1, w1, h1 = box1
            x2, y2, w2, h2 = box2
            
            # Tính overlap theo chiều X và Y
            x_overlap = max(0, min(x1 + w1, x2 + w2) - max(x1, x2))
            y_overlap = max(0, min(y1 + h1, y2 + h2) - max(y1, y2))
            
            # Overlap ratio theo từng chiều
            overlap_x_ratio = x_overlap / min(w1, w2) if min(w1, w2) > 0 else 0
            overlap_y_ratio = y_overlap / min(h1, h2) if min(h1, h2) > 0 else 0
            
            # Nếu overlap đáng kể theo CẢ HAI chiều → chắc chắn cùng một ô
            if overlap_x_ratio >= overlap_threshold and overlap_y_ratio >= overlap_threshold:
                return True
            
            # Kiểm tra khoảng cách gần nhau (cho trường hợp text nhiều dòng trong cùng một ô)
            # Tính khoảng cách giữa các tâm
            center1_x, center1_y = x1 + w1 / 2, y1 + h1 / 2
            center2_x, center2_y = x2 + w2 / 2, y2 + h2 / 2
            
            # Khoảng cách tuyệt đối
            dist_x_abs = abs(center1_x - center2_x)
            dist_y_abs = abs(center1_y - center2_y)
            
            # Khoảng cách chuẩn hóa theo kích thước box
            avg_w = (w1 + w2) / 2
            avg_h = (h1 + h2) / 2
            
            dist_x_ratio = dist_x_abs / avg_w if avg_w > 0 else float('inf')
            dist_y_ratio = dist_y_abs / avg_h if avg_h > 0 else float('inf')
            
            # Trường hợp 1: Cùng cột (dist_x rất nhỏ) và khoảng cách Y hợp lý (text nhiều dòng trong cùng ô)
            # Chỉ merge nếu khoảng cách Y không quá lớn (<1.5x chiều cao trung bình)
            if (dist_x_ratio < 0.25 and dist_y_ratio < 1.5 and 
                overlap_x_ratio > 0.15):  # Có ít nhất 15% overlap X
                return True
            
            # Trường hợp 2: Overlap X đáng kể (>40%) và khoảng cách Y nhỏ (<1.5x) - text nhiều dòng trong cùng ô
            if overlap_x_ratio >= overlap_threshold and dist_y_ratio < 1.5:
                return True
            
            return False
        
        # Nhóm các box thành các cell groups
        cell_groups = []
        used_boxes = set()
        
        for i, box in enumerate(boxes):
            if i in used_boxes:
                continue
            
            # Tạo nhóm mới với box này
            group = [box]
            used_boxes.add(i)
            
            # Tìm tất cả các box khác có overlap hoặc gần với box này
            changed = True
            while changed:
                changed = False
                for j, other_box in enumerate(boxes):
                    if j in used_boxes:
                        continue
                    
                    # Kiểm tra overlap với bất kỳ box nào trong nhóm
                    for group_box in group:
                        if _boxes_overlap_or_near(group_box, other_box):
                            group.append(other_box)
                            used_boxes.add(j)
                            changed = True
                            break
            
            cell_groups.append(group)
        
        # Tính bounding box tổng hợp cho mỗi nhóm (để gán vào row/col chính xác hơn)
        # CẢI TIẾN: Filter các merged cells quá lớn (có thể span nhiều cột/hàng)
        merged_cells = []
        avg_cell_width = np.median([b[2] for b in boxes]) if boxes else med_w
        avg_cell_height = np.median([b[3] for b in boxes]) if boxes else med_h
        
        for group in cell_groups:
            if not group:
                continue
            
            # Tính bounding box bao phủ toàn bộ nhóm
            min_x = min(b[0] for b in group)
            min_y = min(b[1] for b in group)
            max_x = max(b[0] + b[2] for b in group)
            max_y = max(b[1] + b[3] for b in group)
            merged_box = (min_x, min_y, max_x - min_x, max_y - min_y)
            mw, mh = max_x - min_x, max_y - min_y
            
            # CẢI TIẾN: Filter các merged cells quá lớn
            # Nếu merged cell có width > 3x avg_cell_width hoặc height > 3x avg_cell_height
            # → có thể là nhiều cells bị merge nhầm → skip hoặc split
            if mw > avg_cell_width * 3.5 or mh > avg_cell_height * 3.5:
                # Cell quá lớn → có thể là header/footer hoặc nhiều cells bị merge nhầm
                # Chỉ giữ lại nếu có ít parts (có thể là cell thực sự lớn)
                if len(group) <= 2:
                    # Có thể là cell lớn hợp lệ (ví dụ: header)
                    merged_cells.append({
                        "merged_box": merged_box,
                        "parts": group
                    })
                else:
                    # Quá nhiều parts → có thể là merge nhầm → skip
                    logger.debug(f"Skipping merged cell quá lớn: {mw}x{mh} với {len(group)} parts")
                    continue
            else:
                merged_cells.append({
                    "merged_box": merged_box,
                    "parts": group  # Các box gốc trong nhóm
                })
        
        # Tạo bảng với nội dung OCR và metadata tọa độ dựa trên row/column bins
        table_cells = [
            [
                {"text": "", "cell_text_parts": []}
                for _ in range(num_cols)
            ]
            for _ in range(num_rows)
        ]

        # CẢI TIẾN: Tracking các cells đã được gán để tránh duplicate
        # Sử dụng set để track (row_idx, col_idx) đã được sử dụng bởi merged_cell nào
        cell_assignment_map = {}  # {(row_idx, col_idx): merged_cell_index}
        
        # Xử lý từng merged cell
        for merged_cell_idx, merged_cell in enumerate(merged_cells):
            merged_box = merged_cell["merged_box"]
            parts = merged_cell["parts"]
            
            x, y, cw, ch = merged_box
            
            # WORKFLOW MỚI: OCR trên toàn bộ merged box (ô đã ghép) thay vì từng phần riêng lẻ
            # Điều này đảm bảo context đầy đủ và spell check chính xác hơn
            cell_img = gray[max(0, y):min(gray.shape[0], y + ch), max(0, x):min(gray.shape[1], x + cw)]
            if cell_img.size == 0:
                continue

            combined_text = ""
            cell_text_parts_list = []
            
            if pytesseract is not None:
                try:
                    # OCR trên toàn bộ merged cell (ô đã ghép)
                    work = cell_img.copy()
                    if use_clahe:
                        try:
                            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
                            work = clahe.apply(work)
                        except Exception:
                            pass
                    try:
                        work = cv2.medianBlur(work, 3)
                    except Exception:
                        pass
                    try:
                        _, work_bin = cv2.threshold(work, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                        work = work_bin
                    except Exception:
                        pass
                    if upscale_factor and upscale_factor > 1:
                        try:
                            work = cv2.resize(work, None, fx=upscale_factor, fy=upscale_factor, interpolation=cv2.INTER_CUBIC)
                        except Exception:
                            pass
                    padding = 6
                    padded = cv2.copyMakeBorder(work, padding, padding, padding, padding, cv2.BORDER_CONSTANT, value=255)
                    
                    # Dùng PSM phù hợp với ô nhiều dòng
                    cfg_primary = f'--oem 1 --psm {primary_psm}'
                    text1 = pytesseract.image_to_string(padded, lang=lang_normalized, config=cfg_primary).strip()

                    def _clean_multiline(t: str) -> str:
                        lines = [ln.strip() for ln in t.splitlines() if ln.strip()]
                        return "\n".join(lines)

                    best_text = _clean_multiline(text1)

                    def _is_poor(t: str) -> bool:
                        if not t:
                            return True
                        alnum = sum(ch.isalnum() for ch in t)
                        return len(t) < 3 or alnum < max(1, len(t) // 3)

                    if _is_poor(best_text):
                        cfg_fb = f'--oem 1 --psm {fallback_psm}'
                        text2 = pytesseract.image_to_string(padded, lang=lang_normalized, config=cfg_fb).strip()
                        text2 = _clean_multiline(text2)
                        if len(text2) > len(best_text):
                            best_text = text2
                    if try_numeric_whitelist and _is_poor(best_text):
                        cfg_num = f'--oem 1 --psm {fallback_psm} -c tessedit_char_whitelist=0123456789.,-/%()'
                        text3 = pytesseract.image_to_string(padded, lang=lang_normalized, config=cfg_num).strip()
                        text3 = _clean_multiline(text3)
                        if len(text3) > len(best_text):
                            best_text = text3
                    
                    combined_text = best_text
                    
                    # CẢI TIẾN: Validation text quality trước khi lưu
                    # Filter các text quá ngắn hoặc có quá nhiều ký tự đặc biệt (có thể là noise)
                    def _is_valid_cell_text(text: str) -> bool:
                        """Kiểm tra text có hợp lệ không"""
                        if not text or len(text.strip()) < 2:
                            return False
                        
                        text_clean = text.strip()
                        
                        # Đếm số ký tự alphanumeric và tiếng Việt
                        vietnamese_chars = 'àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđĐ'
                        alnum_count = sum(1 for c in text_clean if c.isalnum() or c in vietnamese_chars)
                        
                        # Đếm các ký tự đặc biệt lạ (noise từ OCR)
                        noise_chars = '°‹Ÿ¬`^«£ÀÊSẼẴaAaPẨ€+—_*•…"\''
                        noise_count = sum(1 for c in text_clean if c in noise_chars)
                        
                        # Đếm các ký tự hợp lệ (chữ cái, số, dấu câu thông thường, khoảng trắng)
                        valid_chars = '.,;:!?()[]{}-/\\'
                        valid_count = alnum_count + sum(1 for c in text_clean if c in valid_chars or c.isspace())
                        
                        # Nếu có quá nhiều ký tự noise (>25% text) → có thể là noise
                        if len(text_clean) > 0 and noise_count / len(text_clean) > 0.25:
                            return False
                        
                        # Nếu có quá ít ký tự hợp lệ (<25% text) → có thể là noise
                        if len(text_clean) > 0 and valid_count / len(text_clean) < 0.25:
                            return False
                        
                        # Nếu có quá ít ký tự alphanumeric (<15% text) → có thể là noise
                        if len(text_clean) > 0 and alnum_count / len(text_clean) < 0.15:
                            return False
                        
                        # Kiểm tra nếu text chỉ chứa các ký tự đặc biệt và số ít chữ cái
                        # Ví dụ: "|||'‹ ° Cu|toi PÓ Tô Ÿ CC ST 4ì||"
                        if len(text_clean) > 10 and alnum_count < len(text_clean) * 0.2:
                            # Nếu có nhiều ký tự đặc biệt và ít chữ cái → noise
                            return False
                        
                        return True
                    
                    # Chỉ lưu nếu text hợp lệ
                    if _is_valid_cell_text(combined_text):
                        cell_text_parts_list.append({
                            "text": combined_text,
                            "x": int(x),
                            "y": int(y),
                            "w": int(cw),
                            "h": int(ch)
                        })
                    else:
                        # Text không hợp lệ → skip cell này
                        logger.debug(f"Skipping merged_cell {merged_cell_idx}: Text không hợp lệ (quá ngắn hoặc noise): '{combined_text[:50]}'")
                        continue
                except Exception as ocr_err:
                    logger.debug(f"OCR merged cell ({x},{y}) thất bại: {ocr_err}")
                    combined_text = ""
            
            # CẢI TIẾN: Nếu không có text hợp lệ từ OCR, skip cell này
            if not cell_text_parts_list:
                logger.debug(f"Skipping merged_cell {merged_cell_idx}: Không có text hợp lệ từ OCR")
                continue
            
            # Lấy combined_text từ cell_text_parts_list
            combined_text = " ".join([part["text"] for part in cell_text_parts_list])
            
            # Xác định row/col theo bins dựa trên merged box
            # CẢI TIẾN: Sử dụng cả x_center và phạm vi của cell để gán cột chính xác hơn
            x_center = x + cw // 2
            y_center = y + ch // 2
            x_left = x
            x_right = x + cw
            y_top = y
            y_bottom = y + ch
            
            # CẢI TIẾN: Tính overlap với từng cột để tìm cột có overlap lớn nhất
            # Thay vì chỉ tìm col_bin trong phạm vi, tính overlap thực tế
            col_scores = []
            for k in range(num_cols):
                col_bin_x = col_bins[k]
                # Tính overlap giữa cell và cột (giả định cột có width = khoảng cách đến cột kế tiếp)
                if k < num_cols - 1:
                    col_width = col_bins[k + 1] - col_bin_x
                else:
                    # Cột cuối: dùng khoảng cách từ cột trước
                    col_width = col_bin_x - col_bins[k - 1] if k > 0 else cw
                
                col_left = col_bin_x - col_width // 2
                col_right = col_bin_x + col_width // 2
                
                # Tính overlap
                overlap_left = max(x_left, col_left)
                overlap_right = min(x_right, col_right)
                overlap_width = max(0, overlap_right - overlap_left)
                overlap_ratio = overlap_width / max(cw, col_width) if max(cw, col_width) > 0 else 0
                
                # CẢI TIẾN: Penalty nếu cell quá rộng so với cột (có thể span nhiều cột)
                width_ratio = cw / col_width if col_width > 0 else 1.0
                if width_ratio > 2.5:
                    # Cell quá rộng so với cột → có thể span nhiều cột → giảm điểm số
                    overlap_ratio *= 0.5
                
                # Điểm số = overlap_ratio - distance_penalty
                distance = abs(x_center - col_bin_x)
                distance_penalty = distance / (cw + col_width) if (cw + col_width) > 0 else 0
                score = overlap_ratio - distance_penalty * 0.3
                
                col_scores.append((score, k))
            
            # Chọn cột có điểm số cao nhất
            col_scores.sort(reverse=True)
            best_score = col_scores[0][0] if col_scores else -1
            
            # CẢI TIẾN: Nếu điểm số quá thấp (<0.2), có thể cell này không phù hợp với bất kỳ cột nào
            # → có thể là noise hoặc cell quá lớn → skip
            if best_score < 0.2:
                logger.debug(f"Skipping merged_cell {merged_cell_idx}: Điểm số cột quá thấp ({best_score:.2f}), có thể là noise")
                continue
            
            col_idx = col_scores[0][1] if col_scores else 0
            
            # Tương tự cho row: tính overlap với từng hàng
            row_scores = []
            for k in range(num_rows):
                row_bin_y = row_bins[k]
                # Tính overlap giữa cell và hàng
                if k < num_rows - 1:
                    row_height = row_bins[k + 1] - row_bin_y
                else:
                    row_height = row_bin_y - row_bins[k - 1] if k > 0 else ch
                
                row_top = row_bin_y - row_height // 2
                row_bottom = row_bin_y + row_height // 2
                
                # Tính overlap
                overlap_top = max(y_top, row_top)
                overlap_bottom = min(y_bottom, row_bottom)
                overlap_height = max(0, overlap_bottom - overlap_top)
                overlap_ratio = overlap_height / max(ch, row_height) if max(ch, row_height) > 0 else 0
                
                # CẢI TIẾN: Penalty nếu cell quá cao so với hàng (có thể span nhiều hàng)
                height_ratio = ch / row_height if row_height > 0 else 1.0
                if height_ratio > 2.5:
                    # Cell quá cao so với hàng → có thể span nhiều hàng → giảm điểm số
                    overlap_ratio *= 0.5
                
                # Điểm số = overlap_ratio - distance_penalty
                distance = abs(y_center - row_bin_y)
                distance_penalty = distance / (ch + row_height) if (ch + row_height) > 0 else 0
                score = overlap_ratio - distance_penalty * 0.3
                
                row_scores.append((score, k))
            
            # Chọn hàng có điểm số cao nhất
            row_scores.sort(reverse=True)
            best_row_score = row_scores[0][0] if row_scores else -1
            
            # CẢI TIẾN: Nếu điểm số quá thấp (<0.2), có thể cell này không phù hợp với bất kỳ hàng nào
            # → có thể là noise hoặc cell quá lớn → skip
            if best_row_score < 0.2:
                logger.debug(f"Skipping merged_cell {merged_cell_idx}: Điểm số hàng quá thấp ({best_row_score:.2f}), có thể là noise")
                continue
            
            row_idx = row_scores[0][1] if row_scores else 0

            # QUAN TRỌNG: Kiểm tra conflict và xử lý
            # CẢI TIẾN: Kiểm tra xem vị trí này đã được gán cho merged_cell khác chưa
            conflict_key = (row_idx, col_idx)
            has_conflict = conflict_key in cell_assignment_map
            if has_conflict:
                # Vị trí này đã được gán → kiểm tra xem có phải cùng một cell không
                existing_cell_idx = cell_assignment_map[conflict_key]
                existing_cell = merged_cells[existing_cell_idx]
                existing_box = existing_cell["merged_box"]
                ex, ey, ecw, ech = existing_box
                existing_x_center = ex + ecw // 2
                existing_y_center = ey + ech // 2
                
                # Tính overlap giữa hai cells
                overlap_x = max(0, min(x_right, ex + ecw) - max(x_left, ex))
                overlap_y = max(0, min(y_bottom, ey + ech) - max(y_top, ey))
                overlap_area = overlap_x * overlap_y
                current_area = cw * ch
                existing_area = ecw * ech
                overlap_ratio = overlap_area / min(current_area, existing_area) if min(current_area, existing_area) > 0 else 0
                
                # Nếu overlap > 50% → có thể là cùng một cell → merge
                # Nếu không → tìm vị trí khác
                if overlap_ratio < 0.5:
                    # Đây là cell khác → tìm vị trí trống phù hợp
                    # Ưu tiên 1: Tìm cột trống trong cùng hàng có overlap tốt nhất
                    best_empty_col = None
                    best_col_score = -1
                    for c_idx in range(num_cols):
                        if not table_cells[row_idx][c_idx]["text"] and (row_idx, c_idx) not in cell_assignment_map:
                            col_bin_x = col_bins[c_idx]
                            # Tính overlap với cột này
                            if x_left <= col_bin_x <= x_right:
                                # Có overlap → điểm số cao
                                score = 1.0
                            else:
                                # Không overlap → điểm số thấp hơn (dựa trên khoảng cách)
                                dist = abs(x_center - col_bin_x)
                                score = max(0, 1.0 - dist / (cw * 2))
                            
                            if score > best_col_score:
                                best_col_score = score
                                best_empty_col = c_idx
                    
                    # Ưu tiên 2: Nếu không có cột trống tốt trong cùng hàng, tìm hàng trống trong cùng cột
                    if best_empty_col is None or best_col_score < 0.3:
                        best_empty_row = None
                        best_row_score = -1
                        for r_idx in range(num_rows):
                            if not table_cells[r_idx][col_idx]["text"] and (r_idx, col_idx) not in cell_assignment_map:
                                row_bin_y = row_bins[r_idx]
                                # Tính overlap với hàng này
                                if y_top <= row_bin_y <= y_bottom:
                                    score = 1.0
                                else:
                                    dist = abs(y_center - row_bin_y)
                                    score = max(0, 1.0 - dist / (ch * 2))
                                
                                if score > best_row_score:
                                    best_row_score = score
                                    best_empty_row = r_idx
                        
                        if best_empty_row is not None and best_row_score >= 0.3:
                            row_idx = best_empty_row
                        elif best_empty_col is not None:
                            col_idx = best_empty_col
                        else:
                            # Không tìm được vị trí tốt → skip cell này
                            logger.debug(f"Skipping merged_cell {merged_cell_idx}: Không tìm được vị trí trống phù hợp")
                            continue
                    else:
                        col_idx = best_empty_col
            
            # Đảm bảo row_idx và col_idx hợp lệ
            if row_idx < 0 or row_idx >= num_rows:
                logger.debug(f"Row index {row_idx} out of range [0, {num_rows}), skipping cell")
                continue
            if col_idx < 0 or col_idx >= num_cols:
                logger.debug(f"Col index {col_idx} out of range [0, {num_cols}), skipping cell")
                continue
            
            entry = table_cells[row_idx][col_idx]
            
            # CẢI TIẾN: Kiểm tra xem vị trí này đã có cell khác chưa
            # Nếu có và không phải cùng cell → đã được xử lý ở trên (tìm vị trí trống)
            # Nếu có và là cùng cell → merge text
            if combined_text:
                combined_text_normalized = " ".join(combined_text.split())
                
                if entry["text"]:
                    # Kiểm tra xem đây có phải là cùng một cell không
                    existing_x_center = None
                    existing_y_center = None
                    existing_x_left = None
                    existing_x_right = None
                    existing_y_top = None
                    existing_y_bottom = None
                    
                    if entry["cell_text_parts"]:
                        for part in entry["cell_text_parts"]:
                            if "x" in part and "w" in part:
                                px = part["x"]
                                pw = part["w"]
                                if existing_x_left is None or px < existing_x_left:
                                    existing_x_left = px
                                if existing_x_right is None or px + pw > existing_x_right:
                                    existing_x_right = px + pw
                                existing_x_center = px + pw // 2
                            if "y" in part and "h" in part:
                                py = part["y"]
                                ph = part["h"]
                                if existing_y_top is None or py < existing_y_top:
                                    existing_y_top = py
                                if existing_y_bottom is None or py + ph > existing_y_bottom:
                                    existing_y_bottom = py + ph
                                existing_y_center = py + ph // 2
                    
                    # Tính overlap để xác định có phải cùng cell không
                    if existing_x_left is not None and existing_x_right is not None and existing_y_top is not None and existing_y_bottom is not None:
                        # Tính overlap
                        overlap_x = max(0, min(x_right, existing_x_right) - max(x_left, existing_x_left))
                        overlap_y = max(0, min(y_bottom, existing_y_bottom) - max(y_top, existing_y_top))
                        overlap_area = overlap_x * overlap_y
                        current_area = cw * ch
                        existing_area = (existing_x_right - existing_x_left) * (existing_y_bottom - existing_y_top)
                        overlap_ratio = overlap_area / min(current_area, existing_area) if min(current_area, existing_area) > 0 else 0
                        
                        # Nếu overlap > 50% → cùng một cell → merge text
                        if overlap_ratio >= 0.5:
                            entry["text"] = entry["text"] + " " + combined_text_normalized
                            entry["cell_text_parts"].extend(cell_text_parts_list)
                        else:
                            # Khác cell → không nên xảy ra vì đã xử lý ở trên, nhưng log để debug
                            logger.debug(f"Conflict: Cell at ({x_center}, {y_center}) conflicts with existing, overlap={overlap_ratio:.2f}")
                            # Giữ cell có text dài hơn hoặc diện tích lớn hơn
                            existing_area_sum = sum(part.get("w", 0) * part.get("h", 0) for part in entry["cell_text_parts"])
                            if len(combined_text_normalized) > len(entry["text"]) * 1.2 or current_area > existing_area_sum * 1.2:
                                entry["text"] = combined_text_normalized
                                entry["cell_text_parts"] = cell_text_parts_list
                    else:
                        # Không có metadata cũ → thay thế
                        entry["text"] = combined_text_normalized
                        entry["cell_text_parts"] = cell_text_parts_list
                else:
                    # Ô trống → gán trực tiếp
                    entry["text"] = combined_text_normalized
                    entry["cell_text_parts"] = cell_text_parts_list
            else:
                # Không có text nhưng vẫn lưu metadata (để đánh dấu vị trí cell)
                if not entry["cell_text_parts"]:
                    entry["cell_text_parts"] = cell_text_parts_list
            
            # Đánh dấu vị trí này đã được sử dụng (sau khi đã gán thành công)
            cell_assignment_map[(row_idx, col_idx)] = merged_cell_idx

        # Chỉ giữ lại những hàng có nội dung
        # QUAN TRỌNG: Giữ nguyên vị trí cột để tránh tịnh tiến dữ liệu khi có ô trống
        # CẢI TIẾN: Filter các hàng có quá nhiều ô trống ở đầu (có thể là noise)
        table_data = []
        metadata_cells = []
        row_index_map = {}
        for original_row_idx, row_entries in enumerate(table_cells):
            # Giữ nguyên vị trí cột: row_values[col_idx] = entry[col_idx]["text"]
            # Điều này đảm bảo ô trống ở giữa không làm dịch chuyển các ô bên phải
            row_values = []
            for col_idx in range(num_cols):
                if col_idx < len(row_entries):
                    cell_text = row_entries[col_idx]["text"]
                    # CẢI TIẾN: Cleanup text trước khi thêm vào
                    if cell_text:
                        # Loại bỏ các ký tự đặc biệt lạ ở đầu/cuối
                        cell_text = cell_text.strip()
                        
                        # Đếm các ký tự hợp lệ và invalid
                        vietnamese_chars = 'àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđĐ'
                        noise_chars = '°‹Ÿ¬`^«£ÀÊSẼẴaAaPẨ€+—_*•…"\''
                        invalid_chars = noise_chars + '~…®©œƒŠšŸŒŽžÀÁÂÃÄÅÆÇÈÉÊËÌÍÎÏÐÑÒÓÔÕÖØÙÚÛÜÝÞßàáâãäåæçèéêëìíîïðñòóôõöøùúûüýþÿ'
                        
                        alnum_count = sum(1 for c in cell_text if c.isalnum() or c in vietnamese_chars)
                        invalid_count = sum(1 for c in cell_text if c in invalid_chars)
                        
                        # Nếu text chỉ chứa ký tự đặc biệt lạ → coi như trống
                        if len(cell_text) > 0:
                            alnum_ratio = alnum_count / len(cell_text)
                            invalid_ratio = invalid_count / len(cell_text)
                            
                            # Quá ít ký tự hợp lệ (<15%) → coi như trống
                            if alnum_ratio < 0.15:
                                cell_text = ""
                            # Hoặc có quá nhiều ký tự invalid (>20%) và ít alphanumeric (<30%)
                            elif invalid_ratio > 0.2 and alnum_ratio < 0.3:
                                cell_text = ""
                            # Hoặc với cell ngắn (<15 ký tự), nếu có nhiều invalid (>15%) → coi như trống
                            elif len(cell_text) < 15 and invalid_ratio > 0.15:
                                cell_text = ""
                    row_values.append(cell_text)
                else:
                    # Nếu thiếu cột → thêm ô trống (không làm dịch chuyển)
                    row_values.append("")
            
            # CẢI TIẾN: Filter các hàng có quá nhiều ô trống ở đầu (>= 50% số cột)
            # và chỉ có ít nội dung → có thể là noise
            non_empty_count = sum(1 for val in row_values if val and val.strip())
            empty_prefix_count = sum(1 for val in row_values[:num_cols//2] if not val or not val.strip())
            
            # Nếu có >= 50% cột đầu trống và chỉ có <= 1 ô có nội dung → có thể là noise
            if empty_prefix_count >= num_cols * 0.5 and non_empty_count <= 1:
                logger.debug(f"Skipping row {original_row_idx}: Quá nhiều ô trống ở đầu ({empty_prefix_count}/{num_cols//2}), chỉ có {non_empty_count} ô có nội dung")
                continue
            
            # CẢI TIẾN: Filter các hàng có quá nhiều ký tự đặc biệt lạ (noise từ OCR)
            # Kiểm tra tổng hợp toàn bộ hàng
            row_text = " ".join(row_values)
            vietnamese_chars = 'àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđĐ'
            noise_chars = '°‹Ÿ¬`^«£ÀÊSẼẴaAaPẨ€+—_*•…"\''
            # Mở rộng danh sách noise chars để bao gồm các ký tự không hợp lệ khác
            invalid_chars = noise_chars + '~…®©œƒŠšŸŒŽžÀÁÂÃÄÅÆÇÈÉÊËÌÍÎÏÐÑÒÓÔÕÖØÙÚÛÜÝÞßàáâãäåæçèéêëìíîïðñòóôõöøùúûüýþÿ'
            
            alnum_count = sum(1 for c in row_text if c.isalnum() or c in vietnamese_chars)
            noise_count = sum(1 for c in row_text if c in noise_chars)
            invalid_count = sum(1 for c in row_text if c in invalid_chars)
            
            # Đếm các ký tự hợp lệ (alphanumeric + dấu câu thông thường + khoảng trắng)
            valid_punctuation = '.,;:!?()[]{}-/\\'
            valid_count = alnum_count + sum(1 for c in row_text if c in valid_punctuation or c.isspace())
            
            # CẢI TIẾN: Áp dụng filter cho cả hàng ngắn (>= 5 ký tự)
            if len(row_text) >= 5:
                noise_ratio = noise_count / len(row_text) if len(row_text) > 0 else 0
                alnum_ratio = alnum_count / len(row_text) if len(row_text) > 0 else 0
                invalid_ratio = invalid_count / len(row_text) if len(row_text) > 0 else 0
                valid_ratio = valid_count / len(row_text) if len(row_text) > 0 else 0
                
                # CẢI TIẾN: Filter các hàng ngắn có nhiều ký tự đặc biệt
                if len(row_text) < 30:
                    # Với hàng ngắn, threshold thấp hơn vì dễ bị nhiễu
                    # Filter 1: Invalid ratio cao (>15%)
                    if invalid_ratio > 0.15:
                        logger.debug(f"Skipping row {original_row_idx}: Hàng ngắn có nhiều ký tự invalid ({invalid_ratio:.2%})")
                        continue
                    # Filter 2: Noise ratio cao (>10%) và alphanumeric thấp (<50%)
                    if noise_ratio > 0.1 and alnum_ratio < 0.5:
                        logger.debug(f"Skipping row {original_row_idx}: Hàng ngắn có nhiều noise ({noise_ratio:.2%}), alphanumeric thấp ({alnum_ratio:.2%})")
                        continue
                    # Filter 3: Noise ratio cao (>8%) và alphanumeric thấp (<60%)
                    if noise_ratio > 0.08 and alnum_ratio < 0.6:
                        logger.debug(f"Skipping row {original_row_idx}: Hàng ngắn có nhiều noise ({noise_ratio:.2%}), alphanumeric thấp ({alnum_ratio:.2%})")
                        continue
                    # Filter 4: Invalid ratio cao (>12%) và valid ratio thấp (<80%)
                    if invalid_ratio > 0.12 and valid_ratio < 0.8:
                        logger.debug(f"Skipping row {original_row_idx}: Hàng ngắn có nhiều invalid ({invalid_ratio:.2%}), valid thấp ({valid_ratio:.2%})")
                        continue
                
                # CẢI TIẾN: Filter mạnh hơn - nếu có nhiều ký tự invalid (>20%) và ít valid (<60%)
                if invalid_ratio > 0.2 and valid_ratio < 0.6:
                    logger.debug(f"Skipping row {original_row_idx}: Quá nhiều ký tự invalid ({invalid_ratio:.2%}), ít valid ({valid_ratio:.2%})")
                    continue
                
                if noise_ratio > 0.3 and alnum_ratio < 0.2:
                    logger.debug(f"Skipping row {original_row_idx}: Quá nhiều ký tự noise ({noise_ratio:.2%}), ít alphanumeric ({alnum_ratio:.2%})")
                    continue
                
                # Nếu hàng có quá ít alphanumeric (<15%) và có nhiều ký tự đặc biệt
                if alnum_ratio < 0.15 and noise_count > 5:
                    logger.debug(f"Skipping row {original_row_idx}: Quá ít alphanumeric ({alnum_ratio:.2%}), nhiều ký tự đặc biệt ({noise_count})")
                    continue
                
                # CẢI TIẾN: Nếu có nhiều ký tự đặc biệt (>15%) và tỷ lệ valid thấp (<50%)
                # Điều này sẽ catch các hàng như "|||'‹ ° Cu|toi PÓ Tô Ÿ CC ST 4ì||"
                if invalid_ratio > 0.15 and valid_ratio < 0.5:
                    logger.debug(f"Skipping row {original_row_idx}: Nhiều ký tự invalid ({invalid_ratio:.2%}), tỷ lệ valid thấp ({valid_ratio:.2%})")
                    continue
                
                # CẢI TIẾN: Nếu có nhiều ký tự invalid (>18%) và có nhiều ký tự đặc biệt không hợp lệ
                # Ngay cả khi tỷ lệ valid cao, nếu có quá nhiều ký tự đặc biệt thì vẫn là noise
                if invalid_ratio > 0.18 and invalid_count > 10:
                    logger.debug(f"Skipping row {original_row_idx}: Nhiều ký tự invalid ({invalid_ratio:.2%}, {invalid_count} ký tự)")
                    continue
                
                # CẢI TIẾN: Nếu có nhiều ký tự noise (>12%) và tỷ lệ alphanumeric không cao (<40%)
                # Điều này sẽ catch các hàng có nhiều ký tự đặc biệt nhưng vẫn có một số chữ cái
                if noise_ratio > 0.12 and alnum_ratio < 0.4 and len(row_text) > 20:
                    logger.debug(f"Skipping row {original_row_idx}: Nhiều ký tự noise ({noise_ratio:.2%}), alphanumeric thấp ({alnum_ratio:.2%})")
                    continue
                
                # CẢI TIẾN: Filter các hàng có quá nhiều ký tự đặc biệt không hợp lệ ngay cả khi tỷ lệ alphanumeric cao
                # Ví dụ: "||||toi PÓ Tô Ÿ CC ST 4ì||" có alnum_ratio cao nhưng có nhiều ký tự đặc biệt
                if invalid_count > 5 and invalid_ratio > 0.12:
                    # Nếu có nhiều ký tự invalid và tỷ lệ valid không đủ cao (<70%)
                    if valid_ratio < 0.7:
                        logger.debug(f"Skipping row {original_row_idx}: Nhiều ký tự invalid ({invalid_count}, {invalid_ratio:.2%}), valid ratio thấp ({valid_ratio:.2%})")
                        continue
                
                # CẢI TIẾN: Filter các hàng có noise ratio cao (>15%) ngay cả khi alphanumeric ratio trung bình
                # Ví dụ: "|T:€=+ X2 «—t P1 ' 7E 0 tư ƯỜNG|||||" có noise_ratio 20%
                if noise_ratio > 0.15 and alnum_ratio < 0.55:
                    logger.debug(f"Skipping row {original_row_idx}: Noise ratio cao ({noise_ratio:.2%}), alphanumeric thấp ({alnum_ratio:.2%})")
                    continue
                
                # CẢI TIẾN: Filter các hàng có invalid ratio cao (>18%) ngay cả khi valid ratio cao
                # Điều này catch các hàng có nhiều ký tự đặc biệt không hợp lệ
                if invalid_ratio > 0.18:
                    logger.debug(f"Skipping row {original_row_idx}: Invalid ratio quá cao ({invalid_ratio:.2%})")
                    continue
            
            # Chỉ thêm hàng nếu có ít nhất một ô có nội dung hợp lệ
            if any(val.strip() for val in row_values):
                new_row_idx = len(table_data)
                row_index_map[original_row_idx] = new_row_idx
                table_data.append(row_values)
                # Lưu metadata cho các ô có nội dung
                for col_idx, entry in enumerate(row_entries):
                    if col_idx < len(row_values) and entry.get("cell_text_parts"):
                        metadata_cells.append({
                            "row": new_row_idx,
                            "col": col_idx,
                            "cell_text_parts": entry["cell_text_parts"]
                        })

        # Lưu table_data vào kết quả (không xuất CSV trung gian)
        if table_data:
            # Đảm bảo tất cả các hàng có cùng số cột (pad với chuỗi rỗng nếu thiếu)
            max_cols = max(len(row) for row in table_data) if table_data else num_cols
            if max_cols == 0:
                max_cols = num_cols  # Fallback về num_cols từ grid detection
            
            # Chuẩn hóa dữ liệu bảng
            normalized_table = []
            for row in table_data:
                normalized_row = []
                for col_idx in range(max_cols):
                    if col_idx < len(row):
                        cell = row[col_idx]
                        if cell:
                            # Ghép paragraph trong ô thành một paragraph
                            normalized_cell = " ".join(cell.split())
                            normalized_row.append(normalized_cell)
                        else:
                            normalized_row.append("")
                    else:
                        # Pad với chuỗi rỗng nếu thiếu cột
                        normalized_row.append("")
                normalized_table.append(normalized_row)
            
            tables_by_page[i] = {
                "page": i,
                "rows": normalized_table,
                "num_cols": max_cols,
                "metadata": metadata_cells
            }
            exported += 1
            logger.info(f"🗂️  Đã extract bảng (OpenCV+OCR) trang {i}: {len(normalized_table)} hàng, {max_cols} cột")
        else:
            logger.info(f"Trang {i}: Phát hiện bảng nhưng không có nội dung text sau OCR")

    if exported == 0:
        logger.info("ℹ️  OpenCV fallback không tìm thấy bảng nào")
        return {}
    else:
        logger.info(f"✅ OpenCV fallback: Đã extract {exported} bảng từ {len(tables_by_page)} trang")
        return tables_by_page

def convert_pdf_to_docx(pdf_path: str, output_path: str, pages: Optional[List[int]] = None) -> str:
    """
    Convert PDF → DOCX trực tiếp bằng pdf2docx.
    
    Args:
        pdf_path: Đường dẫn file PDF input
        output_path: Đường dẫn file DOCX output
        pages: Danh sách số trang cần convert (1-indexed). None = tất cả trang.
    
    Returns:
        str: Đường dẫn file DOCX đã tạo
    
    Raises:
        RuntimeError: Nếu pdf2docx chưa được cài đặt
    """
    if Converter is None:
        raise RuntimeError("pdf2docx chưa được cài đặt. Cài pdf2docx để dùng convert PDF → DOCX.")
    
    logger.info(f"📄 Đang convert PDF → DOCX: {pdf_path}")
    
    try:
        cv = Converter(pdf_path)
        
        # pdf2docx hỗ trợ pages thông qua start_page và end_page
        # Nếu có pages chỉ định, convert chỉ những trang đó
        if pages:
            valid_pages = sorted(set(pages))
            logger.info(f"Chỉ convert {len(valid_pages)} trang: {valid_pages}")
            
            # Convert từ start đến end của pages
            start_page = min(valid_pages)
            end_page = max(valid_pages)
            cv.convert(output_path, start=start_page, end=end_page)
        else:
            cv.convert(output_path)
        
        cv.close()
        
        # Validate output file
        if not os.path.exists(output_path):
            raise RuntimeError(f"Conversion thất bại: File output không tồn tại: {output_path}")
        
        file_size = os.path.getsize(output_path)
        if file_size < 100:  # DOCX tối thiểu phải > 100 bytes
            raise RuntimeError(f"Conversion thất bại: File output quá nhỏ ({file_size} bytes)")
        
        logger.info(f"✅ Đã convert PDF → DOCX: {output_path} ({file_size} bytes)")
        return output_path
        
    except AttributeError as attr_error:
        # Lỗi do version không tương thích (ví dụ: 'Rect' object has no attribute 'get_area')
        error_msg = str(attr_error)
        logger.error(f"❌ Lỗi version không tương thích: {error_msg}")
        logger.warning("💡 Lỗi do pdf2docx và PyMuPDF không tương thích.")
        logger.warning("💡 Giải pháp: pip install PyMuPDF==1.26.4")
        raise RuntimeError(f"Lỗi version không tương thích (pdf2docx/PyMuPDF): {error_msg}")
    except Exception as e:
        logger.error(f"❌ Lỗi khi convert PDF → DOCX: {e}")
        import traceback
        logger.debug(traceback.format_exc())
        raise


def update_paragraph_in_place(para, new_text: str):
    """
    Update paragraph text nhưng giữ nguyên formatting từ original runs.
    
    Args:
        para: python-docx Paragraph object
        new_text: Text mới để replace
    """
    if not para.runs:
        # No runs → add new text với default formatting
        para.add_run(new_text)
        return
    
    # Strategy: Giữ formatting từ first run (hoặc most common)
    first_run = para.runs[0]
    
    # Save formatting
    font_format = {
        "bold": first_run.font.bold,
        "italic": first_run.font.italic,
        "size": first_run.font.size,
        "name": first_run.font.name
    }
    
    # Clear paragraph
    para.clear()
    
    # Add new text với formatting
    run = para.add_run(new_text)
    
    # Apply formatting
    if font_format["bold"]:
        run.font.bold = True
    if font_format["italic"]:
        run.font.italic = True
    if font_format["size"]:
        run.font.size = font_format["size"]
    if font_format["name"]:
        run.font.name = font_format["name"]


def re_insert_images_to_paragraph(para, images: List[dict]):
    """
    Re-insert images vào paragraph sau khi process text.
    
    Args:
        para: python-docx Paragraph object
        images: List of image dicts từ extract_images_from_paragraph()
    """
    import io
    
    if not images:
        return
    
    for img_info in images:
        try:
            image_bytes = img_info.get("image_data")
            if not image_bytes or len(image_bytes) < 10:
                continue
            
            width_inches = img_info.get("width", 4.0)
            height_inches = img_info.get("height", 3.0)
            
            # Add image vào paragraph
            run = para.add_run()
            img_stream = io.BytesIO(image_bytes)
            img_stream.seek(0)
            
            try:
                if width_inches > 0.1 and height_inches > 0.1:
                    # Use width (maintain aspect ratio)
                    max_width = 6.0  # Max 6 inches
                    run.add_picture(img_stream, width=Inches(min(width_inches, max_width)))
                else:
                    run.add_picture(img_stream, width=Inches(4.0))
            except Exception as pic_error:
                logger.warning(f"Không thể re-insert image vào paragraph: {pic_error}")
                continue
                
        except Exception as e:
            logger.warning(f"Lỗi khi re-insert image: {e}")
            continue


def split_batched_result(batched_text: str, original_count: int) -> List[str]:
    """
    Split batched result về số paragraphs ban đầu (estimate).
    
    Args:
        batched_text: Text đã được process (có thể có nhiều paragraphs)
        original_count: Số paragraphs ban đầu
    
    Returns:
        List[str]: List các paragraphs (có thể không đúng số lượng)
    """
    # Simple: Split by double newlines
    paragraphs = batched_text.split('\n\n')
    paragraphs = [p.strip() for p in paragraphs if p.strip()]
    
    # Nếu số paragraphs không match → distribute đều hoặc keep như cũ
    if len(paragraphs) != original_count:
        # Estimate: Nếu có ít hơn, có thể AI đã merge
        # Nếu có nhiều hơn, có thể AI đã split
        # Tạm thời return như cũ
        pass
    
    return paragraphs


def update_docx_with_processed_text(docx_path: str, processed_paragraphs: List[dict], ocr_cfg: dict) -> str:
    """
    Update DOCX với processed text, giữ nguyên formatting và images.
    
    Args:
        docx_path: Đường dẫn file DOCX input/output (sẽ được update in-place)
        processed_paragraphs: List các processed paragraph dicts
        ocr_cfg: Config dictionary
    
    Returns:
        str: Đường dẫn file DOCX đã update
    """
    if Document is None:
        raise RuntimeError("python-docx chưa được cài đặt.")
    
    logger.info(f"🔄 Đang update DOCX với processed text: {docx_path}")
    
    doc = Document(docx_path)
    
    # Process từng paragraph/batch
    para_idx = 0
    
    for processed in processed_paragraphs:
        if processed["type"] == "batch":
            # Batch: Split result và update từng paragraph
            cleaned_text = processed.get("cleaned_text", processed["text"])
            spell_checked_text = processed.get("spell_checked_text", cleaned_text)
            
            # Split batched result
            para_objects = processed["para_objects"]
            images_list = processed.get("images_list", [])
            
            split_paragraphs = split_batched_result(spell_checked_text, len(para_objects))
            
            # Update từng paragraph trong batch
            for i, para_obj in enumerate(para_objects):
                if i < len(split_paragraphs):
                    para_text = split_paragraphs[i]
                else:
                    # Not enough split → use remaining text hoặc empty
                    para_text = "" if i > 0 else spell_checked_text
                
                # Update paragraph
                if para_text.strip():
                    update_paragraph_in_place(para_obj, para_text)
                    
                    # Re-insert images từ paragraph này
                    if i < len(images_list):
                        re_insert_images_to_paragraph(para_obj, images_list[i])
                else:
                    # Empty paragraph → clear
                    para_obj.clear()
            
        else:
            # Single paragraph
            para_obj = processed["para_object"]
            cleaned_text = processed.get("cleaned_text", processed["text"])
            spell_checked_text = processed.get("spell_checked_text", cleaned_text)
            images = processed.get("images", [])
            
            # Update text
            if spell_checked_text.strip():
                update_paragraph_in_place(para_obj, spell_checked_text)
                
                # Re-insert images
                re_insert_images_to_paragraph(para_obj, images)
            else:
                # Empty → clear
                para_obj.clear()
        
        # Handle merge
        if processed.get("should_merge_with_next") and para_idx + 1 < len(processed_paragraphs):
            # Merge với paragraph sau: update current với merged text, clear next
            next_processed = processed_paragraphs[para_idx + 1]
            merged_text = spell_checked_text + " " + next_processed.get("spell_checked_text", next_processed["text"])
            
            # Update current paragraph với merged text
            if processed["type"] == "single":
                update_paragraph_in_place(processed["para_object"], merged_text)
            elif processed["type"] == "batch" and processed["para_objects"]:
                # Update last paragraph trong batch
                update_paragraph_in_place(processed["para_objects"][-1], merged_text)
            
            # Clear next paragraph
            if next_processed["type"] == "single":
                next_processed["para_object"].clear()
            elif next_processed["type"] == "batch" and next_processed["para_objects"]:
                # Clear first paragraph trong batch
                next_processed["para_objects"][0].clear()
        
        para_idx += 1
    
    # Save updated DOCX
    try:
        doc.save(docx_path)
        logger.info(f"✅ Đã update DOCX: {docx_path}")
        return docx_path
    except Exception as e:
        logger.error(f"❌ Không thể save updated DOCX: {e}")
        raise


def convert_docx_to_epub(docx_path: str, epub_path: str, ocr_cfg: dict) -> str:
    """
    Convert DOCX → EPUB using pypandoc.
    
    Args:
        docx_path: Đường dẫn file DOCX input
        epub_path: Đường dẫn file EPUB output
        ocr_cfg: Config dictionary
    
    Returns:
        str: Đường dẫn file EPUB đã tạo
    
    Raises:
        RuntimeError: Nếu pypandoc chưa được cài đặt hoặc conversion fail
    """
    try:
        import pypandoc
    except ImportError:
        raise RuntimeError("pypandoc chưa được cài đặt. Cài pypandoc để convert DOCX → EPUB.")
    
    logger.info(f"📚 Đang convert DOCX → EPUB: {docx_path}")
    
    try:
        pypandoc.convert_file(
            docx_path,
            'epub',
            outputfile=epub_path,
            extra_args=['--standalone']
        )
        
        # Validate output
        if not os.path.exists(epub_path):
            raise RuntimeError(f"Conversion thất bại: File EPUB không tồn tại: {epub_path}")
        
        file_size = os.path.getsize(epub_path)
        if file_size < 1000:  # EPUB tối thiểu > 1KB
            raise RuntimeError(f"Conversion thất bại: File EPUB quá nhỏ ({file_size} bytes)")
        
        logger.info(f"✅ Đã convert DOCX → EPUB: {epub_path} ({file_size} bytes)")
        return epub_path
        
    except Exception as e:
        logger.error(f"❌ Lỗi khi convert DOCX → EPUB: {e}")
        import traceback
        logger.debug(traceback.format_exc())
        raise
def hybrid_workflow_pdf_to_docx(pdf_path: str, output_path: str, ocr_cfg: dict, pages: Optional[List[int]] = None) -> str:
    """
    Hybrid workflow: PDF → DOCX → Cleanup & Spell Check → DOCX
    
    Args:
        pdf_path: Đường dẫn file PDF input
        output_path: Đường dẫn file DOCX output
        ocr_cfg: Config dictionary
        pages: Danh sách số trang cần process (1-indexed). None = tất cả trang.
    
    Returns:
        str: Đường dẫn file DOCX output
    """
    logger.info("=" * 80)
    logger.info("🚀 BẮT ĐẦU HYBRID WORKFLOW")
    logger.info("=" * 80)
    
    # Step 1: Convert PDF → DOCX
    temp_docx_path = output_path.replace(".docx", "_temp.docx")
    try:
        convert_pdf_to_docx(pdf_path, temp_docx_path, pages)
    except Exception as e:
        logger.error(f"❌ Không thể convert PDF → DOCX: {e}")
        raise
    
    # Step 2: Extract paragraphs với hints
    try:
        paragraphs_data = extract_paragraphs_with_hints(temp_docx_path)
        if not paragraphs_data:
            raise RuntimeError("Không extract được paragraphs từ DOCX")
    except Exception as e:
        logger.error(f"❌ Không thể extract paragraphs: {e}")
        raise
    
    # Step 3: Batch small paragraphs
    batched_paragraphs = batch_small_paragraphs(paragraphs_data, min_chars=50)
    
    # Step 4: Process từng paragraph/batch (Cleanup + Spell Check)
    processed_paragraphs = []
    
    logger.info(f"🔄 Đang process {len(batched_paragraphs)} paragraphs/batches...")
    for idx, para_data in enumerate(batched_paragraphs, 1):
        logger.info(f"Processing {idx}/{len(batched_paragraphs)}...")
        
        # Cleanup
        cleanup_result = cleanup_paragraph_with_hints(para_data, ocr_cfg)
        para_data["cleaned_text"] = cleanup_result["cleaned_text"]
        para_data["should_merge_with_next"] = cleanup_result.get("should_merge_with_next", False)
        
        # Spell check
        spell_checked_text = spell_check_paragraph(para_data, ocr_cfg)
        para_data["spell_checked_text"] = spell_checked_text
        
        processed_paragraphs.append(para_data)
    
    # Step 5: Update DOCX với processed text
    try:
        update_docx_with_processed_text(temp_docx_path, processed_paragraphs, ocr_cfg)
    except Exception as e:
        logger.error(f"❌ Không thể update DOCX: {e}")
        raise
    
    # Step 6: Move temp file to final output
    try:
        if os.path.exists(output_path):
            os.remove(output_path)
        os.rename(temp_docx_path, output_path)
        logger.info(f"✅ Đã tạo DOCX: {output_path}")
    except Exception as e:
        logger.error(f"❌ Không thể move file: {e}")
        raise
    
    # Step 7: Cleanup intermediate files
    _cleanup_intermediate_files(output_path)
    
    # Return DOCX path
    return output_path


def extract_text_and_images_from_pdf(pdf_path: str, ocr_cfg: dict, pages: Optional[List[int]] = None) -> tuple[List[dict], int]:
    """
    Extract text và images từ PDF có text layer.
    
    Args:
        pdf_path: Đường dẫn file PDF
        ocr_cfg: Config dictionary
        pages: Danh sách số trang cần extract (1-indexed). None = tất cả trang.
    
    Returns:
        tuple: (pages_data, total_pages) trong đó pages_data là list of dict với keys:
            - page_num: số trang (1-indexed)
            - text: text content
            - images: list of image data (bytes) với position info
    """
    if fitz is None:
        raise RuntimeError("PyMuPDF (fitz) chưa được cài đặt. Cài PyMuPDF để hỗ trợ extract images.")
    
    pages_data: List[dict] = []
    
    try:
        doc = fitz.open(pdf_path)
        total = len(doc)
        logger.info(f"Extract text và images: Tổng số trang: {total}")
        
        # Filter pages nếu có chỉ định
        if pages:
            valid_pages = [p for p in pages if 1 <= p <= total]
            invalid_pages = [p for p in pages if p < 1 or p > total]
            if invalid_pages:
                logger.warning(f"Các trang không hợp lệ (nằm ngoài 1-{total}): {invalid_pages}. Bỏ qua.")
            if not valid_pages:
                logger.error("Không có trang hợp lệ nào để extract.")
                return [], 0
            logger.info(f"Extract text và images: Chỉ extract {len(valid_pages)} trang: {valid_pages}")
            pages_to_extract = sorted(set(valid_pages))
        else:
            pages_to_extract = list(range(1, total + 1))
        
        show_progress = bool(ocr_cfg.get("show_progress", True))
        
        if show_progress and tqdm is not None and len(pages_to_extract) > 1:
            iterator = tqdm(pages_to_extract, desc="Extract text & images", unit="trang")
        else:
            iterator = pages_to_extract
        
        for page_num in iterator:
            page = doc[page_num - 1]  # fitz dùng 0-indexed
            
            # Extract text
            text = page.get_text().strip()
            
            # Extract images với position
            images = []
            image_list = page.get_images(full=True)
            page_rect = page.rect  # Page dimensions
            
            for img_idx, img in enumerate(image_list):
                try:
                    # img là tuple: (xref, smask, width, height, bpc, colorspace, alt. colorspace, name, filter, referencer)
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # Extract image position từ PDF
                    # PyMuPDF không trực tiếp cho image position, nhưng ta có thể estimate từ image rect
                    # Hoặc dùng image_list với position nếu có
                    y_position = 0  # Default, sẽ được cập nhật nếu có thông tin
                    x_position = 0
                    
                    # Thử extract image rect từ page (nếu có)
                    try:
                        # Get image rectangles từ page
                        image_rects = page.get_image_rects(xref)
                        if image_rects:
                            # Lấy rect đầu tiên (thường chỉ có 1)
                            rect = image_rects[0]
                            y_position = rect.y0  # Top Y position
                            x_position = rect.x0  # Left X position
                    except Exception:
                        # Estimate position: images thường ở giữa hoặc theo thứ tự
                        y_position = page_rect.height * 0.3 + (img_idx * page_rect.height * 0.2)
                        x_position = page_rect.width * 0.5  # Center
                    
                    images.append({
                        "data": image_bytes,
                        "ext": image_ext,
                        "xref": xref,
                        "width": img[2],
                        "height": img[3],
                        "y_position": y_position,
                        "x_position": x_position
                    })
                except Exception as e:
                    logger.debug(f"Không thể extract image {img_idx} từ trang {page_num}: {e}")
                    continue
            
            pages_data.append({
                "page_num": page_num,
                "text": text,
                "images": images
            })
            
            if not show_progress and len(pages_data) % 50 == 0:
                logger.info(f"Extract text và images: {len(pages_data)}/{len(pages_to_extract)} trang")
        
        doc.close()
        return pages_data, total
        
    except Exception as e:
        logger.error(f"PyMuPDF failed: {e}")
        raise


def create_docx_from_processed_text(pdf_path: str, output_path: str, processed_text: str, ocr_cfg: dict, pages: Optional[List[int]] = None) -> str:
    """
    Tạo file DOCX từ text đã được xử lý (cleanup/spell check) và images từ PDF.
    
    Args:
        pdf_path: Đường dẫn file PDF input (để extract images)
        output_path: Đường dẫn file DOCX output
        processed_text: Text đã được cleanup và spell check
        ocr_cfg: Config dictionary
        pages: Danh sách số trang cần extract (1-indexed). None = tất cả trang.
    
    Returns:
        str: Đường dẫn file DOCX đã tạo
    """
    if Document is None:
        raise RuntimeError("python-docx chưa được cài đặt. Cài python-docx để tạo DOCX output.")
    
    logger.info(f"📄 Tạo DOCX từ text đã xử lý và images: {pdf_path}")
    
    # Extract images từ PDF (không cần text vì đã có processed_text)
    pages_data, total_pages = extract_text_and_images_from_pdf(pdf_path, ocr_cfg, pages)
    
    if not pages_data:
        raise ValueError("Không có dữ liệu nào được extract từ PDF.")
    
    # Tạo images_dict từ pages_data và collect all_images
    images_dict = {}
    all_images = []
    for page_info in pages_data:
        page_num = page_info["page_num"]
        images_dict[page_num] = page_info.get("images", [])
        all_images.extend([(page_num, img_info) for img_info in page_info["images"]])
    
    # Kiểm tra xem có sử dụng HTML intermediate workflow không
    use_html_workflow = ocr_cfg.get("use_html_intermediate_workflow", True)  # Default: True
    
    # Nếu có images và use_html_workflow, thử dùng HTML workflow để đảm bảo vị trí chính xác
    if use_html_workflow and all_images and len(all_images) > 0:
        logger.info("📄 Sử dụng HTML intermediate workflow để đảm bảo vị trí text và images chính xác...")
        
        # Extract text blocks với Y-position từ PDF gốc để có position chính xác
        logger.info("🔍 Đang extract text blocks với position từ PDF...")
        text_blocks_by_page, _ = extract_text_blocks_with_position(pdf_path, ocr_cfg, pages)
        
        # Tạo all_items_with_position: kết hợp text paragraphs và images với position chính xác
        all_items_with_position = []
        
        # Chia processed_text thành paragraphs (giữ thứ tự)
        text_paragraphs = processed_text.split('\n\n')
        text_paragraphs = [p.strip() for p in text_paragraphs if p.strip()]
        
        # Map processed paragraphs với original text blocks dựa trên similarity
        # Strategy: Match processed paragraphs với original text blocks bằng fuzzy matching
        processed_para_idx = 0
        
        for page_info in pages_data:
            page_num = page_info["page_num"]
            original_blocks = text_blocks_by_page.get(page_num, [])
            
            # Nếu có original blocks với position, map processed paragraphs với chúng
            if original_blocks:
                # Match processed paragraphs với original blocks
                # Simple strategy: giả định thứ tự paragraphs được giữ nguyên sau cleanup/spell check
                # Có thể cải thiện bằng fuzzy matching nếu cần
                
                # Estimate số paragraphs per block
                if len(original_blocks) > 0:
                    # Chia processed paragraphs đều cho các blocks
                    paras_per_block = max(1, len(text_paragraphs) // max(1, sum(len(text_blocks_by_page.get(p, [])) for p in pages_data if pages_data)))
                else:
                    paras_per_block = len(text_paragraphs)
                
                block_idx = 0
                for original_block in original_blocks:
                    # Tìm processed paragraph tương ứng (simple sequential matching)
                    if processed_para_idx < len(text_paragraphs):
                        # Có thể cải thiện bằng fuzzy matching
                        para_text = text_paragraphs[processed_para_idx]
                        processed_para_idx += 1
                        
                        all_items_with_position.append({
                            "type": "text",
                            "content": para_text,
                            "page_num": page_num,
                            "y_position": original_block.get("y_position", 0),
                            "x_position": original_block.get("x_position", 0)
                        })
                    
                    # Thêm images nằm sau block này (nếu có)
                    # Images với Y-position < block Y-position đã được xử lý
                    block_y = original_block.get("y_position", 0)
                    
                    # Tìm images của trang này có Y-position gần block này
                    for img_info in page_info.get("images", []):
                        img_y = img_info.get("y_position", 0)
                        img_x = img_info.get("x_position", 0)
                        
                        # Nếu image có Y-position trong khoảng hợp lý với block này, thêm vào
                        # (Images thường nằm giữa các text blocks)
                        # Chỉ thêm nếu chưa được thêm (check bằng xref)
                        if "xref" in img_info:
                            # Check xem image này đã được thêm chưa
                            already_added = any(
                                item.get("type") == "image" and 
                                item.get("img_info", {}).get("xref") == img_info["xref"]
                                for item in all_items_with_position
                            )
                            
                            if not already_added:
                                # Estimate: nếu image Y gần block Y, thêm ngay sau block
                                if abs(img_y - block_y) < 200:  # 200px threshold
                                    all_items_with_position.append({
                                        "type": "image",
                                        "img_info": img_info,
                                        "page_num": page_num,
                                        "y_position": img_y,
                                        "x_position": img_x
                                    })
            else:
                # Không có original blocks với position → dùng estimate cũ
                # Chia processed paragraphs đều cho các trang
                paragraphs_per_page = max(1, len(text_paragraphs) // len(pages_data)) if pages_data else len(text_paragraphs)
                start_idx = (page_num - (pages_data[0]["page_num"] if pages_data else 1)) * paragraphs_per_page
                end_idx = min(start_idx + paragraphs_per_page, len(text_paragraphs))
                
                for para_text in text_paragraphs[start_idx:end_idx]:
                    all_items_with_position.append({
                        "type": "text",
                        "content": para_text,
                        "page_num": page_num,
                        "y_position": 100 * (page_num - (pages_data[0]["page_num"] if pages_data else 1)) * 50,  # Estimate
                        "x_position": 0
                    })
            
            # Thêm các images còn lại của trang này (chưa được thêm)
            for img_info in page_info.get("images", []):
                if "xref" in img_info:
                    already_added = any(
                        item.get("type") == "image" and 
                        item.get("img_info", {}).get("xref") == img_info["xref"]
                        for item in all_items_with_position
                    )
                    
                    if not already_added:
                        y_position = img_info.get("y_position", 500 * page_num)
                        x_position = img_info.get("x_position", 0)
                        
                        all_items_with_position.append({
                            "type": "image",
                            "img_info": img_info,
                            "page_num": page_num,
                            "y_position": y_position,
                            "x_position": x_position
                        })
        
        # Sort theo (y_position, x_position) để đảm bảo thứ tự đúng
        all_items_with_position.sort(key=lambda x: (x.get("page_num", 0) * 10000 + x.get("y_position", 0), x.get("x_position", 0)))
        
        logger.info(f"📊 Đã thu thập {len([i for i in all_items_with_position if i['type'] == 'text'])} text blocks và {len([i for i in all_items_with_position if i['type'] == 'image'])} images với position")
        
        # Thử tạo DOCX qua HTML workflow
        try:
            html_path = _create_html_from_items(all_items_with_position, output_path)
            success = _convert_html_to_docx_with_pandoc(html_path, output_path, ocr_cfg)
            
            if success:
                logger.info(f"✅ Đã tạo DOCX qua HTML intermediate workflow: {output_path}")
                # Cleanup HTML temp file
                try:
                    if os.path.exists(html_path):
                        os.remove(html_path)
                        logger.debug(f"Đã xóa file HTML temp: {html_path}")
                except Exception:
                    pass
                return output_path
            else:
                logger.warning("⚠️ HTML → DOCX conversion thất bại, fallback về python-docx trực tiếp...")
        except Exception as html_error:
            logger.warning(f"⚠️ HTML intermediate workflow thất bại: {html_error}")
            logger.warning("⚠️ Fallback về python-docx trực tiếp...")
            import traceback
            logger.debug(traceback.format_exc())
    
    # Fallback: Dùng python-docx trực tiếp (workflow cũ)
    # Tạo DOCX document
    try:
        doc = Document()
    except Exception as e:
        raise RuntimeError(f"Không thể tạo Document object: {e}")
    
    # Set document properties (optional)
    try:
        doc.core_properties.title = os.path.splitext(os.path.basename(pdf_path))[0]
    except Exception:
        pass
    import io
    show_progress = bool(ocr_cfg.get("show_progress", True))
    
    # Bước 1: Chèn tất cả images từ tất cả các trang (giữ nguyên thứ tự)
    images_added_count = 0
    if all_images:
        logger.info(f"🖼️  Đang chèn {len(all_images)} images vào DOCX...")
        for page_num, img_info in (tqdm(all_images, desc="Chèn images", unit="ảnh") if (show_progress and tqdm and len(all_images) > 1) else all_images):
            try:
                image_bytes = img_info.get("data")
                if not image_bytes or len(image_bytes) == 0:
                    logger.warning(f"Image data rỗng từ trang {page_num}, bỏ qua")
                    continue
                
                if len(image_bytes) < 10:
                    logger.warning(f"Image từ trang {page_num} quá nhỏ ({len(image_bytes)} bytes), bỏ qua")
                    continue
                
                image_ext = img_info.get("ext", "png")
                
                # Validate image data - kiểm tra magic bytes
                is_valid = False
                magic_msg = ""
                if image_ext.lower() in ("jpeg", "jpg"):
                    if len(image_bytes) >= 2 and image_bytes[:2] == b'\xff\xd8':
                        is_valid = True
                    else:
                        magic_msg = f"Magic bytes: {image_bytes[:2].hex() if len(image_bytes) >= 2 else 'too short'} (expected: ffd8)"
                elif image_ext.lower() == "png":
                    if len(image_bytes) >= 8 and image_bytes[:8] == b'\x89PNG\r\n\x1a\n':
                        is_valid = True
                    else:
                        magic_msg = f"Magic bytes: {image_bytes[:8].hex() if len(image_bytes) >= 8 else 'too short'} (expected: 89504e47...)"
                elif image_ext.lower() in ("gif", "bmp", "tiff", "webp"):
                    is_valid = True
                else:
                    # Format khác, chấp nhận nhưng log warning
                    logger.debug(f"Image từ trang {page_num} có format không chuẩn: {image_ext}")
                    is_valid = True
                
                if not is_valid:
                    logger.warning(f"Image từ trang {page_num} không phải {image_ext.upper()} hợp lệ. {magic_msg}")
                    continue
                
                # Log thông tin image trước khi thử add
                logger.debug(f"Thử thêm image từ trang {page_num}: ext={image_ext}, size={len(image_bytes)} bytes, width={img_width}, height={img_height}")
                
                # Kiểm tra magic bytes của image data
                if len(image_bytes) >= 2:
                    first_bytes = image_bytes[:8] if len(image_bytes) >= 8 else image_bytes[:2]
                    logger.debug(f"Image magic bytes: {first_bytes.hex()}")
                
                # Thử load image với PIL để validate trước khi add vào DOCX
                try:
                    from PIL import Image as PILImage
                    pil_img = PILImage.open(io.BytesIO(image_bytes))
                    pil_img.verify()  # Verify image integrity
                    pil_img.close()
                    logger.debug(f"✅ Image từ trang {page_num} đã được validate bằng PIL")
                except Exception as pil_error:
                    logger.warning(f"⚠️  Image từ trang {page_num} không thể validate bằng PIL: {pil_error}")
                    # Vẫn thử add vào DOCX vì có thể PIL không hỗ trợ format này nhưng python-docx có thể
                
                # Tạo image từ bytes
                img_stream = io.BytesIO(image_bytes)
                img_stream.seek(0)
                
                # Validate stream có data (đã check ở trên nhưng double-check)
                if len(image_bytes) == 0:
                    logger.warning(f"Image bytes rỗng từ trang {page_num}")
                    continue
                
                # Thêm image vào document
                para = doc.add_paragraph()
                run = para.add_run()
                
                # Tính toán kích thước image (giữ tỷ lệ)
                img_width = img_info.get("width", 0)
                img_height = img_info.get("height", 0)
                
                try:
                    if img_width > 0 and img_height > 0:
                        max_width_inches = 6.0
                        aspect_ratio = img_height / img_width
                        
                        if img_width > 500:
                            width_inches = min(max_width_inches, img_width / 96.0)
                            height_inches = width_inches * aspect_ratio
                        else:
                            width_inches = img_width / 96.0
                            height_inches = img_height / 96.0
                        
                        if width_inches > 0.1 and height_inches > 0.1:
                            run.add_picture(img_stream, width=Inches(min(width_inches, max_width_inches)))
                            images_added_count += 1
                        else:
                            run.add_picture(img_stream, width=Inches(4.0))
                            images_added_count += 1
                    else:
                        run.add_picture(img_stream, width=Inches(4.0))
                        images_added_count += 1
                    
                    para_format = para.paragraph_format
                    para_format.space_after = Pt(6)
                except Exception as pic_error:
                    import traceback
                    error_details = traceback.format_exc()
                    error_msg = str(pic_error)
                    logger.warning(f"❌ Không thể add picture vào DOCX từ trang {page_num}: {error_msg}")
                    logger.warning(f"   Image info: ext={image_ext}, size={len(image_bytes)} bytes, width={img_width}, height={img_height}")
                    # Log full traceback ở debug level
                    if logger.isEnabledFor(logging.DEBUG):
                        logger.debug(f"Chi tiết lỗi add_picture:\n{error_details}")
                    continue
                
            except Exception as e:
                logger.warning(f"Không thể thêm image vào DOCX từ trang {page_num}: {e}")
                continue
    
    if all_images and images_added_count == 0:
        logger.warning(f"Không có image nào được thêm vào DOCX (tổng {len(all_images)} images)")
    
    # Bước 2: Chèn text đã được xử lý (processed_text)
    text_added = False
    if processed_text and processed_text.strip():
        logger.info("📝 Đang chèn text đã được xử lý vào DOCX...")
        paragraphs = processed_text.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                para = doc.add_paragraph(para_text)
                para_format = para.paragraph_format
                para_format.space_after = Pt(6)
                text_added = True
    
    # Đảm bảo document có nội dung
    has_content = text_added or images_added_count > 0
    
    if not has_content:
        logger.warning("Document không có nội dung. Thêm paragraph mặc định...")
        try:
            doc.add_paragraph("(Không có nội dung từ PDF)")
        except Exception:
            pass
    
    # Save DOCX
    try:
        doc.save(output_path)
        logger.info(f"✅ Đã tạo DOCX: {output_path}")
        logger.info(f"   📊 Thống kê: {images_added_count} images, {len(processed_text.split())} từ")
        return output_path
    except Exception as e:
        logger.error(f"❌ Không thể lưu DOCX: {e}")
        raise


def create_docx_from_pdf(pdf_path: str, output_path: str, ocr_cfg: dict, pages: Optional[List[int]] = None, apply_cleanup: bool = True, apply_spell_check: bool = True) -> str:
    """
    Tạo file DOCX từ PDF có text layer, giữ lại cả text và images.
    Có thể áp dụng cleanup và spell check cho text trước khi tạo DOCX.
    
    LƯU Ý: Hàm này được giữ lại để backward compatibility, nhưng workflow mới nên dùng:
    - extract_text_and_images_from_pdf() để extract
    - ocr_file() để xử lý text
    - create_docx_from_processed_text() để tạo DOCX
    
    Args:
        pdf_path: Đường dẫn file PDF input
        output_path: Đường dẫn file DOCX output
        ocr_cfg: Config dictionary
        pages: Danh sách số trang cần extract (1-indexed). None = tất cả trang.
        apply_cleanup: Có áp dụng AI cleanup cho text không (mặc định: True)
        apply_spell_check: Có áp dụng AI spell check cho text không (mặc định: True)
    
    Returns:
        str: Đường dẫn file DOCX đã tạo
    """
    if Document is None:
        raise RuntimeError("python-docx chưa được cài đặt. Cài python-docx để tạo DOCX output.")
    
    logger.info(f"Tạo DOCX từ PDF: {pdf_path}")
    
    # Extract text và images
    pages_data, total_pages = extract_text_and_images_from_pdf(pdf_path, ocr_cfg, pages)
    
    if not pages_data:
        raise ValueError("Không có dữ liệu nào được extract từ PDF.")
    
    # Extract images với thông tin vị trí Y để chèn đúng vị trí
    # Sử dụng images từ pages_data (đã có trong extract_text_and_images_from_pdf)
    # Tạo images_dict từ pages_data để có structure tương thích
    images_dict = {}
    for page_info in pages_data:
        page_num = page_info["page_num"]
        images_dict[page_num] = page_info.get("images", [])
    
    # Xử lý text: ghép text từ tất cả các trang và apply cleanup/spell check nếu cần
    all_text = "\n\n".join([page_info["text"] for page_info in pages_data])
    processed_text = all_text  # Mặc định là text gốc
    
    if apply_cleanup or apply_spell_check:
        logger.info("Đang xử lý text với AI (cleanup và spell check)...")
        
        # Apply cleanup nếu enabled
        cleanup_cfg = ocr_cfg.get("ai_cleanup", {})
        if apply_cleanup and cleanup_cfg.get("enabled", False):
            logger.info("🧹 Đang chạy AI Cleanup...")
            result = ai_cleanup_text(all_text, ocr_cfg)
            if isinstance(result, tuple):
                all_text, cleanup_failed_indices, cleanup_original_chunks = result
                cleanup_failed = len(cleanup_failed_indices)
                if cleanup_failed > 0:
                    logger.warning(f"AI Cleanup: {cleanup_failed} chunks failed")
            else:
                all_text = result
            logger.info("✅ Hoàn tất AI Cleanup")
        
        # Apply spell check nếu enabled
        spell_check_cfg = ocr_cfg.get("ai_spell_check", {})
        if apply_spell_check and spell_check_cfg.get("enabled", False):
            logger.info("✍️  Đang chạy AI Spell Check...")
            result = ai_spell_check_and_paragraph_restore(all_text, ocr_cfg)
            if isinstance(result, tuple):
                all_text, spell_check_failed_indices, spell_check_original_chunks = result
                spell_check_failed = len(spell_check_failed_indices)
                if spell_check_failed > 0:
                    logger.warning(f"AI Spell Check: {spell_check_failed} chunks failed")
            else:
                all_text = result
            logger.info("✅ Hoàn tất AI Spell Check")
        
        # Sau khi xử lý, text đã được cleanup và spell check
        # Ta sẽ lưu toàn bộ text đã xử lý và chèn vào DOCX sau tất cả images
        # Để giữ liên kết giữa images và text, ta sẽ:
        # 1. Chèn images từ tất cả các trang (giữ nguyên thứ tự)
        # 2. Chèn toàn bộ text đã xử lý sau tất cả images
        processed_text = all_text
    
    # Tạo DOCX document
    try:
        doc = Document()
    except Exception as e:
        raise RuntimeError(f"Không thể tạo Document object: {e}")
    
    # Set document properties (optional)
    try:
        doc.core_properties.title = os.path.splitext(os.path.basename(pdf_path))[0]
    except Exception:
        pass
    
    # python-docx tự động tạo một paragraph trống khi khởi tạo Document()
    # Ta sẽ để nó như vậy và chỉ thêm nội dung khi cần
    
    show_progress = bool(ocr_cfg.get("show_progress", True))
    total_items = len(pages_data)
    
    if show_progress and tqdm is not None:
        iterator = tqdm(pages_data, desc="Tạo DOCX", unit="trang")
    else:
        iterator = pages_data
    
    import io
    
    # Nếu có text đã được xử lý (cleanup/spell check), ta sẽ chèn text sau tất cả images
    has_processed_text = apply_cleanup or apply_spell_check
    
    # Bước 1: Chèn tất cả images từ tất cả các trang (giữ nguyên thứ tự)
    all_images = []
    for page_info in pages_data:
        all_images.extend([(page_info["page_num"], img_info) for img_info in page_info["images"]])
    
    # Chèn images
    images_added_count = 0
    if all_images:
        for page_num, img_info in (tqdm(all_images, desc="Chèn images", unit="ảnh") if (show_progress and tqdm and len(all_images) > 1) else all_images):
            try:
                image_bytes = img_info.get("data")
                if not image_bytes or len(image_bytes) == 0:
                    logger.warning(f"Image data rỗng từ trang {page_num}, bỏ qua")
                    continue
                
                if len(image_bytes) < 10:  # Image quá nhỏ, có thể không hợp lệ
                    logger.warning(f"Image từ trang {page_num} quá nhỏ ({len(image_bytes)} bytes), bỏ qua")
                    continue
                
                image_ext = img_info.get("ext", "png")
                
                # Validate image data - kiểm tra magic bytes
                is_valid = False
                if image_ext.lower() in ("jpeg", "jpg"):
                    if len(image_bytes) >= 2 and image_bytes[:2] == b'\xff\xd8':
                        is_valid = True
                elif image_ext.lower() == "png":
                    if len(image_bytes) >= 8 and image_bytes[:8] == b'\x89PNG\r\n\x1a\n':
                        is_valid = True
                elif image_ext.lower() in ("gif", "bmp", "tiff", "webp"):
                    # Chấp nhận các format khác mà không validate magic bytes (để python-docx xử lý)
                    is_valid = True
                else:
                    # Format khác, thử add anyway
                    is_valid = True
                
                if not is_valid:
                    logger.warning(f"Image từ trang {page_num} không phải {image_ext.upper()} hợp lệ, bỏ qua")
                    continue
                
                # Tạo image từ bytes
                img_stream = io.BytesIO(image_bytes)
                img_stream.seek(0)  # Đảm bảo stream ở đầu file
                
                # Thêm image vào document
                # Giới hạn kích thước để vừa với page width
                para = doc.add_paragraph()
                run = para.add_run()
                
                # Tính toán kích thước image (giữ tỷ lệ)
                img_width = img_info.get("width", 0)
                img_height = img_info.get("height", 0)
                
                try:
                    if img_width > 0 and img_height > 0:
                        # Giới hạn width tối đa là 6 inches (khoảng 15cm)
                        max_width_inches = 6.0
                        aspect_ratio = img_height / img_width
                        
                        if img_width > 500:  # Nếu image lớn, scale xuống
                            width_inches = min(max_width_inches, img_width / 96.0)  # Giả định 96 DPI
                            height_inches = width_inches * aspect_ratio
                        else:
                            width_inches = img_width / 96.0
                            height_inches = img_height / 96.0
                        
                        # Đảm bảo kích thước hợp lệ (min 0.1 inches, max 7 inches)
                        if width_inches > 0.1 and height_inches > 0.1:
                            run.add_picture(img_stream, width=Inches(min(width_inches, max_width_inches)))
                            images_added_count += 1
                        else:
                            run.add_picture(img_stream, width=Inches(4.0))
                            images_added_count += 1
                    else:
                        # Nếu không có thông tin size, dùng default
                        run.add_picture(img_stream, width=Inches(4.0))
                        images_added_count += 1
                    
                    # Thêm spacing sau image
                    para_format = para.paragraph_format
                    para_format.space_after = Pt(6)
                except Exception as pic_error:
                    # Nếu add_picture thất bại, xóa paragraph trống
                    logger.warning(f"Không thể add picture vào DOCX từ trang {page_num}: {pic_error}")
                    # Không cần xóa paragraph vì python-docx sẽ xử lý
                    continue
                
            except Exception as e:
                logger.warning(f"Không thể thêm image vào DOCX từ trang {page_num}: {e}")
                import traceback
                logger.debug(traceback.format_exc())
                continue
    
    if all_images and images_added_count == 0:
        logger.warning(f"Không có image nào được thêm vào DOCX (tổng {len(all_images)} images)")
    
    # Bước 2: Chèn text đã được xử lý (nếu có) hoặc text gốc từ từng trang
    text_added = False
    if has_processed_text:
        # Chèn toàn bộ text đã được xử lý sau tất cả images
        if processed_text and processed_text.strip():
            logger.info("Đang chèn text đã được xử lý vào DOCX...")
            # Chia text thành paragraphs (dựa trên double newlines)
            paragraphs = processed_text.split('\n\n')
            
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    para = doc.add_paragraph(para_text)
                    para_format = para.paragraph_format
                    para_format.space_after = Pt(6)
                    text_added = True
    else:
        # Không có xử lý → chèn text gốc từ từng trang cùng với images
        for page_info in iterator:
            page_num = page_info["page_num"]
            text = page_info["text"]
            
            # Thêm text của trang này
            if text and text.strip():
                # Chia text thành paragraphs (dựa trên double newlines)
                paragraphs = text.split('\n\n')
                
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        para = doc.add_paragraph(para_text)
                        para_format = para.paragraph_format
                        para_format.space_after = Pt(6)
                        text_added = True
            
            # Thêm page break sau mỗi trang (trừ trang cuối) - chỉ khi có text hoặc images
            if page_num < pages_data[-1]["page_num"]:
                doc.add_page_break()
            
            if not show_progress and (page_num % 50 == 0 or page_num == total_items):
                logger.info(f"Đã xử lý {page_num}/{total_items} trang")
    
    # Đảm bảo document có ít nhất một paragraph hợp lệ (nếu không có gì cả)
    # Kiểm tra xem có nội dung gì không (images hoặc text)
    has_content = text_added or images_added_count > 0
    
    if not has_content:
        logger.warning("Document không có nội dung (không có text và images). Thêm paragraph mặc định...")
        try:
            # Thêm paragraph có nội dung
            # Không xóa paragraph trống vì có thể gây lỗi cấu trúc DOCX
            doc.add_paragraph("(Không có nội dung từ PDF)")
        except Exception as e:
            logger.warning(f"Không thể thêm paragraph mặc định: {e}")
            # Fallback: thử thêm vào paragraph đầu tiên nếu có
            try:
                if len(doc.paragraphs) > 0:
                    doc.paragraphs[0].text = "(Không có nội dung)"
                else:
                    doc.add_paragraph("(Không có nội dung)")
            except Exception:
                pass
    
    # Validate document trước khi save
    try:
        # Kiểm tra xem document có hợp lệ không
        # Document phải có ít nhất một element (paragraph với text hoặc images được embed)
        total_elements = len(doc.paragraphs)
        if total_elements == 0:
            logger.warning("Document trống, thêm paragraph mặc định")
            doc.add_paragraph("(Không có nội dung)")
        else:
            # Kiểm tra xem có paragraph nào có nội dung không (text hoặc runs - có thể chứa images)
            has_any_content = False
            for para in doc.paragraphs:
                # Kiểm tra text
                if para.text.strip():
                    has_any_content = True
                    break
                # Kiểm tra runs (có thể chứa images hoặc inline shapes)
                if len(para.runs) > 0:
                    # Nếu có runs, giả định là có nội dung (images hoặc text)
                    has_any_content = True
                    break
            
            if not has_any_content:
                logger.warning("Tất cả paragraphs đều trống, thêm paragraph mặc định")
                # Sử dụng paragraph đầu tiên nếu có, hoặc tạo mới
                try:
                    if len(doc.paragraphs) > 0:
                        # Thêm text vào paragraph đầu tiên
                        para = doc.paragraphs[0]
                        if not para.text.strip():
                            para.add_run("(Không có nội dung)")
                        else:
                            doc.add_paragraph("(Không có nội dung)")
                    else:
                        doc.add_paragraph("(Không có nội dung)")
                except Exception:
                    # Fallback: chỉ tạo paragraph mới
                    try:
                        doc.add_paragraph("(Không có nội dung)")
                    except Exception:
                        pass
    except Exception as e:
        logger.warning(f"Không thể validate document: {e}")
    
    # Lưu file
    logger.info(f"Đang lưu DOCX: {output_path}")
    try:
        doc.save(output_path)
        logger.info(f"✅ Đã tạo DOCX thành công: {output_path}")
        
        # Validate file sau khi save (kiểm tra file size)
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            if file_size == 0:
                raise ValueError("File DOCX có kích thước 0 bytes - không hợp lệ")
            logger.info(f"📄 File size: {file_size:,} bytes")
    except Exception as e:
        logger.error(f"❌ Lỗi khi lưu DOCX: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise
    
    return output_path
def _fix_docx_leading_tabs_and_soft_wraps(docx_path: str) -> None:
    """
    Khắc phục các tab thừa đầu dòng do soft-wrap khi convert PDF→DOCX (pdf2docx).
    Quy tắc an toàn (không xâm phạm layout nhiều):
    - Nếu một paragraph bắt đầu bằng tab (\t) và ký tự đầu tiên có nghĩa sau tab là chữ thường/không phải số/bullet,
      và paragraph trước đó KHÔNG kết thúc bằng dấu câu (., !, ?), thì loại bỏ các tab/space đầu paragraph đó.
    - Không merge/xóa paragraph để tránh rủi ro layout; chỉ loại bỏ tab đầu dòng gây xấu văn bản.
    """
    try:
        if Document is None:
            return
        doc = Document(docx_path)
        prev_text = ""
        for para in doc.paragraphs:
            full_text = para.text or ""
            try:
                import re
            except Exception:
                re = None
            
            # Bỏ qua đoạn có URL để tránh phá hyperlink/format
            if re and re.search(r"https?://\S+", full_text):
                prev_text = full_text
                continue
            
            # 1) Loại tab nội tuyến trong từng run để giữ hyperlink và format
            for run in getattr(para, 'runs', []):
                if not run.text:
                    continue
                if re:
                    new_run_text = re.sub(r"\s*\t+\s*", " ", run.text)
                else:
                    new_run_text = run.text.replace("\t", " ")
                if new_run_text != run.text:
                    run.text = new_run_text
            
            # Cập nhật lại full_text sau bước (1)
            full_text = para.text or ""
            stripped = full_text.lstrip("\t ")
            starts_with_tab = (full_text != stripped)
            prev_ends_with_punct = bool(re.search(r"[.!?]$", prev_text.strip())) if (re and prev_text) else False
            is_bullet_like = bool(re.match(r"^[•·\-*]\s", stripped)) or bool(re.match(r"^\d+[.)]\s", stripped)) if re else False
            
            # 2) Loại tab/space đầu đoạn (continuation của câu trước), chỉ tác động lên runs đầu
            if starts_with_tab and not prev_ends_with_punct and not is_bullet_like:
                remaining_to_strip = len(full_text) - len(stripped)
                # Bỏ qua nếu stripping làm rỗng hoàn toàn (an toàn)
                if remaining_to_strip > 0 and stripped:
                    for run in getattr(para, 'runs', []):
                        if remaining_to_strip <= 0:
                            break
                        if not run.text:
                            continue
                        run_len = len(run.text)
                        # Tính số ký tự whitespace đầu run có thể cắt
                        prefix = 0
                        while prefix < run_len and remaining_to_strip > 0 and run.text[prefix] in ('\t', ' '):
                            prefix += 1
                            remaining_to_strip -= 1
                        if prefix > 0:
                            run.text = run.text[prefix:]
                        # Nếu run không còn whitespace đầu và vẫn còn remaining_to_strip, tiếp tục sang run kế
            
            prev_text = para.text or ""
        doc.save(docx_path)
    except Exception:
        # Không chặn pipeline nếu chỉnh sửa thất bại
        pass


async def _cleanup_chunk_async(chunk: str, api_key: str, model_name: str, prompt: str, chunk_idx: int, total_chunks: int, timeout_s: float, safety_settings: Optional[List[dict]] = None) -> str:
    """
    Cleanup một chunk text bằng AI (async).
    
    Args:
        safety_settings: Optional safety settings để pass vào GenerativeModel (nếu None sẽ dùng default)
    """
    # Suppress logs TRƯỚC khi import
    _suppress_google_logs()
    # Đảm bảo stderr filter đang active
    if not isinstance(sys.stderr, NoisyMessageFilter):
        original_stderr = sys.stderr if not isinstance(sys.stderr, NoisyMessageFilter) else getattr(sys.stderr, 'original_stream', sys.stderr)
        sys.stderr = NoisyMessageFilter(original_stderr)
    import google.generativeai as genai
    genai.configure(api_key=api_key)
    # Pass safety_settings vào GenerativeModel (nếu có)
    model = genai.GenerativeModel(model_name, safety_settings=safety_settings) if safety_settings else genai.GenerativeModel(model_name)
    
    # Run trong thread pool và áp timeout để tránh treo vô hạn
    loop = asyncio.get_event_loop()
    response = await asyncio.wait_for(
        loop.run_in_executor(
            None,
            lambda: model.generate_content(prompt + chunk)
        ),
        timeout=timeout_s
    )
    
    # Kiểm tra response có hợp lệ không
    if not response or not response.candidates or len(response.candidates) == 0:
        raise ValueError(f"AI cleanup chunk {chunk_idx}/{total_chunks}: No candidates returned")
    
    # Kiểm tra prompt_feedback nếu có
    if hasattr(response, 'prompt_feedback') and response.prompt_feedback:
        if hasattr(response.prompt_feedback, 'block_reason') and response.prompt_feedback.block_reason:
            raise ValueError(f"AI cleanup chunk {chunk_idx}/{total_chunks}: Blocked by safety filter: {response.prompt_feedback.block_reason}")
    
    result = response.text.strip()
    return result


def _format_table_with_coordinates(metadata_path: str) -> str:
    """
    Format bảng với marker tọa độ để AI xử lý.
    
    Format: [CELL R=row C=col]text_part_1[SEP]text_part_2[END CELL]
    
    Args:
        metadata_path: Đường dẫn file JSON metadata
        
    Returns:
        str: Text đã format với marker tọa độ
    """
    try:
        import json
        with open(metadata_path, "r", encoding="utf-8") as f:
            metadata = json.load(f)
        
        cells = metadata.get("cells", [])
        if not cells:
            return ""
        
        # Nhóm theo (row, col)
        cell_dict = {}
        for cell in cells:
            key = (cell["row"], cell["col"])
            if key not in cell_dict:
                cell_dict[key] = []
            cell_dict[key].extend(cell["cell_text_parts"])
        
        # Format: [CELL R=row C=col]text1[SEP]text2[END CELL]
        formatted_lines = []
        for (row, col), parts in sorted(cell_dict.items()):
            cell_texts = [part["text"].strip() for part in parts if part.get("text", "").strip()]
            if cell_texts:
                combined = " ".join(cell_texts)  # Ghép các phần text trong cùng ô
                formatted_lines.append(f"[CELL R={row} C={col}]{combined}[END CELL]")
        
        return "\n".join(formatted_lines)
    except Exception as e:
        logger.warning(f"Không thể format bảng với tọa độ: {e}")
        return ""


def _format_table_row_with_markers(row: List[str], num_cols: int) -> str:
    """
    Format một hàng bảng với marker | phân cách giữa các ô.
    
    Đảm bảo:
    - Mỗi hàng có đúng num_cols cột (pad với chuỗi rỗng nếu thiếu)
    - Marker | chỉ ở giữa các ô, không có ở đầu và cuối
    - Format: cell1|cell2|cell3
    - Xử lý các trường hợp đặc biệt: None, empty list, non-string values
    
    Args:
        row: List các ô trong hàng (có thể chứa None hoặc non-string)
        num_cols: Số cột mong muốn
        
    Returns:
        str: Hàng đã format với marker | (ví dụ: "cell1|cell2|cell3")
    """
    # Đảm bảo row là list
    if not isinstance(row, list):
        row = []
    
    # Convert tất cả cell thành string và loại bỏ None
    row_str_list = [str(cell) if cell is not None else "" for cell in row]
    
    # Đảm bảo hàng có đúng num_cols cột
    if len(row_str_list) < num_cols:
        # Pad với chuỗi rỗng nếu thiếu
        row_str_list = row_str_list + [""] * (num_cols - len(row_str_list))
    elif len(row_str_list) > num_cols:
        # Cắt bớt nếu thừa (chỉ lấy num_cols cột đầu)
        row_str_list = row_str_list[:num_cols]
    
    # Format với marker | (không có | ở đầu và cuối)
    # Đảm bảo không có cell nào là None hoặc không phải string
    result = "|".join(row_str_list)
    
    # Debug log nếu cần
    if not result or "|" not in result:
        logger.debug(f"⚠️  Format table row: row={row}, num_cols={num_cols}, result='{result}'")
    
    return result


def _update_csv_with_cleaned_cells(csv_path: str, cleaned_cells: List[dict]) -> None:
    """
    Cập nhật CSV với text đã được AI cleanup.
    Format: cell1|cell2|cell3 (marker | phân cách cột)
    
    Args:
        csv_path: Đường dẫn file CSV
        cleaned_cells: List các dict {row, col, cleaned_text}
    """
    try:
        # Đọc CSV hiện tại (format: cell1|cell2|cell3)
        rows = []
        with open(csv_path, "r", encoding="utf-8", newline="") as f:
            for line in f:
                line = line.rstrip('\n\r')
                if line:
                    cells = line.split('|')
                    rows.append(cells)
        
        # Tạo dict mapping (row, col) -> cleaned_text
        cell_map = {(cell["row"], cell["col"]): cell["cleaned_text"] for cell in cleaned_cells}
        
        # Cập nhật các ô có trong cell_map
        for row_idx, row in enumerate(rows):
            for col_idx in range(len(row)):
                if (row_idx, col_idx) in cell_map:
                    # Ghép paragraph trong ô thành một paragraph
                    cleaned_text = cell_map[(row_idx, col_idx)]
                    rows[row_idx][col_idx] = " ".join(cleaned_text.split())
        
        # Ghi lại CSV với format marker |
        # Tìm số cột tối đa để đảm bảo nhất quán
        max_cols = max(len(row) for row in rows) if rows else 0
        
        with open(csv_path, "w", encoding="utf-8", newline="") as f:
            for row in rows:
                # Format với helper function để đảm bảo đúng số cột
                row_str = _format_table_row_with_markers(row, max_cols)
                f.write(row_str + "\n")
    except Exception as e:
        logger.warning(f"Không thể cập nhật CSV: {e}")


def ai_cleanup_table_with_coordinates(metadata_path: str, ocr_cfg: dict) -> dict:
    """
    Sử dụng AI để cleanup và ghép text trong các ô bảng có nhiều dòng.
    
    Args:
        metadata_path: Đường dẫn file JSON metadata
        ocr_cfg: Config OCR
        
    Returns:
        dict: {"cells": [{row, col, cleaned_text}], "success": bool}
    """
    cleanup_cfg = ocr_cfg.get("ai_cleanup", {})
    if not cleanup_cfg.get("enabled", False):
        return {"cells": [], "success": False}
    
    api_keys = cleanup_cfg.get("api_keys", [])
    if not api_keys:
        api_keys = ocr_cfg.get("_root_api_keys", [])
    if not api_keys:
        logger.warning("AI cleanup enabled nhưng không có API keys. Bỏ qua cleanup bảng.")
        return {"cells": [], "success": False}
    
    try:
        import json
        with open(metadata_path, "r", encoding="utf-8") as f:
            metadata = json.load(f)
        
        # Format bảng với marker tọa độ
        formatted_text = _format_table_with_coordinates(metadata_path)
        if not formatted_text:
            return {"cells": [], "success": False}
        
        # Prompt cho AI
        prompt = """Bạn là AI chuyên xử lý bảng dữ liệu từ OCR.

NHIỆM VỤ:
1. Ghép các phần text có cùng tọa độ (R=row, C=col) thành một ô hoàn chỉnh
2. GHÉP TẤT CẢ PARAGRAPH TRONG CÙNG MỘT Ô THÀNH MỘT PARAGRAPH (ưu tiên tính toàn vẹn nội dung)
3. Loại bỏ ký tự rác, lỗi OCR (ký tự đặc biệt không có nghĩa, ký tự lạ)
4. Chuẩn hóa khoảng trắng (một khoảng trắng giữa các từ)
5. Giữ nguyên số liệu, ngày tháng, đơn vị, dấu phẩy, dấu chấm
6. Sửa lỗi chính tả phổ biến (ví dụ: "chíh" → "chính", "Hanh phúc" → "Hạnh phúc")

ĐỊNH DẠNG ĐẦU VÀO:
[CELL R=row C=col]text_part_1 text_part_2[END CELL]

ĐỊNH DẠNG ĐẦU RA (BẮT BUỘC):
[CELL R=row C=col]text_đã_ghép_và_cleanup[END CELL]

QUY TẮC QUAN TRỌNG:
- BẮT BUỘC giữ nguyên format [CELL R=... C=...]...[END CELL] trong output
- GHÉP TẤT CẢ PARAGRAPH TRONG CÙNG MỘT Ô THÀNH MỘT PARAGRAPH (thay \n bằng space)
- Ưu tiên tính toàn vẹn nội dung của ô hơn là tính chính xác của việc phân paragraph
- Ghép tất cả text trong cùng một ô (cùng R và C) thành một chuỗi liên tục (không có xuống dòng)
- Loại bỏ ký tự đặc biệt không cần thiết (nhưng giữ: số, dấu phẩy, dấu chấm, dấu gạch ngang, dấu ngoặc)
- Chuẩn hóa khoảng trắng (một khoảng trắng giữa các từ, loại bỏ khoảng trắng thừa)
- Sửa lỗi OCR phổ biến: "chíh" → "chính", "Hanh" → "Hạnh", "Tư do" → "Tự do", "đông" → "đồng"

Bảng cần xử lý:
"""
        
        # Gọi AI với API key đầu tiên
        api_key = api_keys[0]
        model_name = cleanup_cfg.get("model", "gemini-2.5-flash")
        timeout_s = float(cleanup_cfg.get("ai_timeout_seconds", 120))
        
        _suppress_google_logs()
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        
        response = model.generate_content(prompt + formatted_text)
        if not response or not response.candidates:
            logger.warning("AI cleanup bảng: Không có response")
            return {"cells": [], "success": False}
        
        # Kiểm tra finish_reason
        candidate = response.candidates[0]
        if hasattr(candidate, 'finish_reason') and candidate.finish_reason:
            if candidate.finish_reason == 1:  # SAFETY hoặc BLOCKED
                logger.warning(f"AI cleanup bảng: Response bị block (finish_reason={candidate.finish_reason})")
                return {"cells": [], "success": False}
        
        # Lấy text từ response
        try:
            cleaned_text = response.text.strip()
        except Exception as text_err:
            logger.warning(f"AI cleanup bảng: Không thể lấy text từ response: {text_err}")
            # Thử lấy từ parts
            if candidate.parts:
                cleaned_text = " ".join([part.text for part in candidate.parts if hasattr(part, 'text') and part.text]).strip()
            else:
                return {"cells": [], "success": False}
        
        # Parse kết quả: tìm [CELL R=... C=...]...[END CELL]
        import re
        # Hỗ trợ cả [END CELL] và [/END CELL], và cả trường hợp không có marker
        # Pattern 1: [CELL R=... C=...]...[END CELL]
        pattern1 = r'\[CELL R=(\d+) C=(\d+)\](.*?)\[END CELL\]'
        matches1 = re.findall(pattern1, cleaned_text, re.DOTALL)
        # Pattern 2: [CELL R=... C=...]...[/END CELL]
        pattern2 = r'\[CELL R=(\d+) C=(\d+)\](.*?)\[/END CELL\]'
        matches2 = re.findall(pattern2, cleaned_text, re.DOTALL)
        # Gộp kết quả, ưu tiên pattern1
        matches = matches1 if matches1 else matches2
        
        cleaned_cells = []
        for row_str, col_str, text in matches:
            try:
                row = int(row_str)
                col = int(col_str)
                cleaned_cells.append({
                    "row": row,
                    "col": col,
                    "cleaned_text": text.strip()
                })
            except ValueError:
                continue
        
        logger.info(f"✅ AI cleanup bảng: Đã xử lý {len(cleaned_cells)} ô")
        return {"cells": cleaned_cells, "success": True}
        
    except Exception as e:
        logger.warning(f"Lỗi khi cleanup bảng bằng AI: {e}")
        import traceback
        logger.debug(traceback.format_exc())
        return {"cells": [], "success": False}


def ai_cleanup_text(text: str, ocr_cfg: dict) -> str:
    """
    Sử dụng AI để dọn rác text (header/footer, vệt đen từ scan, noise...).
    Hỗ trợ nhiều API keys để xử lý song song.
    """
    cleanup_cfg = ocr_cfg.get("ai_cleanup", {})
    cleanup_enabled = cleanup_cfg.get("enabled", False)
    if not cleanup_enabled:
        return text
    
    # Lấy API keys (ưu tiên từ ai_cleanup.api_keys, fallback về api_keys từ root config)
    api_keys = cleanup_cfg.get("api_keys", [])
    if not api_keys:
        # Đọc từ _root_api_keys đã lưu khi load config
        api_keys = ocr_cfg.get("_root_api_keys", [])
    
    if not api_keys:
        logger.warning("AI cleanup enabled nhưng không có API keys. Bỏ qua cleanup.")
        return text
    
    model_name = cleanup_cfg.get("model", "gemini-2.5-flash")
    max_parallel = cleanup_cfg.get("max_parallel_workers", 5)
    # Giới hạn worker theo số API keys sẵn có
    if api_keys:
        max_parallel = max(1, min(max_parallel, len(api_keys)))
    chunk_size = cleanup_cfg.get("chunk_size", 50000)
    delay = cleanup_cfg.get("delay_between_requests", 0.5)
    max_retries = cleanup_cfg.get("max_retries", 3)
    timeout_s = float(cleanup_cfg.get("ai_timeout_seconds", 120))
    show_progress = bool(ocr_cfg.get("show_progress", True))
    progress_interval = float(ocr_cfg.get("progress_log_interval_seconds", 60))
    
    prompt = """Bạn là một AI chuyên dọn dẹp văn bản OCR/scan. Nhiệm vụ:
1. Loại bỏ header/footer lặp lại ở mỗi trang
2. Loại bỏ các ký tự rác, vệt đen vô nghĩa từ quá trình scan
3. Loại bỏ số trang, watermark
4. Giữ nguyên nội dung chính của văn bản
5. Chuẩn hóa khoảng trắng thừa
6. Giữ nguyên định dạng đoạn văn

Trả về chỉ văn bản đã được dọn dẹp, không giải thích thêm.

Văn bản cần dọn dẹp:
"""
    
    try:
        # Chia nhỏ text nếu quá dài
        if len(text) <= chunk_size:
            # Text ngắn, xử lý trực tiếp
            logger.info("AI Cleanup: Text ngắn, xử lý trực tiếp (1 chunk)")
            # Suppress logs TRƯỚC khi import
            _suppress_google_logs()
            # Đảm bảo stderr filter đang active
            if not isinstance(sys.stderr, NoisyMessageFilter):
                original_stderr = sys.stderr if not isinstance(sys.stderr, NoisyMessageFilter) else getattr(sys.stderr, 'original_stream', sys.stderr)
                sys.stderr = NoisyMessageFilter(original_stderr)
            # Build safety settings từ config
            safety_level = ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
            safety_settings = _build_safety_settings(safety_level)
            
            import google.generativeai as genai
            genai.configure(api_key=api_keys[0])
            model = genai.GenerativeModel(model_name, safety_settings=safety_settings)
            response = model.generate_content(prompt + text)
            
            # Kiểm tra nếu response bị block (mặc dù đã set BLOCK_NONE, nhưng vẫn check để an toàn)
            if hasattr(response, 'prompt_feedback') and response.prompt_feedback:
                block_reason = getattr(response.prompt_feedback, 'block_reason', None)
                if block_reason:
                    logger.warning(f"AI Cleanup bị block: {block_reason}. Sử dụng text gốc.")
                    return (text, [0], [text])  # Return text gốc với failed index
            
            if not hasattr(response, 'candidates') or not response.candidates:
                logger.warning("AI Cleanup không có candidates. Sử dụng text gốc.")
                return (text, [0], [text])  # Return text gốc với failed index
            
            logger.info("AI Cleanup: Hoàn tất. Thành công: 1/1 chunk, Thất bại: 0/1 chunk.")
            cleaned_text = response.text.strip()
            return (cleaned_text, [], [text])  # (result_text, failed_indices, original_chunks)
        
        # Build safety settings từ config
        safety_level = ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
        safety_settings = _build_safety_settings(safety_level)
        
        # Text dài, chia nhỏ ở ranh giới câu và xử lý song song
        text_chunks = _split_text_at_sentence_boundaries(text, chunk_size)
        total_chunks = len(text_chunks)
        logger.info(f"AI Cleanup: Chia thành {total_chunks} chunks (ở ranh giới câu), xử lý song song với {len(api_keys)} API keys")
        logger.info(f"AI Cleanup: Safety level: {safety_level}")
        logger.info("AI Cleanup: Bắt đầu xử lý...")
        
        # Chạy async cleanup với safety settings
        result_text, success_count, failure_count, failed_indices = asyncio.run(_ai_cleanup_parallel(text_chunks, api_keys, model_name, prompt, max_parallel, delay, show_progress, timeout_s, max_retries, progress_interval, safety_settings))
        logger.info(f"AI Cleanup: Hoàn tất. Thành công: {success_count}/{total_chunks} chunks, Thất bại: {failure_count}/{total_chunks} chunks (đã lưu nội dung gốc).")
        
        # Tự động retry các chunks failed sau khi hoàn tất tất cả chunks khác
        if failure_count > 0:
            auto_retry = cleanup_cfg.get("auto_retry_failed", True)  # Mặc định: true
            if auto_retry:
                logger.info(f"AI Cleanup: Tự động retry {failure_count} chunks failed...")
                retry_results, still_failed = _retry_failed_chunks_cleanup(
                    failed_indices,
                    text_chunks,
                    api_keys,
                    model_name,
                    prompt,
                    ocr_cfg
                )
                
                # Merge lại text từ retry results
                if retry_results:
                    cleanup_chunks_list = list(text_chunks)
                    for idx, retry_text in retry_results.items():
                        if idx < len(cleanup_chunks_list):
                            cleanup_chunks_list[idx] = retry_text
                    result_text = "\n\n".join(cleanup_chunks_list)
                    
                    retry_success = len(retry_results) - len(still_failed)
                    logger.info(f"AI Cleanup Auto Retry: {retry_success}/{failure_count} chunks retry thành công.")
                    if still_failed:
                        logger.warning(f"AI Cleanup Auto Retry: {len(still_failed)} chunks vẫn failed sau retry.")
                        # Cập nhật failed_indices với still_failed
                        failed_indices = still_failed
                    else:
                        logger.info(f"AI Cleanup Auto Retry: Tất cả chunks failed đã được retry thành công!")
                        failed_indices = []  # Tất cả đã thành công
                else:
                    logger.warning("AI Cleanup Auto Retry: Không có kết quả retry.")
        
        # Trả về text đã merge, failed_indices, và toàn bộ chunks (để có thể rebuild sau retry)
        return (result_text, failed_indices, text_chunks)
        
    except Exception as e:
        logger.error(f"AI cleanup failed: {e}. Trả về text gốc.")
        return (text, [], [])  # Trả về tuple nhất quán


async def _ai_cleanup_parallel(text_chunks: List[str], api_keys: List[str], model_name: str, prompt: str, max_parallel: int, delay: float, show_progress: bool, timeout_s: float, max_retries: int, progress_interval: float, safety_settings: Optional[List[dict]] = None) -> tuple[str, int, int, List[int]]:
    """
    Xử lý song song nhiều chunks với nhiều API keys.
    
    Args:
        safety_settings: Optional safety settings để pass vào GenerativeModel
    """
    # Tạo queue cho API keys
    key_queue = asyncio.Queue()
    for key in api_keys:
        await key_queue.put(key)
    
    cleaned_chunks: List[tuple[int, str]] = []  # (index, cleaned_text)
    semaphore = asyncio.Semaphore(max_parallel)
    total = len(text_chunks)
    failures = 0
    failed_indices: List[int] = []
    
    async def process_chunk(chunk: str, chunk_idx: int) -> tuple[int, str]:
        nonlocal failures, failed_indices  # Khai báo nonlocal ở đầu function
        async with semaphore:
            retries = 0
            api_key = None
            last_error = None
            while retries < max_retries:
                try:
                    api_key = await key_queue.get()
                    cleaned = await _cleanup_chunk_async(chunk, api_key, model_name, prompt, chunk_idx, len(text_chunks), timeout_s, safety_settings)
                    # Thành công - return ngay
                    return (chunk_idx, cleaned)
                except Exception as e:
                    last_error = e
                    retries += 1
                    if retries < max_retries:
                        logger.debug(f"AI cleanup chunk {chunk_idx} failed (attempt {retries}/{max_retries}): {type(e).__name__}: {e}. Retrying...")
                        await asyncio.sleep(delay * retries)
                    else:
                        # Đã retry hết
                        failures += 1
                        failed_indices.append(chunk_idx)
                        logger.warning(f"AI cleanup chunk {chunk_idx} failed after {max_retries} retries with {type(e).__name__}: {e}")
                        return (chunk_idx, chunk)  # Trả về chunk gốc
                finally:
                    if api_key:
                        try:
                            await key_queue.put(api_key)  # Trả key về queue dù thành công hay lỗi
                        except Exception:
                            pass
                    await asyncio.sleep(delay)
            
            # Nếu đến đây (không nên xảy ra)
            failures += 1
            failed_indices.append(chunk_idx)
            logger.warning(f"AI cleanup chunk {chunk_idx} failed after all retries. Last error: {last_error}")
            return (chunk_idx, chunk)
    
    # Tạo tasks cho tất cả chunks
    tasks = [process_chunk(chunk, idx) for idx, chunk in enumerate(text_chunks)]
    
    # Xử lý và log tiến độ định kỳ
    results = []
    if show_progress:
        start_ts = time.time()
        last_log = start_ts
        completed = 0
        async for result in _as_completed_iter(tasks):
            results.append(result)
            completed += 1
            now = time.time()
            if (now - last_log) >= max(5.0, progress_interval):
                elapsed = now - start_ts
                avg = elapsed / completed if completed > 0 else 0.0
                remaining = max(len(tasks) - completed, 0) * avg
                logger.info(f"AI Cleanup: {completed}/{len(tasks)} chunks • TB {avg:.2f}s/chunk • ETA ~{remaining:.0f}s")
                last_log = now
    else:
        results = await asyncio.gather(*tasks)
    cleaned_chunks = sorted(results, key=lambda x: x[0])
    success_count = total - failures
    if failures > 0:
        logger.warning(f"AI Cleanup: {failures}/{total} chunks failed. Tiếp tục với nội dung gốc cho các chunk lỗi.")
    
    # Ghép các chunks theo thứ tự
    result_text = "\n\n".join([text for _, text in cleaned_chunks])
    return (result_text, success_count, failures, failed_indices)


async def _as_completed_iter(coros):
    for fut in asyncio.as_completed(coros):
        yield await fut


def _split_text_at_sentence_boundaries(text: str, max_chunk_size: int) -> List[str]:
    """
    Chia text thành chunks ở ranh giới câu (kết thúc bằng dấu chấm câu).
    Tham khảo thuật toán từ SmartChunker._split_long_paragraph để đảm bảo không cắt giữa câu.
    
    Args:
        text: Văn bản cần chia
        max_chunk_size: Kích thước tối đa của mỗi chunk (tính theo ký tự)
    
    Returns:
        List[str]: Danh sách các chunks đã được chia ở ranh giới câu
    """
    import re
    
    if not text or len(text) <= max_chunk_size:
        return [text] if text else []
    
    # Pattern để tìm ranh giới câu: . ! ? (cả tiếng Anh) và 。！？ (tiếng Trung)
    # Hỗ trợ các dấu ngoặc kép có thể đi kèm: ["']? (cho tiếng Anh) và » (cho một số ngôn ngữ)
    sentence_pattern = re.compile(r'([.!?。！？]["\'»]?\s*)')
    
    # Tìm tất cả các vị trí kết thúc câu
    parts = sentence_pattern.split(text)
    
    # Ghép lại các phần để tạo sentences (mỗi sentence bao gồm nội dung + dấu câu)
    sentences = []
    for i in range(0, len(parts) - 1, 2):
        if i + 1 < len(parts):
            sentence = (parts[i] + parts[i + 1]).strip()
            if sentence:
                sentences.append(sentence)
    
    # Xử lý phần cuối cùng nếu không kết thúc bằng dấu câu
    if len(parts) % 2 == 1 and parts[-1].strip():
        sentences.append(parts[-1].strip())
    
    # Lọc bỏ các câu rỗng
    sentences = [sent for sent in sentences if sent.strip()]
    
    if not sentences:
        return [text]
    
    # Gom các sentences thành chunks, đảm bảo không vượt quá max_chunk_size
    chunks = []
    current_chunk = []
    current_size = 0
    
    for sentence in sentences:
        sent_size = len(sentence)
        
        # Nếu sentence đơn lẻ quá dài, phải cắt (trường hợp hiếm)
        if sent_size > max_chunk_size:
            # Nếu đang có chunk tích lũy, lưu nó trước
            if current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = []
                current_size = 0
            
            # Chia sentence dài thành nhiều phần nhỏ hơn
            # Ưu tiên cắt ở khoảng trắng nếu có thể
            words = sentence.split()
            temp_chunk = []
            temp_size = 0
            
            for word in words:
                word_size = len(word) + 1  # +1 cho space
                if temp_size + word_size > max_chunk_size and temp_chunk:
                    # Lưu chunk hiện tại
                    chunks.append(' '.join(temp_chunk))
                    temp_chunk = [word]
                    temp_size = len(word)
                else:
                    temp_chunk.append(word)
                    temp_size += word_size
            
            if temp_chunk:
                chunks.append(' '.join(temp_chunk))
        else:
            # Kiểm tra nếu thêm sentence này vào chunk hiện tại có vượt quá max_chunk_size không
            # Nếu đã có sentences trong chunk, cần thêm 1 ký tự cho space khi join
            space_needed = 1 if current_chunk else 0
            if current_size + sent_size + space_needed > max_chunk_size and current_chunk:
                # Lưu chunk hiện tại và bắt đầu chunk mới
                chunks.append(' '.join(current_chunk))
                current_chunk = [sentence]
                current_size = sent_size
            else:
                # Thêm sentence vào chunk hiện tại
                current_chunk.append(sentence)
                current_size += sent_size + space_needed
    
    # Lưu chunk cuối cùng nếu có
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    # Nếu không chia được gì (trường hợp hiếm), trả về toàn bộ text
    if not chunks:
        return [text]
    
    return chunks
def _preprocess_line_breaks(text: str) -> str:
    """
    Preprocessing: Nối lại các câu bị ngắt do line breaks khi convert PDF → TXT.
    Chỉ xử lý các trường hợp rõ ràng, các trường hợp phức tạp sẽ để AI xử lý.
    """
    import re
    
    lines = text.split('\n')
    if not lines:
        return text
    
    result_lines = []
    i = 0
    
    while i < len(lines):
        current_line = lines[i].strip()
        
        # Nếu dòng rỗng → giữ nguyên (đây là paragraph break)
        if not current_line:
            result_lines.append('')
            i += 1
            continue
        
        # Bắt đầu từ dòng hiện tại, cố gắng nối các dòng tiếp theo nếu thỏa điều kiện
        merged_line = current_line
        
        # Kiểm tra và nối các dòng tiếp theo liên tục
        while i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            
            # Nếu dòng tiếp theo rỗng
            if not next_line:
                # Kiểm tra xem có phải paragraph break thực sự không
                # Nếu dòng hiện tại kết thúc bằng dấu câu → paragraph break thực sự
                if re.search(r'[.!?]$', merged_line):
                    break
                
                # Nếu không, có thể là line break do PDF format, kiểm tra dòng sau nữa
                if i + 2 < len(lines):
                    next_next_line = lines[i + 2].strip()
                    if next_next_line:
                        # Kiểm tra dòng sau dòng trống
                        first_char = next_next_line[0]
                        next_starts_with_upper = first_char.isupper() and first_char.isalpha()
                        next_starts_with_number = bool(re.match(r'^\d+', next_next_line))
                        next_starts_with_bullet = bool(re.match(r'^[•·\-*]\s', next_next_line))
                        
                        # Nếu dòng sau dòng trống bắt đầu bằng chữ hoa/số/bullet → paragraph break thực sự
                        if next_starts_with_upper or next_starts_with_number or next_starts_with_bullet:
                            break
                        
                        # Nếu dòng sau dòng trống bắt đầu bằng chữ thường → có thể là câu bị ngắt
                        # Bỏ qua dòng trống và tiếp tục với dòng sau
                        next_line = next_next_line
                        i += 1  # Skip dòng trống
                    else:
                        # Không còn dòng nào → dừng
                        break
                else:
                    # Không còn dòng nào → dừng
                    break
            
            # Dòng hiện tại (đã merged) KHÔNG kết thúc bằng dấu kết thúc câu
            ends_with_punctuation = bool(re.search(r'[.!?]$', merged_line))
            
            if not ends_with_punctuation:
                # Kiểm tra nếu dòng tiếp theo bắt đầu bằng chữ hoa
                # Dùng phương pháp đơn giản: kiểm tra ký tự đầu tiên có phải chữ hoa không
                if next_line:
                    first_char = next_line[0]
                    # Kiểm tra nếu là chữ cái và viết hoa (hỗ trợ Unicode)
                    next_starts_with_upper = first_char.isupper() and first_char.isalpha()
                else:
                    next_starts_with_upper = False
                
                next_starts_with_number = bool(re.match(r'^\d+', next_line))
                next_starts_with_bullet = bool(re.match(r'^[•·\-*]\s', next_line))
                
                # Nếu dòng tiếp theo KHÔNG bắt đầu bằng chữ hoa VÀ không phải số/bullet
                # → Có thể là câu bị ngắt, nối lại
                if not next_starts_with_upper and not next_starts_with_number and not next_starts_with_bullet:
                    # Nối với dòng tiếp theo
                    merged_line = merged_line.rstrip() + ' ' + next_line.lstrip()
                    i += 1
                    # Tiếp tục kiểm tra dòng tiếp theo
                    continue
            
            # Không thỏa điều kiện nối → dừng
            break
        
        # Lưu dòng đã merged (hoặc dòng gốc nếu không merge)
        result_lines.append(merged_line)
        i += 1
    
    return '\n'.join(result_lines)


async def _spell_check_chunk_async(chunk: str, api_key: str, model_name: str, prompt: str, chunk_idx: int, total_chunks: int, timeout_s: float, safety_settings: Optional[List[dict]] = None) -> str:
    """
    Soát lỗi chính tả và phục hồi paragraph cho một chunk text bằng AI (async).
    
    Args:
        safety_settings: Optional safety settings để pass vào GenerativeModel (nếu None sẽ dùng default)
    """
    # Suppress logs TRƯỚC khi import
    _suppress_google_logs()
    # Đảm bảo stderr filter đang active
    if not isinstance(sys.stderr, NoisyMessageFilter):
        original_stderr = sys.stderr if not isinstance(sys.stderr, NoisyMessageFilter) else getattr(sys.stderr, 'original_stream', sys.stderr)
        sys.stderr = NoisyMessageFilter(original_stderr)
    import google.generativeai as genai
    genai.configure(api_key=api_key)
    # Pass safety_settings vào GenerativeModel (nếu có)
    model = genai.GenerativeModel(model_name, safety_settings=safety_settings) if safety_settings else genai.GenerativeModel(model_name)
    
    # Run trong thread pool và áp timeout để tránh treo vô hạn
    loop = asyncio.get_event_loop()
    response = await asyncio.wait_for(
        loop.run_in_executor(
            None,
            lambda: model.generate_content(prompt + chunk)
        ),
        timeout=timeout_s
    )
    # Kiểm tra response có hợp lệ không
    if not response or not response.candidates or len(response.candidates) == 0:
        raise ValueError(f"AI spell check chunk {chunk_idx}/{total_chunks}: No candidates returned")
    
    # Kiểm tra prompt_feedback nếu có
    if hasattr(response, 'prompt_feedback') and response.prompt_feedback:
        if hasattr(response.prompt_feedback, 'block_reason') and response.prompt_feedback.block_reason:
            raise ValueError(f"AI spell check chunk {chunk_idx}/{total_chunks}: Blocked by safety filter: {response.prompt_feedback.block_reason}")
    
    result = response.text.strip()
    return result


async def _ai_spell_check_parallel(text_chunks: List[str], api_keys: List[str], model_name: str, prompt: str, max_parallel: int, delay: float, show_progress: bool, timeout_s: float, max_retries: int, progress_interval: float, safety_settings: Optional[List[dict]] = None) -> tuple[str, int, int, List[int]]:
    """
    Xử lý song song nhiều chunks với nhiều API keys cho spell check.
    
    Args:
        safety_settings: Optional safety settings để pass vào GenerativeModel
    """
    # Tạo queue cho API keys
    key_queue = asyncio.Queue()
    for key in api_keys:
        await key_queue.put(key)
    
    processed_chunks: List[tuple[int, str]] = []  # (index, processed_text)
    semaphore = asyncio.Semaphore(max_parallel)
    total = len(text_chunks)
    failures = 0
    failed_indices: List[int] = []
    
    async def process_chunk(chunk: str, chunk_idx: int) -> tuple[int, str]:
        nonlocal failures, failed_indices  # Khai báo nonlocal ở đầu function
        async with semaphore:
            retries = 0
            api_key = None
            last_error = None
            while retries < max_retries:
                try:
                    api_key = await key_queue.get()
                    processed = await _spell_check_chunk_async(chunk, api_key, model_name, prompt, chunk_idx, len(text_chunks), timeout_s, safety_settings)
                    # Thành công - return ngay
                    return (chunk_idx, processed)
                except Exception as e:
                    last_error = e
                    retries += 1
                    if retries < max_retries:
                        logger.debug(f"AI spell check chunk {chunk_idx} failed (attempt {retries}/{max_retries}): {type(e).__name__}: {e}. Retrying...")
                        await asyncio.sleep(delay * retries)
                    else:
                        # Đã retry hết
                        failures += 1
                        failed_indices.append(chunk_idx)
                        logger.warning(f"AI spell check chunk {chunk_idx} failed after {max_retries} retries with {type(e).__name__}: {e}")
                        return (chunk_idx, chunk)  # Trả về chunk gốc
                finally:
                    if api_key:
                        try:
                            await key_queue.put(api_key)  # Trả key về queue dù thành công hay lỗi
                        except Exception:
                            pass
                    await asyncio.sleep(delay)
            
            # Nếu đến đây (không nên xảy ra)
            failures += 1
            failed_indices.append(chunk_idx)
            logger.warning(f"AI spell check chunk {chunk_idx} failed after all retries. Last error: {last_error}")
            return (chunk_idx, chunk)
    
    # Tạo tasks cho tất cả chunks
    tasks = [process_chunk(chunk, idx) for idx, chunk in enumerate(text_chunks)]
    
    # Xử lý và log tiến độ định kỳ
    results = []
    if show_progress:
        start_ts = time.time()
        last_log = start_ts
        completed = 0
        async for result in _as_completed_iter(tasks):
            results.append(result)
            completed += 1
            now = time.time()
            if (now - last_log) >= max(5.0, progress_interval):
                elapsed = now - start_ts
                avg = elapsed / completed if completed > 0 else 0.0
                remaining = max(len(tasks) - completed, 0) * avg
                logger.info(f"AI Spell Check: {completed}/{len(tasks)} chunks • TB {avg:.2f}s/chunk • ETA ~{remaining:.0f}s")
                last_log = now
    else:
        results = await asyncio.gather(*tasks)
    processed_chunks = sorted(results, key=lambda x: x[0])
    success_count = total - failures
    
    if failures > 0:
        logger.warning(f"AI Spell Check: {failures}/{total} chunks failed. Tiếp tục với nội dung gốc cho các chunk lỗi.")
    else:
        logger.info(f"AI Spell Check: Tất cả {total} chunks đã được xử lý thành công.")
    
    # Ghép các chunks theo thứ tự
    result_text = "\n\n".join([text for _, text in processed_chunks])
    return (result_text, success_count, failures, failed_indices)


def _merge_split_table_rows(tables_by_page: dict, ocr_cfg: dict) -> dict:
    """
    Merge các hàng bị cắt từ ô phía trên trong bảng sau spell check.
    
    Logic:
    - Kiểm tra từng hàng trong bảng
    - Nếu hàng có nhiều ô trống ở đầu và chỉ có 1-2 ô có nội dung → có thể là hàng bị cắt
    - Sử dụng AI để đánh giá xem hàng này có phải là phần tiếp theo của hàng trên không
    - Nếu đúng, merge vào ô tương ứng của hàng trên
    
    Args:
        tables_by_page: {page_num: {"rows": [[cell1, cell2, ...], ...], "num_cols": int}}
        ocr_cfg: Config dictionary
    
    Returns:
        dict: Tables đã được merge (cùng format)
    """
    if not tables_by_page:
        return {}
    
    spell_check_cfg = ocr_cfg.get("ai_spell_check", {})
    if not spell_check_cfg.get("enabled", False):
        return tables_by_page
    
    # Lấy API keys
    api_keys = spell_check_cfg.get("api_keys", [])
    if not api_keys:
        api_keys = ocr_cfg.get("_root_api_keys", [])
    if not api_keys:
        logger.warning("Không có API keys để merge hàng bị cắt → bỏ qua")
        return tables_by_page
    
    model_name = spell_check_cfg.get("model", "gemini-2.5-flash")
    safety_level = ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
    
    merged_tables = {}
    
    for page_num, table_info in tables_by_page.items():
        rows = table_info.get("rows", [])
        num_cols = table_info.get("num_cols", 0)
        
        if not rows or num_cols == 0:
            merged_tables[page_num] = table_info
            continue
        
        # Phân tích từng hàng để tìm hàng bị cắt
        merged_rows = []
        i = 0
        while i < len(rows):
            current_row = rows[i]
            
            # Kiểm tra xem hàng này có phải là hàng bị cắt không
            # Dấu hiệu: nhiều ô trống ở đầu, chỉ có 1-2 ô cuối có nội dung
            non_empty_count = sum(1 for cell in current_row if cell and cell.strip())
            empty_prefix_count = sum(1 for cell in current_row if not cell or not cell.strip())
            
            # Nếu có >= 2 ô trống ở đầu và chỉ có 1-2 ô có nội dung → có thể là hàng bị cắt
            if empty_prefix_count >= 2 and non_empty_count <= 2 and i > 0:
                # Kiểm tra với hàng trên bằng AI
                prev_row = merged_rows[-1] if merged_rows else rows[i-1]
                
                # Tạo prompt để AI đánh giá
                prompt = f"""Bạn là AI chuyên phân tích cấu trúc bảng. Nhiệm vụ: Đánh giá xem hàng sau có phải là phần tiếp theo (bị cắt) của hàng trước không.

HÀNG TRƯỚC: {'|'.join(prev_row)}
HÀNG SAU: {'|'.join(current_row)}

QUY TẮC:
- Nếu hàng sau là phần tiếp theo của một ô trong hàng trước (bị cắt xuống dòng) → Trả về "MERGE"
- Nếu hàng sau là hàng mới độc lập → Trả về "KEEP"

Chỉ trả về một từ: "MERGE" hoặc "KEEP", không giải thích thêm."""
                
                try:
                    _suppress_google_logs()
                    if not isinstance(sys.stderr, NoisyMessageFilter):
                        original_stderr = sys.stderr if not isinstance(sys.stderr, NoisyMessageFilter) else getattr(sys.stderr, 'original_stream', sys.stderr)
                        sys.stderr = NoisyMessageFilter(original_stderr)
                    
                    safety_settings = _build_safety_settings(safety_level)
                    
                    import google.generativeai as genai
                    genai.configure(api_key=api_keys[0])
                    model = genai.GenerativeModel(model_name, safety_settings=safety_settings)
                    response = model.generate_content(prompt)
                    
                    decision = response.text.strip().upper()
                    
                    if decision == "MERGE":
                        # Merge hàng sau vào hàng trước
                        # Tìm ô nào trong hàng trước cần merge (thường là ô cuối cùng có nội dung)
                        last_non_empty_col = -1
                        for col_idx in range(len(prev_row) - 1, -1, -1):
                            if prev_row[col_idx] and prev_row[col_idx].strip():
                                last_non_empty_col = col_idx
                                break
                        
                        # Merge nội dung từ hàng sau vào ô tương ứng của hàng trước
                        if last_non_empty_col >= 0:
                            # Tìm ô đầu tiên có nội dung trong hàng sau
                            first_non_empty_col = -1
                            for col_idx in range(len(current_row)):
                                if current_row[col_idx] and current_row[col_idx].strip():
                                    first_non_empty_col = col_idx
                                    break
                            
                            if first_non_empty_col >= 0:
                                # Merge vào ô cuối cùng có nội dung của hàng trước
                                content_to_merge = current_row[first_non_empty_col].strip()
                                if content_to_merge:
                                    prev_row[last_non_empty_col] = (prev_row[last_non_empty_col] + " " + content_to_merge).strip()
                                    logger.debug(f"Đã merge hàng {i+1} vào hàng {i} (cột {last_non_empty_col})")
                                # Bỏ qua hàng hiện tại (đã merge)
                                i += 1
                                continue
                except Exception as e:
                    logger.debug(f"Không thể dùng AI để merge hàng {i+1}: {e}")
                    # Nếu AI fail, giữ nguyên hàng
            
            # Không merge → thêm hàng vào kết quả
            merged_rows.append(current_row)
            i += 1
        
        merged_tables[page_num] = {
            "page": page_num,
            "rows": merged_rows,
            "num_cols": num_cols
        }
        logger.info(f"Đã xử lý merge hàng cho bảng trang {page_num}: {len(rows)} → {len(merged_rows)} hàng")
    
    return merged_tables


def ai_spell_check_and_paragraph_restore(text: str, ocr_cfg: dict) -> str:
    """
    Sử dụng AI để soát lỗi chính tả và phục hồi cấu trúc paragraph.
    Đặc biệt chú ý bảo vệ toàn vẹn nội dung (không thay đổi ý nghĩa).
    Hỗ trợ nhiều API keys để xử lý song song.
    """
    spell_check_cfg = ocr_cfg.get("ai_spell_check", {})
    spell_check_enabled = spell_check_cfg.get("enabled", False)
    if not spell_check_enabled:
        return text
    
    # Ghi chú: Không dùng preprocessing rule-based vì AI sẽ phân tích ngữ cảnh tốt hơn
    # Hàm _preprocess_line_breaks vẫn được giữ lại nếu cần dùng trong tương lai
    # text = _preprocess_line_breaks(text)  # Tạm tắt để AI làm toàn bộ
    
    # Lấy API keys (ưu tiên từ ai_spell_check.api_keys, fallback về api_keys từ root config)
    api_keys = spell_check_cfg.get("api_keys", [])
    if not api_keys:
        # Đọc từ _root_api_keys đã lưu khi load config
        api_keys = ocr_cfg.get("_root_api_keys", [])
    
    if not api_keys:
        logger.warning("AI spell check enabled nhưng không có API keys. Bỏ qua spell check.")
        return text
    
    model_name = spell_check_cfg.get("model", "gemini-2.5-flash")
    max_parallel = spell_check_cfg.get("max_parallel_workers", 5)
    # Giới hạn worker theo số API keys sẵn có
    if api_keys:
        max_parallel = max(1, min(max_parallel, len(api_keys)))
    chunk_size = spell_check_cfg.get("chunk_size", 50000)
    delay = spell_check_cfg.get("delay_between_requests", 0.5)
    max_retries = spell_check_cfg.get("max_retries", 3)
    timeout_s = float(spell_check_cfg.get("ai_timeout_seconds", 120))
    show_progress = bool(ocr_cfg.get("show_progress", True))
    progress_interval = float(ocr_cfg.get("progress_log_interval_seconds", 60))
    
    prompt = """Bạn là một AI chuyên soát lỗi chính tả và phục hồi cấu trúc văn bản OCR. Nhiệm vụ chính của bạn là PHÂN TÍCH NGỮ CẢNH và QUYẾT ĐỊNH THÔNG MINH.

=== NHIỆM VỤ CHÍNH: PHÂN TÍCH VÀ PHỤC HỒI CÂU BỊ NGẮT (Ưu tiên cao nhất) ===

Bạn cần ĐỌC KỸ NỘI DUNG và PHÂN TÍCH để phân biệt:

A. CÂU BỊ NGẮT DO CONVERT PDF → TXT (CẦN NỐI LẠI):
   - Đọc ngữ cảnh: Nếu dòng trước chưa hoàn thành ý và dòng sau tiếp nối ý đó → nối lại
   - Ví dụ: 
     * "Our client is also the owner of Vietnam Trade Mark Registration No. 315843 for "MICROBAN"
       in Class 5 covering..." 
     → Phân tích: "in Class 5" tiếp nối câu trước → NỐI LẠI thành một câu
   
   - Dấu hiệu cần nối:
     * Dòng trước không kết thúc bằng dấu câu (. ! ?) HOẶC kết thúc bằng dấu phẩy, hai chấm
     * Dòng sau bắt đầu bằng chữ thường (tiếp nối câu trước)
     * Nội dung dòng sau về mặt ngữ pháp và ngữ nghĩa là phần tiếp theo của câu trước
     * Đọc toàn bộ ngữ cảnh để hiểu rõ mối quan hệ

B. NGẮT PARAGRAPH CÓ CHỦ ĐÍCH (KHÔNG NỐI):
   - Đọc ngữ cảnh: Nếu dòng sau là ý mới, chủ đề mới, hoặc đoạn văn mới → KHÔNG nối
   - Ví dụ:
     * "...attached as Exhibit 1.
       
       Khách hàng của chúng tôi là chủ sở hữu..."
     → Phân tích: Đây là đoạn mới (chuyển từ tiếng Anh sang tiếng Việt) → KHÔNG NỐI
   
   - Dấu hiệu KHÔNG nối:
     * Dòng trước kết thúc bằng dấu chấm (. ! ?) và dòng sau bắt đầu bằng chữ hoa
     * Dòng sau là câu đầu tiên của một đoạn mới (ý tưởng mới, chủ đề mới)
     * Có sự thay đổi rõ ràng về ngữ cảnh (ví dụ: chuyển từ phần này sang phần khác)
     * Đọc toàn bộ ngữ cảnh để xác định đây là ngắt đoạn có chủ đích

QUY TRÌNH PHÂN TÍCH:
1. ĐỌC toàn bộ văn bản để hiểu cấu trúc và ngữ cảnh
2. PHÂN TÍCH từng vị trí ngắt dòng:
   - Xem xét nội dung trước và sau dòng ngắt
   - Đánh giá mối quan hệ ngữ pháp và ngữ nghĩa
   - Xác định đây là câu bị ngắt hay ngắt đoạn có chủ đích
3. QUYẾT ĐỊNH:
   - Nếu là câu bị ngắt → NỐI lại (thay line break bằng space)
   - Nếu là ngắt đoạn có chủ đích → GIỮ NGUYÊN (có thể thêm dòng trống nếu cần)
4. ÁP DỤNG nhất quán cho toàn bộ văn bản

=== CÁC NHIỆM VỤ KHÁC ===

1. SOÁT LỖI CHÍNH TẢ:
   - Sửa các lỗi chính tả do OCR (ví dụ: "Kíng" → "Kính", "hang" → "hàng")
   - Sửa các lỗi chính tả thông thường
   - KHÔNG thay đổi từ ngữ chuyên ngành, tên riêng, địa danh
   - KHÔNG thay đổi số liệu, ngày tháng, địa chỉ

2. PHỤC HỒI CẤU TRÚC PARAGRAPH:
   - Sau khi đã nối các câu bị ngắt, xác định các ngắt đoạn hợp lý
   - Mỗi đoạn văn nên có một ý chính hoàn chỉnh
   - Giữ nguyên các dòng trống giữa các đoạn đã được xác định là có chủ đích
   - Đảm bảo các câu trong một đoạn có liên quan với nhau

3. BẢO VỆ TOÀN VẸN NỘI DUNG:
   - TUYỆT ĐỐI KHÔNG thay đổi ý nghĩa của văn bản
   - KHÔNG thêm, bớt, hoặc diễn giải lại nội dung
   - KHÔNG thay đổi thứ tự từ trong câu (chỉ nối lại khi cần)
   - GIỮ NGUYÊN định dạng đặc biệt (bullet points, numbered lists, bảng)
   - GIỮ NGUYÊN các từ viết hoa nếu chúng là tên riêng, thuật ngữ

4. ĐỊNH DẠNG:
   - Giữ nguyên định dạng văn bản song ngữ (nếu có)
   - Giữ nguyên các dấu câu quan trọng
   - Chuẩn hóa khoảng trắng thừa giữa các từ (nhưng không thay đổi paragraph breaks hợp lý)
   - Đảm bảo mỗi câu kết thúc bằng dấu câu thích hợp

=== NGUYÊN TẮC QUAN TRỌNG ===

- SỬ DỤNG SỨC MẠNH PHÂN TÍCH NGỮ CẢNH: Đọc và hiểu nội dung, không chỉ dựa vào quy tắc cú pháp
- QUYẾT ĐỊNH THÔNG MINH: Mỗi quyết định nối hay không nối phải dựa trên phân tích ngữ cảnh cụ thể
- NHẤT QUÁN: Áp dụng cùng một tiêu chuẩn phân tích cho toàn bộ văn bản
- BẢO TOÀN Ý NGHĨA: Chỉ điều chỉnh cấu trúc, KHÔNG thay đổi nội dung hoặc ý nghĩa

Trả về chỉ văn bản đã được soát và phục hồi, không giải thích thêm.

Văn bản cần phân tích và xử lý:
"""
    
    try:
        # Chia nhỏ text nếu quá dài
        if len(text) <= chunk_size:
            # Text ngắn, xử lý trực tiếp
            logger.info("AI Spell Check: Text ngắn, xử lý trực tiếp (1 chunk)")
            # Suppress logs TRƯỚC khi import
            _suppress_google_logs()
            # Đảm bảo stderr filter đang active
            if not isinstance(sys.stderr, NoisyMessageFilter):
                original_stderr = sys.stderr if not isinstance(sys.stderr, NoisyMessageFilter) else getattr(sys.stderr, 'original_stream', sys.stderr)
                sys.stderr = NoisyMessageFilter(original_stderr)
            # Build safety settings từ config
            safety_level = ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
            safety_settings = _build_safety_settings(safety_level)
            
            import google.generativeai as genai
            genai.configure(api_key=api_keys[0])
            model = genai.GenerativeModel(model_name, safety_settings=safety_settings)
            response = model.generate_content(prompt + text)
            
            # Kiểm tra nếu response bị block (mặc dù đã set BLOCK_NONE, nhưng vẫn check để an toàn)
            if hasattr(response, 'prompt_feedback') and response.prompt_feedback:
                block_reason = getattr(response.prompt_feedback, 'block_reason', None)
                if block_reason:
                    logger.warning(f"AI Spell Check bị block: {block_reason}. Sử dụng text gốc.")
                    return (text, [0], [text])  # Return text gốc với failed index
            
            if not hasattr(response, 'candidates') or not response.candidates:
                logger.warning("AI Spell Check không có candidates. Sử dụng text gốc.")
                return (text, [0], [text])  # Return text gốc với failed index
            
            logger.info("AI Spell Check: Hoàn tất. Thành công: 1/1 chunk, Thất bại: 0/1 chunk.")
            checked_text = response.text.strip()
            return (checked_text, [], [text])  # (result_text, failed_indices, original_chunks)
        
        # Build safety settings từ config
        safety_level = ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
        safety_settings = _build_safety_settings(safety_level)
        
        # Text dài, chia nhỏ ở ranh giới câu và xử lý song song
        text_chunks = _split_text_at_sentence_boundaries(text, chunk_size)
        total_chunks = len(text_chunks)
        logger.info(f"AI Spell Check: Chia thành {total_chunks} chunks (ở ranh giới câu), xử lý song song với {len(api_keys)} API keys")
        logger.info(f"AI Spell Check: Safety level: {safety_level}")
        logger.info("AI Spell Check: Bắt đầu xử lý...")
        
        # Chạy async spell check với safety settings
        result_text, success_count, failure_count, failed_indices = asyncio.run(_ai_spell_check_parallel(text_chunks, api_keys, model_name, prompt, max_parallel, delay, show_progress, timeout_s, max_retries, progress_interval, safety_settings))
        logger.info(f"AI Spell Check: Hoàn tất. Thành công: {success_count}/{total_chunks} chunks, Thất bại: {failure_count}/{total_chunks} chunks (đã lưu nội dung gốc).")
        
        # Tự động retry các chunks failed sau khi hoàn tất tất cả chunks khác
        if failure_count > 0:
            auto_retry = spell_check_cfg.get("auto_retry_failed", True)  # Mặc định: true
            if auto_retry:
                logger.info(f"AI Spell Check: Tự động retry {failure_count} chunks failed...")
                retry_results, still_failed = _retry_failed_chunks_spell_check(
                    failed_indices,
                    text_chunks,
                    api_keys,
                    model_name,
                    prompt,
                    ocr_cfg
                )
                
                # Merge lại text từ retry results
                if retry_results:
                    spell_check_chunks_list = list(text_chunks)
                    for idx, retry_text in retry_results.items():
                        if idx < len(spell_check_chunks_list):
                            spell_check_chunks_list[idx] = retry_text
                    result_text = "\n\n".join(spell_check_chunks_list)
                    
                    retry_success = len(retry_results) - len(still_failed)
                    logger.info(f"AI Spell Check Auto Retry: {retry_success}/{failure_count} chunks retry thành công.")
                    if still_failed:
                        logger.warning(f"AI Spell Check Auto Retry: {len(still_failed)} chunks vẫn failed sau retry.")
                        # Cập nhật failed_indices với still_failed
                        failed_indices = still_failed
                    else:
                        logger.info(f"AI Spell Check Auto Retry: Tất cả chunks failed đã được retry thành công!")
                        failed_indices = []  # Tất cả đã thành công
                else:
                    logger.warning("AI Spell Check Auto Retry: Không có kết quả retry.")
        
        # Trả về text đã merge, failed_indices, và toàn bộ chunks (để có thể rebuild sau retry)
        return (result_text, failed_indices, text_chunks)
        
    except Exception as e:
        logger.error(f"AI spell check failed: {e}. Trả về text gốc.")
        return (text, [], [])  # Trả về tuple nhất quán


def _retry_failed_chunks_cleanup(failed_indices: List[int], all_chunks: List[str], api_keys: List[str], model_name: str, prompt: str, ocr_cfg: dict) -> tuple[dict[int, str], List[int]]:
    """Retry các chunk failed cho AI Cleanup. Trả về dict {idx: processed_text} và danh sách still_failed."""
    if not failed_indices or not all_chunks:
        return ({}, [])
    
    cleanup_cfg = ocr_cfg.get("ai_cleanup", {})
    timeout_s = float(cleanup_cfg.get("ai_timeout_seconds", 240))
    
    failed_chunks = [(idx, all_chunks[idx]) for idx in failed_indices if idx < len(all_chunks)]
    logger.info(f"AI Cleanup Retry: Đang retry {len(failed_chunks)} chunks failed...")
    
    # Build safety settings từ config
    safety_level = ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
    safety_settings = _build_safety_settings(safety_level)
    
    async def _retry_chunk(idx: int, chunk: str) -> tuple[int, str]:
        for key in api_keys:
            try:
                result = await _cleanup_chunk_async(chunk, key, model_name, prompt, idx, len(failed_chunks), timeout_s, safety_settings)
                return (idx, result)
            except Exception:
                continue
        return (idx, chunk)  # Fallback về chunk gốc
    
    tasks = [_retry_chunk(idx, chunk) for idx, chunk in failed_chunks]
    results = asyncio.run(asyncio.gather(*tasks))
    
    retry_results = {idx: text for idx, text in results}
    still_failed = [idx for idx in failed_indices if retry_results.get(idx) == all_chunks[idx]]
    
    logger.info(f"AI Cleanup Retry: {len(failed_indices) - len(still_failed)}/{len(failed_indices)} chunks retry thành công.")
    if still_failed:
        logger.warning(f"AI Cleanup Retry: {len(still_failed)} chunks vẫn failed sau retry.")
    
    return (retry_results, still_failed)
def _retry_failed_chunks_spell_check(failed_indices: List[int], all_chunks: List[str], api_keys: List[str], model_name: str, prompt: str, ocr_cfg: dict) -> tuple[dict[int, str], List[int]]:
    """Retry các chunk failed cho AI Spell Check. Trả về dict {idx: processed_text} và danh sách still_failed."""
    if not failed_indices or not all_chunks:
        return ({}, [])
    
    spell_check_cfg = ocr_cfg.get("ai_spell_check", {})
    timeout_s = float(spell_check_cfg.get("ai_timeout_seconds", 240))
    
    failed_chunks = [(idx, all_chunks[idx]) for idx in failed_indices if idx < len(all_chunks)]
    logger.info(f"AI Spell Check Retry: Đang retry {len(failed_chunks)} chunks failed...")
    
    # Build safety settings từ config
    safety_level = ocr_cfg.get("safety_level", "BLOCK_ONLY_HIGH")
    safety_settings = _build_safety_settings(safety_level)
    
    async def _retry_chunk(idx: int, chunk: str) -> tuple[int, str]:
        for key in api_keys:
            try:
                result = await _spell_check_chunk_async(chunk, key, model_name, prompt, idx, len(failed_chunks), timeout_s, safety_settings)
                return (idx, result)
            except Exception:
                continue
        return (idx, chunk)  # Fallback về chunk gốc
    
    tasks = [_retry_chunk(idx, chunk) for idx, chunk in failed_chunks]
    results = asyncio.run(asyncio.gather(*tasks))
    
    retry_results = {idx: text for idx, text in results}
    still_failed = [idx for idx in failed_indices if retry_results.get(idx) == all_chunks[idx]]
    
    logger.info(f"AI Spell Check Retry: {len(failed_indices) - len(still_failed)}/{len(failed_indices)} chunks retry thành công.")
    if still_failed:
        logger.warning(f"AI Spell Check Retry: {len(still_failed)} chunks vẫn failed sau retry.")
    
    return (retry_results, still_failed)


def _get_intermediate_file_path(output_path: str, suffix: str) -> str:
    """Tạo đường dẫn file tạm thời dựa trên output_path và suffix."""
    output_dir = os.path.dirname(output_path) if os.path.dirname(output_path) else "."
    output_basename = os.path.basename(output_path)
    output_name_without_ext = os.path.splitext(output_basename)[0]
    return os.path.join(output_dir, output_name_without_ext + suffix)


def _cleanup_intermediate_files(output_path: str):
    """
    Xóa các file trung gian (_ocred.txt, _cleanup.txt) sau khi đã tạo file final.
    
    Args:
        output_path: Đường dẫn file output final (để tạo tên file trung gian)
    """
    intermediate_files = [
        _get_intermediate_file_path(output_path, "_ocred.txt"),
        _get_intermediate_file_path(output_path, "_cleanup.txt")
    ]
    
    for file_path in intermediate_files:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.debug(f"🗑️  Đã xóa file trung gian: {file_path}")
        except Exception as e:
            logger.debug(f"Không thể xóa file trung gian {file_path}: {e}")


def _check_existing_files(output_path: str) -> dict:
    """Kiểm tra các file đã tồn tại từ phiên làm việc trước."""
    results = {
        "ocred": None,
        "cleanup": None,
        "output": None,
        "all_exist": False
    }
    
    ocred_path = _get_intermediate_file_path(output_path, "_ocred.txt")
    cleanup_path = _get_intermediate_file_path(output_path, "_cleanup.txt")
    
    if os.path.exists(ocred_path):
        results["ocred"] = ocred_path
    if os.path.exists(cleanup_path):
        results["cleanup"] = cleanup_path
    if os.path.exists(output_path):
        results["output"] = output_path
    
    results["all_exist"] = any([results["ocred"], results["cleanup"], results["output"]])
    return results


def _load_resume_file(file_path: str, step_name: str) -> Optional[str]:
    """Load file từ phiên trước để resume."""
    if not os.path.exists(file_path):
        return None
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        logger.info(f"✅ Đã load file {step_name}: {file_path}")
        return content
    except Exception as e:
        logger.warning(f"Không thể load file {step_name} ({file_path}): {e}")
        return None


def _show_completion_menu(cleanup_failed: int, spell_check_failed: int, output_path: str = None) -> str:
    """Hiển thị menu lựa chọn sau khi OCR hoàn tất. Trả về 'retry', 'save', hoặc 'exit'."""
    import threading
    
    has_failures = cleanup_failed > 0 or spell_check_failed > 0
    user_choice = None
    user_choice_lock = threading.Lock()
    user_choice_done = threading.Event()
    
    def _auto_save_timer():
        nonlocal user_choice
        time.sleep(600)  # 10 phút = 600 giây
        with user_choice_lock:
            if user_choice is None:
                logger.info("\n⏰ Tự động lưu file sau 10 phút...")
                user_choice = "save"
                user_choice_done.set()
    
    auto_save_thread = threading.Thread(target=_auto_save_timer, daemon=True)
    auto_save_thread.start()
    
    if not has_failures:
        # Không có lỗi, chỉ có option save/exit
        logger.info("\n" + "=" * 80)
        logger.info("✅ OCR hoàn tất không có lỗi!")
        logger.info("=" * 80)
        logger.info("Lựa chọn:")
        logger.info("  1. Lưu file (tự động lưu sau 10 phút nếu không chọn)")
        logger.info("  2. Thoát không lưu")
        logger.info("=" * 80)
        
        while not user_choice_done.is_set():
            try:
                choice = input("\nNhập lựa chọn (1/2): ").strip()
                with user_choice_lock:
                    if choice == "1":
                        user_choice = "save"
                        user_choice_done.set()
                        break
                    elif choice == "2":
                        user_choice = "exit"
                        user_choice_done.set()
                        break
                    else:
                        logger.warning("Lựa chọn không hợp lệ. Vui lòng nhập 1 hoặc 2.")
            except (EOFError, KeyboardInterrupt):
                with user_choice_lock:
                    user_choice = "save"
                    user_choice_done.set()
                break
    else:
        # Có lỗi, hiển thị đầy đủ 3 options
        logger.info("\n" + "=" * 80)
        logger.info("⚠️  OCR hoàn tất với một số lỗi:")
        if cleanup_failed > 0:
            logger.info(f"  - AI Cleanup: {cleanup_failed} chunks failed")
        if spell_check_failed > 0:
            logger.info(f"  - AI Spell Check: {spell_check_failed} chunks failed")
        logger.info("=" * 80)
        logger.info("Lựa chọn:")
        logger.info("  1. Retry các chunk failed")
        logger.info("  2. Lưu file (tự động lưu sau 10 phút nếu không chọn)")
        logger.info("  3. Thoát không lưu")
        logger.info("=" * 80)
        
        while not user_choice_done.is_set():
            try:
                choice = input("\nNhập lựa chọn (1/2/3): ").strip()
                with user_choice_lock:
                    if choice == "1":
                        user_choice = "retry"
                        user_choice_done.set()
                        break
                    elif choice == "2":
                        user_choice = "save"
                        user_choice_done.set()
                        break
                    elif choice == "3":
                        user_choice = "exit"
                        user_choice_done.set()
                        break
                    else:
                        logger.warning("Lựa chọn không hợp lệ. Vui lòng nhập 1, 2 hoặc 3.")
            except (EOFError, KeyboardInterrupt):
                with user_choice_lock:
                    user_choice = "save"
                    user_choice_done.set()
                break
    
    # Đợi user chọn hoặc auto-save
    user_choice_done.wait()
    return user_choice if user_choice else "save"


def ocr_image(image_path: str, config_path: str = "config/config.yaml") -> str:
    ocr_cfg = _detect_bundled_binaries(load_ocr_config(config_path))
    _ensure_dependencies(ocr_cfg)
    if Image is None:
        raise RuntimeError("Pillow not installed. Please install pillow.")
    _apply_tesseract_cfg(ocr_cfg)
    if not os.path.exists(image_path):
        raise FileNotFoundError(image_path)
    logger.info(f"OCR: Đang nhận dạng ảnh: {image_path}")
    img = Image.open(image_path)
    # Auto-detect language/variant nếu cần
    raw_lang = ocr_cfg.get("lang", "vie")
    normalized_lang = _normalize_lang_code(raw_lang)
    
    # Chỉ detect Chinese variant nếu lang="CN" hoặc "chi" (không có auto-detect)
    needs_chinese_variant_detection = ("chi" in normalized_lang.lower() and 
                                       "chi_sim" not in normalized_lang and 
                                       "chi_tra" not in normalized_lang)
    
    if needs_chinese_variant_detection:
        # Chỉ detect Chinese variant (giản thể/phồn thể)
        resolved_lang = _resolve_language(raw_lang, ocr_cfg, sample_img=img)
        text = _image_to_text(img, ocr_cfg, lang_override=resolved_lang)
    else:
        # Chỉ normalize, không cần detect
        resolved_lang = _resolve_language(raw_lang, ocr_cfg, sample_img=None)
        text = _image_to_text(img, ocr_cfg, lang_override=resolved_lang)
    return text


def ocr_pdf(pdf_path: str, config_path: str = "config/config.yaml", pages: Optional[List[int]] = None) -> tuple[str, int]:
    ocr_cfg = _detect_bundled_binaries(load_ocr_config(config_path))
    _ensure_dependencies(ocr_cfg)
    if convert_from_path is None:
        raise RuntimeError("pdf2image not installed. Please install pdf2image and poppler if needed.")
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(pdf_path)
    _apply_tesseract_cfg(ocr_cfg)

    # Tối ưu DPI: giảm mặc định từ 300 → 250
    dpi = int(ocr_cfg.get("dpi", 250) or 250)
    poppler_path = ocr_cfg.get("poppler_path")
    
    # Config cho batch processing và memory optimization
    max_batch_size = int(ocr_cfg.get("render_batch_size", 20))  # Render tối đa 20 trang/batch
    image_format = ocr_cfg.get("image_format", "jpeg").lower()  # jpeg hoặc png
    jpeg_quality = int(ocr_cfg.get("jpeg_quality", 85))  # Quality 85-90 cho OCR
    memory_optimize = ocr_cfg.get("memory_optimize", True)

    # Resume/caching: sử dụng thư mục cùng tên file input để lưu/trích xuất ảnh các trang
    pdf_p = Path(pdf_path)
    cache_dir = pdf_p.with_suffix("")  # cùng tên với file, bỏ đuôi .pdf
    
    # Helper function để render và save với batch processing, format tối ưu, và memory management
    def _render_and_save_batch(first_page: int, last_page: int, image_format: str, jpeg_quality: int, memory_optimize: bool) -> dict[int, Path]:
        """Render một batch pages và save với format tối ưu. Trả về dict: page_idx → Path."""
        result: dict[int, Path] = {}
        try:
            # CẢI TIẾN: Error handling tốt hơn cho Poppler
            # Dùng biến local để tránh UnboundLocalError khi gán lại poppler_path
            actual_poppler_path = poppler_path
            try:
                if actual_poppler_path and isinstance(actual_poppler_path, str) and actual_poppler_path.strip():
                    # Kiểm tra Poppler path có tồn tại không
                    poppler_bin = os.path.join(actual_poppler_path, "pdftoppm.exe" if sys.platform == "win32" else "pdftoppm")
                    if not os.path.exists(poppler_bin):
                        logger.warning(f"⚠️  Poppler path không hợp lệ: {actual_poppler_path}")
                        logger.warning(f"💡 Đang thử không dùng poppler_path...")
                        actual_poppler_path = None
                    
                    if actual_poppler_path:
                        imgs = convert_from_path(pdf_path, dpi=dpi, poppler_path=actual_poppler_path,
                                                  first_page=first_page, last_page=last_page, thread_count=1)
                    else:
                        imgs = convert_from_path(pdf_path, dpi=dpi, first_page=first_page, last_page=last_page, thread_count=1)
                else:
                    imgs = convert_from_path(pdf_path, dpi=dpi, first_page=first_page, last_page=last_page, thread_count=1)
            except Exception as poppler_err:
                error_msg = str(poppler_err).lower()
                if "poppler" in error_msg or "pdftoppm" in error_msg or "pdftocairo" in error_msg:
                    logger.error(f"❌ Lỗi Poppler khi render batch {first_page}-{last_page}: {poppler_err}")
                    logger.error("💡 Hướng dẫn cài đặt Poppler:")
                    logger.error("   - Windows: Tải từ https://github.com/oschwartz10612/poppler-windows/releases")
                    logger.error("   - Hoặc dùng: choco install poppler")
                    logger.error("   - Sau khi cài, thêm đường dẫn bin vào config.yaml (poppler_path)")
                    raise RuntimeError(f"Poppler không khả dụng: {poppler_err}. Vui lòng cài đặt Poppler và cấu hình poppler_path trong config.yaml")
                else:
                    # Lỗi khác (có thể là PDF corrupt, permission, etc.)
                    logger.error(f"❌ Lỗi khi convert PDF sang ảnh (batch {first_page}-{last_page}): {poppler_err}")
                    raise
            
            # CẢI TIẾN: Memory handling tốt hơn trong loop
            for offset, img in enumerate(imgs):
                idx = first_page + offset
                # Chọn extension dựa trên format
                ext = ".jpg" if image_format == "jpeg" else ".png"
                out_path = cache_dir / f"page_{idx:04d}{ext}"
                
                try:
                    # Save với format và compression tối ưu
                    if image_format == "jpeg":
                        # Convert RGBA/RGB nếu cần (JPEG không hỗ trợ alpha)
                        if img.mode in ('RGBA', 'LA', 'P'):
                            # Tạo nền trắng cho alpha channel
                            rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                            if img.mode == 'P':
                                img = img.convert('RGBA')
                            rgb_img.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                            # CẢI TIẾN: Giải phóng img cũ trước khi gán mới
                            if memory_optimize:
                                del img
                                gc.collect()
                            img = rgb_img
                        img.save(str(out_path), format='JPEG', quality=jpeg_quality, optimize=True)
                    else:
                        # PNG với optimize
                        img.save(str(out_path), format='PNG', optimize=True)
                    
                    result[idx] = out_path
                    
                    # CẢI TIẾN: Memory management - giải phóng ngay sau khi save
                    if memory_optimize:
                        if 'img' in locals():
                            del img
                        if offset % 5 == 0:  # Garbage collect mỗi 5 images
                            gc.collect()
                except Exception as e:
                    logger.warning(f"Không thể lưu ảnh cache {out_path}: {e}")
                    # Giải phóng memory ngay cả khi lỗi
                    if memory_optimize and 'img' in locals():
                        del img
                        gc.collect()
            
            # CẢI TIẾN: Final garbage collect sau batch
            if memory_optimize:
                if 'imgs' in locals():
                    del imgs
                gc.collect()
        except RuntimeError:
            # Re-raise RuntimeError từ Poppler (đã có hướng dẫn)
            raise
        except Exception as e:
            logger.error(f"Render batch {first_page}-{last_page} thất bại: {e}")
            # Giải phóng memory nếu có
            if memory_optimize:
                if 'imgs' in locals():
                    del imgs
                gc.collect()
        return result
    
    def _split_range_into_batches(range_start: int, range_end: int, batch_size: int) -> List[tuple[int, int]]:
        """Chia một range lớn thành các batches nhỏ."""
        batches = []
        current = range_start
        while current <= range_end:
            batch_end = min(current + batch_size - 1, range_end)
            batches.append((current, batch_end))
            current = batch_end + 1
        return batches

    def _list_cached_images(dir_path: Path) -> List[Path]:
        if not dir_path.exists() or not dir_path.is_dir():
            return []
        image_exts = {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}
        files = [p for p in dir_path.iterdir() if p.suffix.lower() in image_exts]
        if not files:
            return []
        def sort_key(p: Path):
            name = p.stem
            digits = "".join(ch for ch in name if ch.isdigit())
            return (int(digits) if digits else 0, name)
        return sorted(files, key=sort_key)

    cached_images = _list_cached_images(cache_dir)

    # Lấy tổng số trang PDF để so sánh cache và thực hiện resume nếu thiếu
    def _get_total_pages(pdf_file: str) -> Optional[int]:
        try:
            if pdfplumber is not None:
                with pdfplumber.open(pdf_file) as pdf:
                    return len(pdf.pages)
        except Exception:
            pass
        try:
            if PyPDF2 is not None:
                with open(pdf_file, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    return len(reader.pages)
        except Exception:
            pass
        return None

    total_pages = _get_total_pages(pdf_path)
    if total_pages is not None:
        logger.info(f"OCR: PDF có {total_pages} trang. Ảnh cache hiện có: {len(cached_images)}")
    else:
        logger.info(f"OCR: Không xác định được tổng số trang. Ảnh cache hiện có: {len(cached_images)}")
    
    # Filter pages nếu có chỉ định
    if pages and total_pages is not None:
        valid_pages = [p for p in pages if 1 <= p <= total_pages]
        invalid_pages = [p for p in pages if p < 1 or p > total_pages]
        if invalid_pages:
            logger.warning(f"Các trang không hợp lệ (nằm ngoài 1-{total_pages}): {invalid_pages}. Bỏ qua.")
        if not valid_pages:
            logger.error("Không có trang hợp lệ nào để OCR.")
            return ("", 0)
        pages_to_ocr = sorted(set(valid_pages))
        logger.info(f"OCR: Chỉ OCR {len(pages_to_ocr)} trang: {pages_to_ocr}")
    elif pages:
        # Không biết total_pages nhưng có pages chỉ định → dùng pages đó
        pages_to_ocr = sorted(set([p for p in pages if p > 0]))
        logger.info(f"OCR: Chỉ OCR {len(pages_to_ocr)} trang (theo chỉ định): {pages_to_ocr}")
    else:
        pages_to_ocr = None  # Tất cả trang

    # Map chỉ số trang → đường dẫn ảnh (cache) hoặc ảnh render mới
    index_to_image_path: dict[int, Path] = {}
    # Parse chỉ số từ tên ảnh cache kiểu page_0001.png
    for p in cached_images:
        name = p.stem
        digits = "".join(ch for ch in name if ch.isdigit())
        if digits:
            try:
                idx = int(digits)
                # Chỉ lấy cache nếu trang đó nằm trong pages_to_ocr (hoặc pages_to_ocr = None)
                if pages_to_ocr is None or idx in pages_to_ocr:
                    index_to_image_path[idx] = p
            except Exception:
                continue

    # Render bổ sung cho các trang thiếu nếu biết total_pages
    if total_pages is not None:
        # Tính missing pages: nếu có pages_to_ocr thì chỉ tính trong đó, ngược lại tính tất cả
        if pages_to_ocr is not None:
            target_pages = set(pages_to_ocr)
            missing = [i for i in target_pages if i not in index_to_image_path]
        else:
            missing = [i for i in range(1, total_pages + 1) if i not in index_to_image_path]
        if missing:
            logger.info(f"OCR: Phát hiện thiếu {len(missing)}/{total_pages} ảnh → render phần còn thiếu")
            # Tạo thư mục cache nếu cần
            try:
                cache_dir.mkdir(parents=True, exist_ok=True)
            except Exception:
                pass
            # Gom missing pages thành các khoảng liên tiếp để render theo range
            ranges: List[tuple[int, int]] = []
            start = prev = None
            for m in missing:
                if start is None:
                    start = prev = m
                elif m == prev + 1:
                    prev = m
                else:
                    ranges.append((start, prev))
                    start = prev = m
            if start is not None:
                ranges.append((start, prev))

            # Render với batch processing để giảm memory usage
            for first, last in ranges:
                # Chia range lớn thành batches nhỏ
                batches = _split_range_into_batches(first, last, max_batch_size)
                for batch_first, batch_last in batches:
                    logger.info(f"OCR: Render bổ sung trang {batch_first}–{batch_last}/{last} (dpi={dpi}, format={image_format})")
                    batch_results = _render_and_save_batch(batch_first, batch_last, image_format, jpeg_quality, memory_optimize)
                    index_to_image_path.update(batch_results)
    # Nếu vẫn chưa có ảnh nào (không có cache và không biết total), render
    if not index_to_image_path:
        if pages_to_ocr is not None and total_pages is not None:
            # Chỉ render các trang được chỉ định
            logger.info(f"OCR: Chuyển PDF → ảnh (dpi={dpi}) cho {len(pages_to_ocr)} trang: {pages_to_ocr}")
            # Gom pages_to_ocr thành ranges để render hiệu quả
            ranges: List[tuple[int, int]] = []
            pages_sorted = sorted(pages_to_ocr)
            start = prev = None
            for p in pages_sorted:
                if start is None:
                    start = prev = p
                elif p == prev + 1:
                    prev = p
                else:
                    ranges.append((start, prev))
                    start = prev = p
            if start is not None:
                ranges.append((start, prev))
            
            try:
                cache_dir.mkdir(parents=True, exist_ok=True)
            except Exception:
                pass
            
            # Render với batch processing để giảm memory usage
            for first, last in ranges:
                # Chia range lớn thành batches nhỏ
                batches = _split_range_into_batches(first, last, max_batch_size)
                for batch_first, batch_last in batches:
                    logger.info(f"OCR: Render trang {batch_first}–{batch_last}/{last} (dpi={dpi}, format={image_format})")
                    batch_results = _render_and_save_batch(batch_first, batch_last, image_format, jpeg_quality, memory_optimize)
                    index_to_image_path.update(batch_results)
        else:
            # Render toàn bộ (không có pages filter) - CHIA THÀNH BATCHES để tránh ngốn RAM
            if total_pages is not None:
                logger.info(f"OCR: Chuyển PDF → ảnh (dpi={dpi}, format={image_format}): {total_pages} trang - render theo batch {max_batch_size} trang/batch")
                try:
                    cache_dir.mkdir(parents=True, exist_ok=True)
                except Exception:
                    pass
                # Chia toàn bộ PDF thành batches
                batches = _split_range_into_batches(1, total_pages, max_batch_size)
                for batch_first, batch_last in batches:
                    logger.info(f"OCR: Render batch {batch_first}–{batch_last}/{total_pages}")
                    batch_results = _render_and_save_batch(batch_first, batch_last, image_format, jpeg_quality, memory_optimize)
                    index_to_image_path.update(batch_results)
            else:
                # Không biết total_pages, phải render hết (vẫn cố gắng dùng thread_count=1 để giảm memory)
                logger.info(f"OCR: Chuyển PDF → ảnh (dpi={dpi}, format={image_format}): không biết số trang, render toàn bộ")
                try:
                    cache_dir.mkdir(parents=True, exist_ok=True)
                except Exception:
                    pass
                if poppler_path and isinstance(poppler_path, str) and poppler_path.strip():
                    all_imgs = convert_from_path(pdf_path, dpi=dpi, poppler_path=poppler_path, thread_count=1)
                else:
                    all_imgs = convert_from_path(pdf_path, dpi=dpi, thread_count=1)
                
                # Save với format tối ưu và memory management
                ext = ".jpg" if image_format == "jpeg" else ".png"
                for idx, img in enumerate(all_imgs, start=1):
                    out_path = cache_dir / f"page_{idx:04d}{ext}"
                    try:
                        if image_format == "jpeg":
                            if img.mode in ('RGBA', 'LA', 'P'):
                                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                                if img.mode == 'P':
                                    img = img.convert('RGBA')
                                rgb_img.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                                img = rgb_img
                            img.save(str(out_path), format='JPEG', quality=jpeg_quality, optimize=True)
                        else:
                            img.save(str(out_path), format='PNG', optimize=True)
                        index_to_image_path[idx] = out_path
                        if memory_optimize:
                            del img
                            if idx % 5 == 0:
                                gc.collect()
                    except Exception as e:
                        logger.warning(f"Không thể lưu ảnh cache {out_path}: {e}")
                if memory_optimize:
                    del all_imgs
                    gc.collect()

    # Tạo danh sách ảnh theo thứ tự trang để OCR
    if pages_to_ocr is not None:
        # Chỉ OCR các trang được chỉ định (và có sẵn ảnh)
        ordered_indices = sorted([idx for idx in pages_to_ocr if idx in index_to_image_path])
    elif total_pages is None:
        # Không biết tổng trang: dùng thứ tự theo index hiện có
        ordered_indices = sorted(index_to_image_path.keys())
    else:
        ordered_indices = list(range(1, total_pages + 1))

    # Auto-detect language/variant nếu cần (chỉ một lần cho toàn bộ PDF)
    resolved_lang = None
    raw_lang = ocr_cfg.get("lang", "vie")
    # Normalize trước để check "auto" và "chi"
    normalized_lang = _normalize_lang_code(raw_lang)
    
    # Chỉ detect Chinese variant nếu lang="CN" hoặc "chi" (không có auto-detect)
    needs_chinese_variant_detection = ("chi" in normalized_lang.lower() and 
                                       "chi_sim" not in normalized_lang and 
                                       "chi_tra" not in normalized_lang)
    
    if needs_chinese_variant_detection and ordered_indices:
        # Chinese variant detection: chỉ cần 1 trang đầu để detect giản thể/phồn thể
        first_page_idx = ordered_indices[0]
        first_page_path = index_to_image_path.get(first_page_idx)
        if first_page_path:
            try:
                logger.info(f"Đang nhận biết Chinese variant (giản thể/phồn thể) từ trang đầu (lang config: {raw_lang})...")
                with Image.open(str(first_page_path)) as sample_img:
                    resolved_lang = _resolve_language(raw_lang, ocr_cfg, sample_img=sample_img)
                logger.info(f"Đã detect Chinese variant: {resolved_lang}")
            except Exception as e:
                logger.warning(f"Không thể detect Chinese variant từ trang đầu: {e}. Dùng mặc định chi_sim.")
                resolved_lang = _resolve_language(raw_lang, ocr_cfg, sample_img=None)
        else:
            resolved_lang = _resolve_language(raw_lang, ocr_cfg, sample_img=None)
    else:
        # Không cần detect variant, chỉ normalize
        resolved_lang = _resolve_language(raw_lang, ocr_cfg, sample_img=None)
    
    texts: List[str] = []
    total = len(ordered_indices)
    logger.info(f"OCR: Tổng số trang cần xử lý: {total}")
    show_progress = bool(ocr_cfg.get("show_progress", True))
    progress_interval = float(ocr_cfg.get("progress_log_interval_seconds", 60))
    if show_progress and tqdm is not None and total > 1:
        start_ts = time.time()
        with tqdm(total=total, desc="OCR PDF", unit="trang") as pbar:
            for i, page_idx in enumerate(ordered_indices, start=1):
                p = index_to_image_path.get(page_idx)
                if p is None:
                    logger.warning(f"Thiếu ảnh cho trang {page_idx}, bỏ qua")
                    pbar.update(1)
                    continue
                try:
                    # Thử mở ảnh với LOAD_TRUNCATED_IMAGES để xử lý ảnh bị truncated
                    try:
                        with Image.open(str(p)) as img:
                            # Thử load full image với LOAD_TRUNCATED_IMAGES nếu bị truncated
                            img.load()  # Load toàn bộ image data
                            text = _image_to_text(img, ocr_cfg, lang_override=resolved_lang)
                    except Exception as load_error:
                        # Nếu vẫn lỗi, thử với LOAD_TRUNCATED_IMAGES flag
                        if "truncated" in str(load_error).lower():
                            logger.warning(f"Ảnh trang {page_idx} bị truncated, thử load với LOAD_TRUNCATED_IMAGES...")
                            with Image.open(str(p)) as img:
                                # Pillow tự động xử lý truncated images nếu có thể
                                try:
                                    # Thử verify=False để bỏ qua một số checks
                                    img.verify()
                                    img = Image.open(str(p))  # Reopen sau verify
                                    text = _image_to_text(img, ocr_cfg, lang_override=resolved_lang)
                                except Exception:
                                    # Nếu vẫn lỗi, thử render lại từ PDF
                                    logger.warning(f"Không thể load ảnh truncated trang {page_idx}, thử render lại từ PDF...")
                                    try:
                                        # Xóa file cache bị lỗi
                                        if p.exists():
                                            try:
                                                p.unlink()
                                                logger.debug(f"Đã xóa file cache bị lỗi: {p}")
                                            except Exception:
                                                pass
                                        
                                        # Render lại từ PDF (single page - dùng format tối ưu)
                                        if poppler_path and isinstance(poppler_path, str) and poppler_path.strip():
                                            imgs = convert_from_path(pdf_path, dpi=dpi, poppler_path=poppler_path, 
                                                                    first_page=page_idx, last_page=page_idx, thread_count=1)
                                        else:
                                            imgs = convert_from_path(pdf_path, dpi=dpi, first_page=page_idx, last_page=page_idx, thread_count=1)
                                        
                                        if imgs and len(imgs) > 0:
                                            img = imgs[0]
                                            # Lưu lại vào cache với format tối ưu
                                            try:
                                                cache_dir.mkdir(parents=True, exist_ok=True)
                                                if image_format == "jpeg":
                                                    if img.mode in ('RGBA', 'LA', 'P'):
                                                        rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                                                        if img.mode == 'P':
                                                            img = img.convert('RGBA')
                                                        rgb_img.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                                                        img = rgb_img
                                                    img.save(str(p), format='JPEG', quality=jpeg_quality, optimize=True)
                                                else:
                                                    img.save(str(p), format='PNG', optimize=True)
                                                logger.info(f"Đã render lại và lưu cache cho trang {page_idx}")
                                            except Exception as save_err:
                                                logger.debug(f"Không thể lưu cache lại: {save_err}")
                                            
                                            # OCR lại (trước khi giải phóng memory)
                                            text = _image_to_text(img, ocr_cfg, lang_override=resolved_lang)
                                            
                                            # Giải phóng memory sau khi OCR xong
                                            if memory_optimize:
                                                del img, imgs
                                                gc.collect()
                                        else:
                                            logger.warning(f"Không thể render lại trang {page_idx} từ PDF")
                                            text = ""
                                    except Exception as render_error:
                                        logger.warning(f"Không thể render lại trang {page_idx}: {render_error}")
                                        text = ""
                        else:
                            raise load_error  # Nếu không phải truncated error, re-raise
                except Exception as e:
                    logger.warning(f"Không thể mở/OCR ảnh cho trang {page_idx} ({p}): {e}")
                    text = ""
                texts.append(text)
                elapsed = time.time() - start_ts
                avg = elapsed / i if i > 0 else 0.0
                remaining = max(total - i, 0) * avg
                pbar.set_postfix(avg_s_per_page=f"{avg:.2f}", eta=f"{remaining:.0f}s")
                pbar.update(1)
    else:
        start_ts = time.time()
        last_log = start_ts
        for i, page_idx in enumerate(ordered_indices, start=1):
            p = index_to_image_path.get(page_idx)
            if p is None:
                logger.warning(f"Thiếu ảnh cho trang {page_idx}, bỏ qua")
                continue
            try:
                # Thử mở ảnh với xử lý truncated images
                try:
                    with Image.open(str(p)) as img:
                        img.load()  # Load toàn bộ image data
                        texts.append(_image_to_text(img, ocr_cfg, lang_override=resolved_lang))
                except Exception as load_error:
                    # Nếu bị truncated, thử các cách khắc phục
                    if "truncated" in str(load_error).lower():
                        logger.warning(f"Ảnh trang {page_idx} bị truncated, thử load với LOAD_TRUNCATED_IMAGES...")
                        try:
                            with Image.open(str(p)) as img:
                                img.verify()
                                img = Image.open(str(p))  # Reopen sau verify
                                texts.append(_image_to_text(img, ocr_cfg, lang_override=resolved_lang))
                        except Exception:
                            # Thử render lại từ PDF
                            logger.warning(f"Không thể load ảnh truncated trang {page_idx}, thử render lại từ PDF...")
                            try:
                                # Xóa file cache bị lỗi
                                if p.exists():
                                    try:
                                        p.unlink()
                                    except Exception:
                                        pass
                                
                                # Render lại từ PDF (single page - dùng format tối ưu)
                                if poppler_path and isinstance(poppler_path, str) and poppler_path.strip():
                                    imgs = convert_from_path(pdf_path, dpi=dpi, poppler_path=poppler_path, 
                                                            first_page=page_idx, last_page=page_idx, thread_count=1)
                                else:
                                    imgs = convert_from_path(pdf_path, dpi=dpi, first_page=page_idx, last_page=page_idx, thread_count=1)
                                
                                if imgs and len(imgs) > 0:
                                    img = imgs[0]
                                    # Lưu lại vào cache với format tối ưu
                                    try:
                                        cache_dir.mkdir(parents=True, exist_ok=True)
                                        if image_format == "jpeg":
                                            if img.mode in ('RGBA', 'LA', 'P'):
                                                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                                                if img.mode == 'P':
                                                    img = img.convert('RGBA')
                                                rgb_img.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                                                img = rgb_img
                                            img.save(str(p), format='JPEG', quality=jpeg_quality, optimize=True)
                                        else:
                                            img.save(str(p), format='PNG', optimize=True)
                                    except Exception:
                                        pass
                                
                                    # OCR lại (trước khi giải phóng memory)
                                    texts.append(_image_to_text(img, ocr_cfg, lang_override=resolved_lang))
                                    
                                    # Giải phóng memory sau khi OCR xong
                                    if memory_optimize:
                                        del img, imgs
                                        gc.collect()
                                else:
                                    logger.warning(f"Không thể render lại trang {page_idx} từ PDF")
                                    texts.append("")
                            except Exception as render_error:
                                logger.warning(f"Không thể render lại trang {page_idx}: {render_error}")
                                texts.append("")
                    else:
                        raise load_error
            except Exception as e:
                logger.warning(f"Không thể mở/OCR ảnh cho trang {page_idx} ({p}): {e}")
                texts.append("")
            now = time.time()
            if now - last_log >= max(5.0, progress_interval):  # báo cáo định kỳ
                elapsed = now - start_ts
                avg = elapsed / i if i > 0 else 0.0
                remaining = max(total - i, 0) * avg
                logger.info(f"OCR: Trang {i}/{total} • TB {avg:.2f}s/trang • ETA ~{remaining:.0f}s")
                last_log = now
        elapsed = time.time() - start_ts
        avg = elapsed / total if total > 0 else 0.0
        logger.info(f"OCR: Hoàn tất {total} trang • TB {avg:.2f}s/trang")
    # Trả về số trang đã thực sự OCR (không phải total_pages)
    pages_processed = len(ordered_indices)
    return ("\n\n".join(texts), pages_processed)
def ocr_file(input_path: str, config_path: str = "config/config.yaml", pages: Optional[List[int]] = None, output_path: Optional[str] = None, skip_steps: Optional[dict] = None, process_mode: str = "process") -> str:
    """
    Extract text từ file PDF hoặc ảnh.
    Tự động detect PDF scan vs text-based để tối ưu.
    
    Args:
        output_path: Đường dẫn file output (để tạo tên file tạm thời)
        skip_steps: Dict với keys 'ocr', 'cleanup', 'spell_check' để skip các bước đã hoàn tất
        process_mode: "fast" = convert trực tiếp PDF→DOCX (chỉ cho text-based), "process" = extract→cleanup→spell check
    """
    _ensure_logger_config()
    pipeline_start_time = time.time()
    total_pages_processed = 0
    cleanup_stats = {"success": 0, "failed": 0}
    spell_check_stats = {"success": 0, "failed": 0}
    extracted_tables = {}
    
    if skip_steps is None:
        skip_steps = {}
    
    if not os.path.exists(input_path):
        raise FileNotFoundError(input_path)
    
    ocr_cfg = _detect_bundled_binaries(load_ocr_config(config_path))
    # Apply environment overrides for tables
    try:
        tables_override = os.environ.get("OCR_TABLES_RECONSTRUCT")
        if tables_override is not None:
            ocr_cfg.setdefault("tables", {})
            ocr_cfg["tables"]["reconstruct"] = tables_override == "1"
    except Exception:
        pass
    _ensure_dependencies(ocr_cfg)
    
    ext = os.path.splitext(input_path)[-1].lower()
    
    # Xử lý PDF
    if ext == ".pdf":
        auto_detect = ocr_cfg.get("auto_detect_pdf_type", True)
        
        if auto_detect:
            logger.info(f"Đang phát hiện loại PDF: {input_path}")
            pdf_type = detect_pdf_type(input_path, ocr_cfg)
            logger.info(f"PDF type: {pdf_type}")
            
            if pdf_type == "text":
                # Mode "fast": Convert trực tiếp PDF → DOCX (chỉ khi output_path là DOCX)
                if process_mode == "fast" and output_path and output_path.endswith(".docx"):
                    logger.info("📄 PDF text-based + Mode 'fast' → Convert trực tiếp PDF → DOCX...")
                    try:
                        # Đảm bảo dependencies đã được load
                        _ensure_dependencies(ocr_cfg)
                        convert_pdf_to_docx(input_path, output_path, pages)
                        
                        # Post-fix: loại bỏ tab thừa đầu dòng do soft-wrap
                        try:
                            _fix_docx_leading_tabs_and_soft_wraps(output_path)
                        except Exception:
                            pass
                        
                        # Return dict với text rỗng (đã convert trực tiếp)
                        return {
                            "text": "",
                            "cleanup_failed": 0,
                            "cleanup_failed_indices": [],
                            "cleanup_original_chunks": [],
                            "cleanup_all_chunks": [],
                            "spell_check_failed": 0,
                            "spell_check_failed_indices": [],
                            "spell_check_original_chunks": [],
                            "spell_check_all_chunks": [],
                            "ocr_cfg": ocr_cfg,
                            "direct_converted": True,  # Flag để web app biết đã convert trực tiếp
                        }
                    except Exception as e:
                        logger.warning(f"Convert trực tiếp thất bại: {e}, fallback về mode 'process'")
                        import traceback
                        logger.debug(traceback.format_exc())
                        # Fallback về mode process
                        process_mode = "process"
                
                # Mode "process": Extract text → cleanup → spell check
                logger.info("PDF có text layer → Extract text trực tiếp (nhanh)")
                text = extract_text_from_pdf(input_path, ocr_cfg, pages)
                # Đếm số trang đã xử lý
                if pages:
                    # Validate và đếm số trang hợp lệ
                    try:
                        if pdfplumber is not None:
                            with pdfplumber.open(input_path) as pdf:
                                total = len(pdf.pages)
                                valid_pages = [p for p in pages if 1 <= p <= total]
                                total_pages_processed = len(valid_pages)
                        elif PyPDF2 is not None:
                            with open(input_path, 'rb') as f:
                                reader = PyPDF2.PdfReader(f)
                                total = len(reader.pages)
                                valid_pages = [p for p in pages if 1 <= p <= total]
                                total_pages_processed = len(valid_pages)
                        else:
                            total_pages_processed = len(pages)  # Fallback
                    except Exception:
                        total_pages_processed = len(pages)  # Fallback
                else:
                    try:
                        if pdfplumber is not None:
                            with pdfplumber.open(input_path) as pdf:
                                total_pages_processed = len(pdf.pages)
                        elif PyPDF2 is not None:
                            with open(input_path, 'rb') as f:
                                reader = PyPDF2.PdfReader(f)
                                total_pages_processed = len(reader.pages)
                    except Exception:
                        total_pages_processed = 0
            else:
                logger.info("📷 PDF scan → Sử dụng OCR")
                text, total_pages_processed = ocr_pdf(input_path, config_path, pages)
                # After OCR, try table reconstruction if enabled
                try:
                    tables_cfg = ocr_cfg.get("tables", {})
                    if tables_cfg.get("reconstruct", False) and output_path:
                        logger.info("🗂️  Bắt đầu extract bảng từ PDF scan...")
                        
                        # Strategy: Ưu tiên unstructured.io (95-98% chính xác)
                        # Fallback về ocrmypdf + pdfplumber, cuối cùng là OpenCV + pytesseract advanced
                        extracted_tables = {}
                        table_mode = tables_cfg.get("mode", "auto")
                        
                        # Strategy 1: unstructured.io (nếu có và mode cho phép)
                        if table_mode in ("auto", "unstructured"):
                            try:
                                extracted_tables = _extract_tables_with_unstructured(input_path, output_path, ocr_cfg, pages)
                                if extracted_tables:
                                    logger.info(f"✅ unstructured.io: Đã extract {len(extracted_tables)} bảng từ PDF")
                            except Exception as e:
                                logger.debug(f"unstructured.io không khả dụng hoặc thất bại: {e}")
                                if table_mode == "unstructured":
                                    # Nếu user chỉ định unstructured nhưng không có → báo lỗi
                                    logger.warning("⚠️  unstructured.io không khả dụng. Cài bằng: pip install unstructured[pdf]")
                        
                        # Strategy 2: ocrmypdf + pdfplumber (fallback nếu unstructured không có kết quả)
                        if not extracted_tables and table_mode in ("auto", "ocrmypdf_then_extract"):
                            try:
                                extracted_tables = _try_extract_tables_from_pdf_via_ocrmypdf(input_path, output_path, ocr_cfg, pages)
                                if extracted_tables:
                                    logger.info(f"✅ ocrmypdf+pdfplumber: Đã extract {len(extracted_tables)} bảng từ PDF")
                            except Exception as e:
                                logger.debug(f"ocrmypdf+pdfplumber thất bại: {e}")
                        
                        # Strategy 3: OpenCV + pytesseract advanced (fallback cuối cùng)
                        if not extracted_tables and table_mode in ("auto", "opencv_grid", "pytesseract_advanced"):
                            try:
                                # Thử dùng pytesseract advanced nếu có sklearn
                                try:
                                    from sklearn.cluster import DBSCAN
                                    # Sử dụng hàm advanced pytesseract trong _extract_tables_from_images_cv
                                    extracted_tables = _extract_tables_from_images_cv(input_path, output_path, ocr_cfg, pages)
                                    if extracted_tables:
                                        logger.info(f"✅ OpenCV+pytesseract advanced: Đã extract {len(extracted_tables)} bảng từ PDF")
                                except ImportError:
                                    # Fallback về OpenCV grid detection cũ
                                    extracted_tables = _extract_tables_from_images_cv(input_path, output_path, ocr_cfg, pages)
                                    if extracted_tables:
                                        logger.info(f"✅ OpenCV grid: Đã extract {len(extracted_tables)} bảng từ PDF")
                            except Exception as e:
                                logger.debug(f"OpenCV fallback thất bại: {e}")
                        
                        if extracted_tables:
                            logger.info(f"✅ Tổng cộng đã extract {len(extracted_tables)} bảng từ PDF")
                    else:
                        logger.debug(f"Table extraction disabled: reconstruct={tables_cfg.get('reconstruct', False)}, output_path={output_path}")
                except Exception as e:
                    logger.warning(f"⚠️  Lỗi khi extract bảng: {e}")
                    import traceback
                    logger.debug(traceback.format_exc())
        else:
            # Force OCR nếu auto_detect = false
            text, total_pages_processed = ocr_pdf(input_path, config_path, pages)
    # Xử lý ảnh
    elif ext in {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}:
        text = ocr_image(input_path, config_path)
        total_pages_processed = 1  # Một ảnh = 1 trang
    else:
        raise ValueError(f"Unsupported input format for OCR: {ext}")
    
    # Lưu file sau bước OCR nếu chưa skip và có output_path
    if not skip_steps.get("ocr", False) and output_path:
        ocred_path = _get_intermediate_file_path(output_path, "_ocred.txt")
        try:
            with open(ocred_path, "w", encoding="utf-8") as f:
                f.write(text)
            logger.info(f"💾 Đã lưu kết quả OCR: {ocred_path}")
        except Exception as e:
            logger.warning(f"Không thể lưu file OCR: {e}")
    
    # Áp dụng AI cleanup nếu enabled
    cleanup_cfg = ocr_cfg.get("ai_cleanup", {})
    cleanup_failed = 0
    cleanup_failed_indices = []
    cleanup_original_chunks = []
    
    if cleanup_cfg.get("enabled", False) and not skip_steps.get("cleanup", False):
        result = ai_cleanup_text(text, ocr_cfg)
        if isinstance(result, tuple):
            text, cleanup_failed_indices, cleanup_original_chunks = result
            cleanup_failed = len(cleanup_failed_indices)
        else:
            text = result
        
        # Lưu file sau bước cleanup nếu có output_path
        if output_path:
            cleanup_path = _get_intermediate_file_path(output_path, "_cleanup.txt")
            try:
                with open(cleanup_path, "w", encoding="utf-8") as f:
                    f.write(text)
                logger.info(f"💾 Đã lưu kết quả Cleanup: {cleanup_path}")
            except Exception as e:
                logger.warning(f"Không thể lưu file Cleanup: {e}")
    elif skip_steps.get("cleanup", False):
        logger.info("⏭️  Bỏ qua bước Cleanup (đã có file từ phiên trước)")
    
    # Áp dụng AI spell check và paragraph restoration nếu enabled
    spell_check_cfg = ocr_cfg.get("ai_spell_check", {})
    spell_check_failed = 0
    spell_check_failed_indices = []
    spell_check_original_chunks = []
    
    if spell_check_cfg.get("enabled", False) and not skip_steps.get("spell_check", False):
        result = ai_spell_check_and_paragraph_restore(text, ocr_cfg)
        if isinstance(result, tuple):
            text, spell_check_failed_indices, spell_check_original_chunks = result
            spell_check_failed = len(spell_check_failed_indices)
        else:
            text = result
    elif skip_steps.get("spell_check", False):
        logger.info("⏭️  Bỏ qua bước Spell Check (đã có file từ phiên trước)")
    
    # Log tổng kết
    total_time = time.time() - pipeline_start_time
    hours = int(total_time // 3600)
    minutes = int((total_time % 3600) // 60)
    seconds = int(total_time % 60)
    time_str = f"{hours:02d}:{minutes:02d}:{seconds:02d}" if hours > 0 else f"{minutes:02d}:{seconds:02d}"
    
    logger.info("=" * 80)
    logger.info("📊 TỔNG KẾT OCR PIPELINE")
    logger.info(f"⏱️  Tổng thời gian: {time_str} ({total_time:.2f} giây)")
    logger.info(f"📄 Số trang đã OCR: {total_pages_processed}")
    if cleanup_cfg.get("enabled", False):
        if cleanup_failed > 0:
            logger.info(f"🧹 AI Cleanup: {cleanup_failed} chunks failed (đã lưu nội dung gốc)")
        else:
            logger.info(f"🧹 AI Cleanup: Hoàn tất không có lỗi")
    if spell_check_cfg.get("enabled", False):
        if spell_check_failed > 0:
            logger.info(f"✅ AI Spell Check: {spell_check_failed} chunks failed (đã lưu nội dung gốc)")
        else:
            logger.info(f"✅ AI Spell Check: Hoàn tất không có lỗi")
    logger.info("=" * 80)
    
    # Lưu lại cấu trúc chunks để có thể merge lại sau retry
    cleanup_all_chunks = cleanup_original_chunks if cleanup_original_chunks else []
    spell_check_all_chunks = spell_check_original_chunks if spell_check_original_chunks else []
    
    # Xử lý merge hàng bị cắt trong bảng sau spell check (nếu có)
    if extracted_tables and spell_check_cfg.get("enabled", False):
        logger.info("🔧 Đang merge các hàng bị cắt trong bảng...")
        extracted_tables = _merge_split_table_rows(extracted_tables, ocr_cfg)
    
    # Trả về text và thông tin failures để menu xử lý
    return {
        "text": text,
        "cleanup_failed": cleanup_failed,
        "cleanup_failed_indices": cleanup_failed_indices,
        "cleanup_original_chunks": cleanup_original_chunks,
        "cleanup_all_chunks": cleanup_all_chunks,  # Tất cả chunks (để merge lại)
        "spell_check_failed": spell_check_failed,
        "spell_check_failed_indices": spell_check_failed_indices,
        "spell_check_original_chunks": spell_check_original_chunks,
        "spell_check_all_chunks": spell_check_all_chunks,  # Tất cả chunks (để merge lại)
        "extracted_tables": extracted_tables,  # Bảng đã extract và xử lý
        "ocr_cfg": ocr_cfg
    }


def _create_html_from_items(all_items_with_position: List[dict], output_path: str) -> str:
    """
    Tạo file HTML từ all_items_with_position (text + images).
    Images được embed dưới dạng base64 để không cần temp files.
    
    Args:
        all_items_with_position: List các items (text hoặc image) đã được sort theo (Y, X)
        output_path: Đường dẫn file DOCX output (để tạo HTML temp file cùng folder)
    
    Returns:
        str: Đường dẫn file HTML đã tạo
    """
    import base64
    import html
    
    html_path = output_path.replace('.docx', '_temp.html')
    if html_path == output_path:  # Nếu không phải .docx
        html_path = output_path + '_temp.html'
    
    logger.info(f"📄 Tạo HTML trung gian: {html_path}")
    
    html_parts = ['<!DOCTYPE html>\n<html>\n<head>\n<meta charset="UTF-8">\n']
    html_parts.append('<style>\n')
    html_parts.append('body { font-family: Arial, sans-serif; margin: 20px; line-height: 1.6; }\n')
    html_parts.append('p { margin-bottom: 6pt; }\n')
    html_parts.append('img { max-width: 6in; height: auto; margin: 6pt 0; display: block; }\n')
    html_parts.append('</style>\n</head>\n<body>\n')
    
    images_count = 0
    text_count = 0
    
    for item in all_items_with_position:
        if item["type"] == "text":
            content = item["content"]
            # Escape HTML và thay thế line breaks
            content = html.escape(content)
            content = content.replace('\n\n', '</p><p>')
            content = content.replace('\n', '<br>')
            html_parts.append(f'<p>{content}</p>\n')
            text_count += 1
        elif item["type"] == "image":
            img_info = item["img_info"]
            image_bytes = img_info.get("data")
            if image_bytes and len(image_bytes) >= 10:
                image_ext = img_info.get("ext", "png").lower()
                if image_ext == "jpg":
                    image_ext = "jpeg"
                
                # Convert image bytes to base64
                try:
                    base64_data = base64.b64encode(image_bytes).decode('utf-8')
                    data_uri = f"data:image/{image_ext};base64,{base64_data}"
                    
                    # Get image dimensions for sizing
                    img_width = img_info.get("width", 0)
                    img_height = img_info.get("height", 0)
                    
                    if img_width > 0 and img_height > 0:
                        max_width_px = 576  # 6 inches at 96 DPI
                        if img_width > max_width_px:
                            aspect_ratio = img_height / img_width
                            display_width = max_width_px
                            display_height = int(max_width_px * aspect_ratio)
                        else:
                            display_width = img_width
                            display_height = img_height
                        
                        html_parts.append(f'<img src="{data_uri}" width="{display_width}" height="{display_height}" alt="Image from page {item["page_num"]}">\n')
                    else:
                        html_parts.append(f'<img src="{data_uri}" alt="Image from page {item["page_num"]}">\n')
                    
                    images_count += 1
                    logger.debug(f"✅ Đã embed image {images_count} vào HTML (trang {item['page_num']}, size: {len(image_bytes)} bytes)")
                except Exception as e:
                    logger.warning(f"⚠️ Không thể convert image từ trang {item['page_num']} sang base64: {e}")
            else:
                logger.warning(f"⚠️ Image từ trang {item['page_num']} không có data hợp lệ (size: {len(image_bytes) if image_bytes else 0} bytes)")
    
    html_parts.append('</body>\n</html>')
    
    html_content = ''.join(html_parts)
    
    try:
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        logger.info(f"✅ Đã tạo HTML: {text_count} paragraphs, {images_count} images")
        return html_path
    except Exception as e:
        logger.error(f"❌ Không thể tạo HTML file: {e}")
        raise


def _convert_html_to_docx_with_pandoc(html_path: str, output_path: str, ocr_cfg: dict) -> bool:
    """
    Convert HTML sang DOCX bằng pandoc.
    
    Args:
        html_path: Đường dẫn file HTML input
        output_path: Đường dẫn file DOCX output
        ocr_cfg: Config dictionary
    
    Returns:
        bool: True nếu thành công, False nếu thất bại
    """
    try:
        import pypandoc
    except ImportError:
        logger.warning("⚠️ pypandoc chưa được cài đặt. Cài pypandoc để dùng HTML intermediate workflow.")
        return False
    
    try:
        logger.info(f"🔄 Đang convert HTML → DOCX bằng pandoc...")
        # Pandoc options để preserve images và formatting
        extra_args = [
            '--standalone',
            '--wrap=none',  # Không wrap lines
        ]
        
        pypandoc.convert_file(
            html_path,
            'docx',
            outputfile=output_path,
            extra_args=extra_args
        )
        
        logger.info(f"✅ Đã convert HTML → DOCX thành công bằng pandoc")
        return True
    except Exception as e:
        logger.warning(f"⚠️ Pandoc conversion thất bại: {e}")
        import traceback
        logger.debug(traceback.format_exc())
        return False
if __name__ == "__main__":
    import argparse
    import threading

    parser = argparse.ArgumentParser(description="OCR scan file (PDF/image) and extract text")
    parser.add_argument("input", help="Path to image or PDF file (scan)")
    parser.add_argument("--config", default="config/config.yaml", help="Path to YAML config")
    parser.add_argument("--output", help="Save recognized text to file")
    parser.add_argument("--pages", help="Chỉ định các trang cần OCR. Ví dụ: '1,2,5,7' hoặc '1-7' hoặc '1-3,5,7-9'")
    parser.add_argument("--format", choices=["txt", "docx"], default=None, help="Định dạng file output (txt hoặc docx). Nếu không chỉ định, sẽ tự động detect từ extension của --output")
    args = parser.parse_args()

    _ensure_logger_config()
    
    # Xác định output_path và output_format
    if args.output:
        output_path = args.output
        # Detect format từ extension nếu --format không được chỉ định
        if args.format:
            output_format = args.format
        else:
            ext = os.path.splitext(output_path)[1].lower()
            if ext == ".docx":
                output_format = "docx"
            else:
                output_format = "txt"
    else:
        # Tạo output_path mặc định từ input_path
        input_dir = os.path.dirname(args.input) if os.path.dirname(args.input) else "."
        input_basename = os.path.basename(args.input)
        input_name_without_ext = os.path.splitext(input_basename)[0]
        if args.format == "docx":
            output_path = os.path.join(input_dir, input_name_without_ext + "_ocr_result.docx")
            output_format = "docx"
        else:
            output_path = os.path.join(input_dir, input_name_without_ext + "_ocr_result.txt")
            output_format = "txt"
    
    # Check & Resume: Kiểm tra file từ phiên trước
    existing_files = _check_existing_files(output_path)
    skip_steps = {}
    
    if existing_files["all_exist"]:
        logger.info("\n" + "=" * 80)
        logger.info("🔍 PHÁT HIỆN FILE TỪ PHIÊN LÀM VIỆC TRƯỚC")
        logger.info("=" * 80)
        
        found_files = []
        if existing_files["ocred"]:
            found_files.append(f"  • File OCR: {existing_files['ocred']}")
        if existing_files["cleanup"]:
            found_files.append(f"  • File Cleanup: {existing_files['cleanup']}")
        if existing_files["output"]:
            found_files.append(f"  • File Output: {existing_files['output']}")
        
        if found_files:
            logger.info("\nCác file đã phát hiện:")
            for f in found_files:
                logger.info(f)
            logger.info("")
            
            logger.info("Bạn có muốn sử dụng các file này để tiếp tục?")
            logger.info("")
            logger.info("  1. Có, tiếp tục từ Cleanup")
            logger.info("  2. Có, tiếp tục từ Spell Check")
            logger.info("  3. Có, nhưng chạy lại toàn bộ (OCR + Cleanup + Spell Check)")
            logger.info("  4. Không, kết thúc tác vụ")
            logger.info("")
            
            while True:
                try:
                    choice = input("Nhập lựa chọn (1/2/3/4): ").strip()
                    if choice == "1":
                        # Resume từ Cleanup (cần file _ocred.txt)
                        if existing_files["ocred"]:
                            skip_steps["ocr"] = True
                            logger.info("⏭️  Sẽ bỏ qua OCR, tiếp tục từ Cleanup")
                        else:
                            logger.warning("⚠️  Không tìm thấy file OCR (_ocred.txt). Không thể tiếp tục từ Cleanup.")
                            logger.info("🔄 Sẽ chạy lại toàn bộ quy trình")
                            skip_steps = {}
                        break
                    elif choice == "2":
                        # Resume từ Spell Check (cần file _cleanup.txt)
                        if existing_files["cleanup"]:
                            skip_steps["ocr"] = True
                            skip_steps["cleanup"] = True
                            logger.info("⏭️  Sẽ bỏ qua OCR và Cleanup, chỉ chạy Spell Check")
                        else:
                            logger.warning("⚠️  Không tìm thấy file Cleanup (_cleanup.txt). Không thể tiếp tục từ Spell Check.")
                            if existing_files["ocred"]:
                                logger.info("💡 Phát hiện file OCR. Sẽ tiếp tục từ Cleanup.")
                                skip_steps["ocr"] = True
                            else:
                                logger.info("🔄 Sẽ chạy lại toàn bộ quy trình")
                                skip_steps = {}
                        break
                    elif choice == "3":
                        # Chạy lại toàn bộ
                        skip_steps = {}
                        logger.info("🔄 Sẽ chạy lại toàn bộ quy trình")
                        break
                    elif choice == "4":
                        logger.info("Kết thúc tác vụ.")
                        sys.exit(0)
                    else:
                        logger.warning("Vui lòng nhập 1, 2, 3 hoặc 4.")
                except (KeyboardInterrupt, EOFError):
                    logger.info("\nKết thúc tác vụ.")
                    sys.exit(0)
    
    logger.info("Bắt đầu OCR pipeline...")
    
    # Load config
    ocr_cfg = _detect_bundled_binaries(load_ocr_config(args.config))
    
    # Khởi tạo initial_text để có thể được set bởi OCRmyPDF workflow (PDF scan)
    initial_text = None
    
    # Parse pages nếu có
    pages_list = None
    if args.pages:
        pages_list = _parse_pages(args.pages)
        if pages_list:
            logger.info(f"Chỉ xử lý các trang: {pages_list}")
        else:
            logger.warning(f"Không parse được pages từ '{args.pages}'. Sẽ xử lý tất cả trang.")
    
    # Step 1: Nhận biết file PDF là scan hay text-based (TRƯỚC KHI gọi _ensure_dependencies để tránh treo)
    input_path_lower = args.input.lower()
    is_pdf = input_path_lower.endswith(".pdf")
    
    if is_pdf:
        logger.info(f"Đang nhận biết loại PDF: {args.input}")
        logger.debug(f"pdfplumber available: {pdfplumber is not None}")
        logger.debug(f"PyPDF2 available: {PyPDF2 is not None}")
        logger.debug("Bắt đầu gọi detect_pdf_type...")
        pdf_type = detect_pdf_type(args.input, ocr_cfg)
        logger.debug("Đã hoàn thành detect_pdf_type")
        logger.info(f"PDF type: {pdf_type}")
        
        # 2.2. Nếu là text-based → convert trực tiếp tùy output format
        if pdf_type == "text":
            if output_format == "docx":
                # 2.2.2. Convert PDF → DOCX → lưu file → hoàn thành (không cleanup/spell check)
                # Gọi _ensure_dependencies TRƯỚC KHI convert để có dependencies cần thiết
                _ensure_dependencies(ocr_cfg)
                logger.info("📄 PDF text-based + DOCX output → Convert trực tiếp PDF → DOCX...")
                try:
                    convert_pdf_to_docx(args.input, output_path, pages_list)
                    # Post-fix: loại bỏ tab thừa đầu dòng do soft-wrap (an toàn, không merge para)
                    try:
                        _fix_docx_leading_tabs_and_soft_wraps(output_path)
                    except Exception:
                        pass
                    logger.info(f"✅ Đã tạo DOCX thành công: {output_path}")
                    logger.info("Hoàn tất OCR pipeline (PDF text-based → DOCX không cần cleanup/spell check).")
                    sys.exit(0)
                except Exception as e:
                    logger.error(f"❌ Lỗi khi convert PDF → DOCX: {e}")
                    error_msg = str(e).lower()
                    
                    # Check nếu là lỗi compatibility với PyMuPDF
                    if "get_area" in error_msg or "rect" in error_msg:
                        logger.warning("⚠️  Phát hiện lỗi tương thích giữa pdf2docx và PyMuPDF.")
                        logger.warning("💡 Giải pháp: Cài PyMuPDF==1.26.4 hoặc thấp hơn:")
                        logger.warning("   pip install PyMuPDF==1.26.4")
                    
                    # Fallback: Thử dùng OCRmyPDF để tạo PDF searchable, rồi extract text và tạo DOCX
                    if ocrmypdf_available:
                        logger.info("🔄 Thử fallback workflow với OCRmyPDF...")
                        try:
                            # Step 1: Dùng OCRmyPDF tạo PDF searchable
                            temp_searchable_pdf = os.path.splitext(output_path)[0] + "_searchable.pdf"
                            logger.info("📄 Bước 1: Dùng OCRmyPDF tạo PDF searchable...")
                            convert_pdf_with_ocrmypdf(args.input, temp_searchable_pdf, ocr_cfg, pages_list)
                            
                            # Step 2: Extract text từ PDF searchable
                            logger.info("📄 Bước 2: Extract text từ PDF searchable...")
                            extracted_text = extract_text_from_pdf(temp_searchable_pdf, ocr_cfg, pages_list)
                            
                            if not extracted_text or len(extracted_text.strip()) < 10:
                                raise RuntimeError("Extract text từ PDF searchable không thành công hoặc text quá ngắn.")
                            
                            # Step 3: Tạo DOCX từ extracted text và images
                            logger.info("📄 Bước 3: Tạo DOCX từ text đã extract...")
                            create_docx_from_processed_text(temp_searchable_pdf, output_path, extracted_text, ocr_cfg, pages_list)
                            
                            # Cleanup temp PDF searchable
                            try:
                                if os.path.exists(temp_searchable_pdf):
                                    os.remove(temp_searchable_pdf)
                                    logger.debug(f"🗑️  Đã xóa file temp: {temp_searchable_pdf}")
                            except Exception:
                                pass
                            
                            logger.info(f"✅ Đã tạo DOCX thành công qua OCRmyPDF fallback: {output_path}")
                            logger.info("Hoàn tất OCR pipeline (PDF text-based → OCRmyPDF → DOCX).")
                            sys.exit(0)
                            
                        except Exception as ocr_fallback_error:
                            error_msg_str = str(ocr_fallback_error)
                            logger.error(f"❌ OCRmyPDF fallback thất bại: {error_msg_str}")
                            
                            # Kiểm tra và thông báo về missing dependencies
                            if 'ghostscript' in error_msg_str.lower() or 'gswin64c' in error_msg_str.lower() or 'gs' in error_msg_str.lower():
                                logger.warning("⚠️  OCRmyPDF cần Ghostscript nhưng không tìm thấy.")
                                logger.warning("💡 Hướng dẫn cài đặt Ghostscript:")
                                logger.warning("   Windows: choco install ghostscript")
                                logger.warning("   Hoặc tải từ: https://www.ghostscript.com/download/gsdnld.html")
                                logger.warning("   Sau khi cài, thêm Ghostscript vào PATH và khởi động lại terminal.")
                            elif 'tesseract' in error_msg_str.lower():
                                logger.warning("⚠️  OCRmyPDF cần Tesseract OCR nhưng không tìm thấy.")
                                logger.warning("💡 Hướng dẫn cài đặt Tesseract:")
                                logger.warning("   Windows: choco install tesseract")
                                logger.warning("   Hoặc tải từ: https://github.com/UB-Mannheim/tesseract/wiki")
                            
                            logger.warning("🔄 Fallback về workflow TXT (Extract text → Cleanup → Spell Check)...")
                            output_format = "txt"
                            output_path = os.path.splitext(output_path)[0] + ".txt"
                            # Fall through để convert sang TXT rồi cleanup/spell check
                    else:
                        logger.warning("⚠️  OCRmyPDF không khả dụng. Không thể dùng fallback workflow.")
                        logger.warning("🔄 Fallback về workflow TXT (Extract text → Cleanup → Spell Check)...")
                        output_format = "txt"
                        output_path = os.path.splitext(output_path)[0] + ".txt"
                        # Fall through để convert sang TXT rồi cleanup/spell check
            elif output_format == "txt":
                # 2.2.1. Convert PDF → TXT → cleanup → spell check
                logger.info("📄 PDF text-based + TXT output → Extract text → Cleanup → Spell Check...")
                # Fall through để chạy standard workflow (extract text → cleanup → spell check)
        else:
            # 2.1. Nếu là scan → thử dùng OCRmyPDF trước, nếu không có thì dùng pytesseract
            logger.info("📷 PDF scan → Xử lý OCR...")
            
            # Ưu tiên dùng OCRmyPDF nếu khả dụng (tạo PDF searchable → extract text nhanh hơn)
            if ocrmypdf_available:
                logger.info("🔍 OCRmyPDF khả dụng → Dùng OCRmyPDF để tạo PDF searchable...")
                try:
                    # Step 1: Dùng OCRmyPDF tạo PDF searchable
                    temp_searchable_pdf = os.path.splitext(output_path)[0] + "_searchable.pdf"
                    logger.info("📄 Bước 1: OCRmyPDF đang tạo PDF searchable...")
                    convert_pdf_with_ocrmypdf(args.input, temp_searchable_pdf, ocr_cfg, pages_list)
                    
                    # Step 2: Extract text từ PDF searchable (nhanh hơn OCR từng ảnh)
                    logger.info("📄 Bước 2: Extract text từ PDF searchable...")
                    extracted_text = extract_text_from_pdf(temp_searchable_pdf, ocr_cfg, pages_list)
                    
                    if not extracted_text or len(extracted_text.strip()) < 10:
                        raise RuntimeError("Extract text từ PDF searchable không thành công hoặc text quá ngắn.")
                    
                    # Lưu kết quả OCR vào initial_text để tiếp tục workflow
                    initial_text = extracted_text
                    
                    # Cleanup temp PDF searchable (giữ lại nếu user muốn dùng sau)
                    cleanup_temp_pdf = ocr_cfg.get("cleanup_temp_searchable_pdf", True)
                    if cleanup_temp_pdf:
                        try:
                            if os.path.exists(temp_searchable_pdf):
                                os.remove(temp_searchable_pdf)
                                logger.debug(f"🗑️  Đã xóa file temp: {temp_searchable_pdf}")
                        except Exception:
                            pass
                    else:
                        logger.info(f"💾 Giữ lại PDF searchable: {temp_searchable_pdf}")
                    
                    logger.info("✅ Đã tạo text từ PDF searchable bằng OCRmyPDF. Tiếp tục Cleanup & Spell Check...")
                    # Skip OCR step vì đã có text từ OCRmyPDF
                    skip_steps["ocr"] = True
                    
                except Exception as ocr_fallback_error:
                    error_msg_str = str(ocr_fallback_error)
                    logger.warning(f"⚠️  OCRmyPDF thất bại: {error_msg_str}")
                    logger.warning("🔄 Fallback về pytesseract OCR workflow (tiêu chuẩn)...")
                    
                    # Kiểm tra và thông báo về missing dependencies
                    if 'ghostscript' in error_msg_str.lower() or 'gswin64c' in error_msg_str.lower() or 'gs' in error_msg_str.lower():
                        logger.warning("⚠️  OCRmyPDF cần Ghostscript nhưng không tìm thấy.")
                        logger.warning("💡 Hướng dẫn: choco install ghostscript")
                    elif 'tesseract' in error_msg_str.lower():
                        logger.warning("⚠️  OCRmyPDF cần Tesseract OCR nhưng không tìm thấy.")
                        logger.warning("💡 Hướng dẫn: choco install tesseract")
                    
                    # Fall through để chạy standard OCR workflow (pytesseract)
                    initial_text = None
            else:
                logger.info("⚠️  OCRmyPDF không khả dụng → Dùng pytesseract OCR workflow (tiêu chuẩn)...")
                # Fall through để chạy standard OCR workflow (pytesseract)
                initial_text = None
            
            # Fall through để chạy standard workflow (OCR → cleanup → spell check)
    else:
        # Không phải PDF → workflow hiện tại
        logger.info("Xử lý file không phải PDF...")
        # Fall through để chạy standard workflow
    
    # Load file từ phiên trước nếu resume
    # Note: initial_text có thể đã được set bởi OCRmyPDF workflow cho PDF scan
    # Nếu chưa có initial_text từ OCRmyPDF, kiểm tra resume files
    if initial_text is None and skip_steps.get("ocr", False):
        if skip_steps.get("cleanup", False):
            # Resume từ Spell Check → load file Cleanup
            if existing_files["cleanup"]:
                initial_text = _load_resume_file(existing_files["cleanup"], "Cleanup")
            else:
                logger.error("❌ Không tìm thấy file Cleanup để resume từ Spell Check!")
                logger.info("🔄 Sẽ chạy lại toàn bộ quy trình")
                skip_steps = {}
        else:
            # Resume từ Cleanup → load file OCR
            if existing_files["ocred"]:
                initial_text = _load_resume_file(existing_files["ocred"], "OCR")
            else:
                logger.error("❌ Không tìm thấy file OCR để resume từ Cleanup!")
                logger.info("🔄 Sẽ chạy lại toàn bộ quy trình")
                skip_steps = {}
    
    # Chạy pipeline với skip_steps
    if initial_text:
        # Resume từ file đã có → chỉ cần chạy các bước còn lại
        ocr_cfg = _detect_bundled_binaries(load_ocr_config(args.config))
        _ensure_dependencies(ocr_cfg)
        text = initial_text
        
        cleanup_cfg = ocr_cfg.get("ai_cleanup", {})
        cleanup_failed = 0
        cleanup_failed_indices = []
        cleanup_original_chunks = []
        
        # Chạy cleanup nếu cần (không skip)
        if cleanup_cfg.get("enabled", False) and not skip_steps.get("cleanup", False):
            result = ai_cleanup_text(text, ocr_cfg)
            if isinstance(result, tuple):
                text, cleanup_failed_indices, cleanup_original_chunks = result
                cleanup_failed = len(cleanup_failed_indices)
            else:
                text = result
            
            # Lưu file sau cleanup
            cleanup_path = _get_intermediate_file_path(output_path, "_cleanup.txt")
            try:
                with open(cleanup_path, "w", encoding="utf-8") as f:
                    f.write(text)
                logger.info(f"💾 Đã lưu kết quả Cleanup: {cleanup_path}")
            except Exception as e:
                logger.warning(f"Không thể lưu file Cleanup: {e}")
        
        # Chạy spell check nếu cần
        spell_check_cfg = ocr_cfg.get("ai_spell_check", {})
        spell_check_failed = 0
        spell_check_failed_indices = []
        spell_check_original_chunks = []
        
        if spell_check_cfg.get("enabled", False):
            result = ai_spell_check_and_paragraph_restore(text, ocr_cfg)
            if isinstance(result, tuple):
                text, spell_check_failed_indices, spell_check_original_chunks = result
                spell_check_failed = len(spell_check_failed_indices)
            else:
                text = result
        
        # Tạo result dict giống format của ocr_file
        result = {
            "text": text,
            "cleanup_failed": cleanup_failed,
            "cleanup_failed_indices": cleanup_failed_indices,
            "cleanup_original_chunks": cleanup_original_chunks,
            "cleanup_all_chunks": cleanup_original_chunks if cleanup_original_chunks else [],
            "spell_check_failed": spell_check_failed,
            "spell_check_failed_indices": spell_check_failed_indices,
            "spell_check_original_chunks": spell_check_original_chunks,
            "spell_check_all_chunks": spell_check_original_chunks if spell_check_original_chunks else [],
            "ocr_cfg": ocr_cfg
        }
    else:
        # Chạy toàn bộ pipeline (OCR → Cleanup → Spell Check)
        # Workflow thống nhất cho cả TXT và DOCX
        result = ocr_file(args.input, config_path=args.config, pages=pages_list, output_path=output_path, skip_steps=skip_steps)
    
    # Xử lý kết quả từ ocr_file
    result_text = result["text"]
    cleanup_failed = result["cleanup_failed"]
    spell_check_failed = result["spell_check_failed"]
    ocr_cfg = result["ocr_cfg"]
    
    # PDF text-based + DOCX đã được xử lý ở trên (convert trực tiếp)
    # Không cần xử lý thêm ở đây vì đã exit(0)
    
    # Nếu còn ở đây → là PDF scan + DOCX hoặc PDF text-based + TXT
    # Cả 2 trường hợp đều xuất TXT (không hỗ trợ DOCX cho scan)
    if output_format == "docx" and is_pdf:
        # PDF scan không thể extract images → fallback về TXT
        logger.warning("⚠️  PDF scan không thể extract images để tạo DOCX. Chỉ có thể xuất text.")
        logger.warning("🔄 Fallback về định dạng TXT...")
        output_format = "txt"
        output_path = os.path.splitext(output_path)[0] + ".txt"
    
    # Hiển thị menu completion (cho TXT hoặc sau khi DOCX failed)
    if output_path:
        logger.info(f"\n📁 File output: {output_path}")
        user_choice = _show_completion_menu(cleanup_failed, spell_check_failed, output_path)
        
        if user_choice == "retry":
            # Retry các chunk failed và merge lại (logic giống phần else)
            cleanup_cfg = ocr_cfg.get("ai_cleanup", {})
            spell_check_cfg = ocr_cfg.get("ai_spell_check", {})
            updated_text = result_text
            
            if result["cleanup_failed"] > 0 and cleanup_cfg.get("enabled", False) and result["cleanup_all_chunks"]:
                logger.info(f"Đang retry {result['cleanup_failed']} chunks AI Cleanup failed...")
                api_keys = cleanup_cfg.get("api_keys", [])
                if not api_keys:
                    api_keys = ocr_cfg.get("_root_api_keys", [])
                model_name = cleanup_cfg.get("model", "gemini-2.5-flash")
                prompt = """Bạn là một AI chuyên dọn dẹp văn bản OCR/scan. Nhiệm vụ:
1. Loại bỏ header/footer lặp lại ở mỗi trang
2. Loại bỏ các ký tự rác, vệt đen vô nghĩa từ quá trình scan
3. Loại bỏ số trang, watermark
4. Giữ nguyên nội dung chính của văn bản
5. Chuẩn hóa khoảng trắng thừa
6. Giữ nguyên định dạng đoạn văn

Trả về chỉ văn bản đã được dọn dẹp, không giải thích thêm.

Văn bản cần dọn dẹp:
"""
                retry_results, still_failed = _retry_failed_chunks_cleanup(
                    result["cleanup_failed_indices"],
                    result["cleanup_all_chunks"],
                    api_keys,
                    model_name,
                    prompt,
                    ocr_cfg
                )
                
                all_chunks = list(result["cleanup_all_chunks"])
                for idx, retry_text in retry_results.items():
                    if idx < len(all_chunks):
                        all_chunks[idx] = retry_text
                updated_text = "\n\n".join(all_chunks)
                logger.info(f"AI Cleanup Retry: {result['cleanup_failed'] - len(still_failed)}/{result['cleanup_failed']} chunks retry thành công.")
            
            if result["spell_check_failed"] > 0 and spell_check_cfg.get("enabled", False):
                logger.info(f"Đang retry {result['spell_check_failed']} chunks AI Spell Check failed...")
                api_keys = spell_check_cfg.get("api_keys", [])
                if not api_keys:
                    api_keys = ocr_cfg.get("_root_api_keys", [])
                model_name = spell_check_cfg.get("model", "gemini-2.5-flash")
                prompt = """Bạn là một AI chuyên soát lỗi chính tả và phục hồi cấu trúc văn bản OCR. Nhiệm vụ chính của bạn là PHÂN TÍCH NGỮ CẢNH và QUYẾT ĐỊNH THÔNG MINH.

=== NHIỆM VỤ CHÍNH: PHÂN TÍCH VÀ PHỤC HỒI CÂU BỊ NGẮT (Ưu tiên cao nhất) ===

Bạn cần ĐỌC KỸ NỘI DUNG và PHÂN TÍCH để phân biệt:

A. CÂU BỊ NGẮT DO CONVERT PDF → TXT (CẦN NỐI LẠI):
   - Đọc ngữ cảnh: Nếu dòng trước chưa hoàn thành ý và dòng sau tiếp nối ý đó → nối lại
   - Ví dụ: 
     * "Our client is also the owner of Vietnam Trade Mark Registration No. 315843 for "MICROBAN"
       in Class 5 covering..." 
     → Phân tích: "in Class 5" tiếp nối câu trước → NỐI LẠI thành một câu
   
   - Dấu hiệu cần nối:
     * Dòng trước không kết thúc bằng dấu câu (. ! ?) HOẶC kết thúc bằng dấu phẩy, hai chấm
     * Dòng sau bắt đầu bằng chữ thường (tiếp nối câu trước)
     * Nội dung dòng sau về mặt ngữ pháp và ngữ nghĩa là phần tiếp theo của câu trước
     * Đọc toàn bộ ngữ cảnh để hiểu rõ mối quan hệ

B. NGẮT PARAGRAPH CÓ CHỦ ĐÍCH (KHÔNG NỐI):
   - Đọc ngữ cảnh: Nếu dòng sau là ý mới, chủ đề mới, hoặc đoạn văn mới → KHÔNG nối
   - Ví dụ:
     * "...attached as Exhibit 1.
       
       Khách hàng của chúng tôi là chủ sở hữu..."
     → Phân tích: Đây là đoạn mới (chuyển từ tiếng Anh sang tiếng Việt) → KHÔNG NỐI
   
   - Dấu hiệu KHÔNG nối:
     * Dòng trước kết thúc bằng dấu chấm (. ! ?) và dòng sau bắt đầu bằng chữ hoa
     * Dòng sau là câu đầu tiên của một đoạn mới (ý tưởng mới, chủ đề mới)
     * Có sự thay đổi rõ ràng về ngữ cảnh (ví dụ: chuyển từ phần này sang phần khác)
     * Đọc toàn bộ ngữ cảnh để xác định đây là ngắt đoạn có chủ đích

QUY TRÌNH PHÂN TÍCH:
1. ĐỌC toàn bộ văn bản để hiểu cấu trúc và ngữ cảnh
2. PHÂN TÍCH từng vị trí ngắt dòng:
   - Xem xét nội dung trước và sau dòng ngắt
   - Đánh giá mối quan hệ ngữ pháp và ngữ nghĩa
   - Xác định đây là câu bị ngắt hay ngắt đoạn có chủ đích
3. QUYẾT ĐỊNH:
   - Nếu là câu bị ngắt → NỐI lại (thay line break bằng space)
   - Nếu là ngắt đoạn có chủ đích → GIỮ NGUYÊN (có thể thêm dòng trống nếu cần)
4. ÁP DỤNG nhất quán cho toàn bộ văn bản

=== CÁC NHIỆM VỤ KHÁC ===

1. SOÁT LỖI CHÍNH TẢ:
   - Sửa các lỗi chính tả do OCR (ví dụ: "Kíng" → "Kính", "hang" → "hàng")
   - Sửa các lỗi chính tả thông thường
   - KHÔNG thay đổi từ ngữ chuyên ngành, tên riêng, địa danh
   - KHÔNG thay đổi số liệu, ngày tháng, địa chỉ
2. PHỤC HỒI CẤU TRÚC PARAGRAPH:
   - Sau khi đã nối các câu bị ngắt, xác định các ngắt đoạn hợp lý
   - Mỗi đoạn văn nên có một ý chính hoàn chỉnh
   - Giữ nguyên các dòng trống giữa các đoạn đã được xác định là có chủ đích
   - Đảm bảo các câu trong một đoạn có liên quan với nhau

3. BẢO VỆ TOÀN VẸN NỘI DUNG:
   - TUYỆT ĐỐI KHÔNG thay đổi ý nghĩa của văn bản
   - KHÔNG thêm, bớt, hoặc diễn giải lại nội dung
   - KHÔNG thay đổi thứ tự từ trong câu (chỉ nối lại khi cần)
   - GIỮ NGUYÊN định dạng đặc biệt (bullet points, numbered lists, bảng)
   - GIỮ NGUYÊN các từ viết hoa nếu chúng là tên riêng, thuật ngữ

4. ĐỊNH DẠNG:
   - Giữ nguyên định dạng văn bản song ngữ (nếu có)
   - Giữ nguyên các dấu câu quan trọng
   - Chuẩn hóa khoảng trắng thừa giữa các từ (nhưng không thay đổi paragraph breaks hợp lý)
   - Đảm bảo mỗi câu kết thúc bằng dấu câu thích hợp

=== NGUYÊN TẮC QUAN TRỌNG ===

- SỬ DỤNG SỨC MẠNH PHÂN TÍCH NGỮ CẢNH: Đọc và hiểu nội dung, không chỉ dựa vào quy tắc cú pháp
- QUYẾT ĐỊNH THÔNG MINH: Mỗi quyết định nối hay không nối phải dựa trên phân tích ngữ cảnh cụ thể
- NHẤT QUÁN: Áp dụng cùng một tiêu chuẩn phân tích cho toàn bộ văn bản
- BẢO TOÀN Ý NGHĨA: Chỉ điều chỉnh cấu trúc, KHÔNG thay đổi nội dung hoặc ý nghĩa

Trả về chỉ văn bản đã được soát và phục hồi, không giải thích thêm.

Văn bản cần phân tích và xử lý:
"""
                
                spell_check_chunks = [updated_text[i:i+spell_check_cfg.get("chunk_size", 10000)] for i in range(0, len(updated_text), spell_check_cfg.get("chunk_size", 10000))]
                retry_results, still_failed = _retry_failed_chunks_spell_check(
                    result["spell_check_failed_indices"],
                    spell_check_chunks,
                    api_keys,
                    model_name,
                    prompt,
                    ocr_cfg
                )
                
                spell_check_chunks_list = list(spell_check_chunks)
                for idx, retry_text in retry_results.items():
                    if idx < len(spell_check_chunks_list):
                        spell_check_chunks_list[idx] = retry_text
                updated_text = "\n\n".join(spell_check_chunks_list)
                logger.info(f"AI Spell Check Retry: {result['spell_check_failed'] - len(still_failed)}/{result['spell_check_failed']} chunks retry thành công.")
            
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(updated_text)
            logger.info(f"OCR: Đã lưu text đã được retry vào: {output_path}")
            # Cleanup intermediate files
            _cleanup_intermediate_files(output_path)
        elif user_choice == "save":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(result_text)
            logger.info(f"OCR: Đã lưu vào: {output_path}")
            # Cleanup intermediate files
            _cleanup_intermediate_files(output_path)
        elif user_choice == "exit":
            logger.info("Thoát không lưu.")
        else:
            # Auto-save
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(result_text)
            logger.info(f"OCR: Đã tự động lưu vào: {output_path}")
            # Cleanup intermediate files
            _cleanup_intermediate_files(output_path)
    
    logger.info("Hoàn tất OCR pipeline.")